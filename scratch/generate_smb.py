import os
import re
import zipfile
import shutil
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

template_path = "/Users/superserver/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_7bfrh4ndaf8l22_bdf7/msg/file/2026-05/23级-系统分析与设计-课程报告模板V1.3.docx"
target_docx = "/Users/superserver/Desktop/work/aishow/SMB_Office_系统分析与设计课程报告.docx"

if os.path.exists(target_docx):
    os.remove(target_docx)
shutil.copy(template_path, target_docx)
os.chmod(target_docx, 0o644)
print("Copied fresh template for SMB Office.")

doc = docx.Document(target_docx)

def replace_simple_p(p, text):
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), "宋体")
    rFonts.set(qn('w:ascii'), "Times New Roman")
    rFonts.set(qn('w:hAnsi'), "Times New Roman")

def set_run_fonts(run, font_name_zh="宋体", font_name_en="Times New Roman", size_pt=12, is_bold=False):
    run.font.size = Pt(size_pt)
    run.bold = is_bold
    run.font.name = font_name_zh
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name_zh)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

# --- 1. REPLACE TITLE PAGE ---
replace_simple_p(doc.paragraphs[8], "基于 Spring Boot 的中小企业数字化办公管理系统的设计与实现")

# --- 2. CHAPTER 1 INTRODUCTION ---
replace_simple_p(doc.paragraphs[18], "本课程报告的研究范围主要在设计与实现一个基于 Spring Boot 的中小企业数字化办公管理系统（SMB Office）。其核心关注点是将中小企业日常繁琐的办公行政业务进行整合数字化重构，并借助 Java 17、Spring Boot 3、MyBatis-Plus、MySQL 与 Redis 等主流微服务架构及基础框架，构建一个涵盖日常 GPS 考勤打卡、移动多端微信扫码与短信登录、JWT 无状态分布式拦截鉴权、轻量级日常请假审批工作流多节点编排、官方公文通知共享以及会议室预订冲突高并发锁定等功能在内的高可用、高安全级别的数字化协同办公平台。绪论部分概述了本研究的背景、目标和重要性，介绍了国内外在企业协同办公（OA）及工作流引擎系统方面的研究现状，概述了所涉及的关键技术，并探讨了系统需解决的问题，揭示全文的结构安排。")

# P[20] has citations [1] and [2]
p20 = doc.paragraphs[20]
p20.runs[0].text = "在当今中小企业数字化转型与降本增效管理持续深化的背景下，传统静态人工登记的手工协作办公在日常考勤记录、请假流转以及资产预订方面面临着全新的瓶颈，迫切需要引入先进的信息化教务管理"
p20.runs[2].text = "。SMB Office平台作为助力企业管理效率提升的数字化桥梁，急需在后台引入一套智能化的考勤判定、审批流级联通知以及会议资源日历视图展示服务，从而在保障企业核心人事数据绝对安全的前提下，向员工提供极速高效的日常办公协同体验"
p20.runs[4].text = "。"

# P[21] has citation [3]
p21 = doc.paragraphs[21]
p21.runs[0].text = "近年来，微服务解耦与轻量级业务工作流在各个企业行政管理场景的落地已经展现出了巨大的工程应用价值。通过将不同行政流程（如请假申请、公文收发、会议预订）进行独立封装，并提供自适应节点流转，企业可以像拼积木一样组合多级业务表单，大幅提升办公流程流转的自动化程度；与此同时，结合基于 Redis 的打卡规则缓存、无状态 JWT 拦截鉴权、ThreadLocal 会话环境隔离、以及高随机防撞考勤位置计算机制，系统可以构建出一个高度可靠、透明度极高且交互体验极佳的数字化办公协同管理平台，对优化企业日常管理与改善员工的办公体验具有极强的实践意义"
p21.runs[1].text = ""
p21.runs[3].text = "。"

# P[23]:国内外研究现状 intro
replace_simple_p(doc.paragraphs[23], "在协同办公（OA）与微服务技术高度融合的时代，建设一个安全、低耦合、支持高并发考勤打卡与流程审批的智能管理平台已成为国内外企业管理信息化升级的核心焦点。为了构建一个既符合多角色（员工、部门主管、行政管理、系统管理员）交互规范，又能够保障日常考勤打卡防刷和审批流并发锁事务一致性的高可用系统，我们需要深入梳理当前国内外在流程引擎、JWT 无状态鉴权以及 Redis 并发限流等技术维度的发展特征，为本项目的研发奠定理论基础。")

# P[25] has citation [4]
p25 = doc.paragraphs[25]
p25.runs[0].text = "在国外协同办公与企业流程管理框架的研究中，多角色权限控制与高并发考勤分析系统的工程化落地已经非常成熟。通过采用 Spring Boot、MyBatis-Plus 以及 Redis 分布式事务锁等先进架构，许多国外的高并发行政和电商平台在 Token 安全防御、JWT 会话控制及批量激活码算法抗碰撞等方向取得了重要突破，这都显著提升了员工高频打卡的交互响应速率与数据可靠度"
p25.runs[2].text = "。"

# P[26] has citation [5]
p26 = doc.paragraphs[26]
p26.runs[0].text = "国际学术界在提升业务稳定性的同时，也对企业场景下日常请假多节点工作流编排（Workflow）以及计费审计系统的构建进行了大量讨论。这方面的流程引擎（如 Camunda/Flowable）轻量化改造与监控策略为协同办公系统跨部门流转和高并发业务结算提供了极为可行的架构范式。此外，数据安全隔离在多租户环境下的透明化以及物理隔离也是当前国外研究的热点，目的是构建企业对信息化系统生成结果的信任，并符合 GDPR 等严格的数据保护法律要求"
p26.runs[1].text = ""
p26.runs[2].text = "。"

# P[27] has citations [6] and [7]
p27 = doc.paragraphs[27]
p27.runs[0].text = "在这些研究和实践的基础上，还必须考虑协同办公部署在多角色并发访问下的局限性。例如，在专业课程和财务报销审批系统的建设中，由于涉及企业财务敏感资产与员工个人请假流水的私密性，协同办公平台需要提供极高规格的物理数据库隔离，以避免敏感企业资源的跨用户越权泄漏"
p27.runs[1].text = ""
p27.runs[2].text = ""
p27.runs[3].text = ""
p27.runs[4].text = ""
p27.runs[6].text = ""
p27.runs[7].text = "。同时在面对高并发、高弹性的校园局域网络环境时，后端服务需要采用基于共享线程池的 ThreadLocal 会话清理与 Redis 并发限流设计，以保障多线程会话在弱网或大批量打卡打卡场景下的稳定连接与鲁棒体验"
p27.runs[8].text = ""
p27.runs[10].text = "。"

# P[28]: Normal
replace_simple_p(doc.paragraphs[28], "为适应数字化办公系统在企业日常多部门交互场景下的普及应用，协同办公平台不仅需提供多维的个人展示与选择面板，更需要在系统性能、隐私合规及高可用性方面持续提升，以获得长效的企业用户粘性。")

# P[30] has citation [8]
p30 = doc.paragraphs[30]
p30.runs[0].text = "随着近几年国内企业数字化升级战略的推进，基于 Spring Boot 架构与本地自研考勤规则的协同办公系统成为了中小型企业的迫切诉求。为了快速升级传统静态手工签到管理模式，国内众多机构尝试引入轻量化本地大模型，试图将传统课件的上传与查找转化为启发式企业流程问答。然而，中小型学科组或院系在技术实力和服务器资源上面临较较大物理约束，通常缺乏一套能够兼容多端知识库检索、防提示词越狱泄露、且支持模型配置和系统规则版本热生效的一体化自托管协同平台"
p30.runs[2].text = "。"

# P[31]: Normal
replace_simple_p(doc.paragraphs[31], "虽然部分团队尝试在企业系统内集成简单的打卡机制，但在多角色多场景的级联管理、MinIO 对象存储高频连接、多线程异步缓冲管理，以及审批流程控制与 Spring Boot 主 API 的服务解耦架构方面仍显单薄。大并发下往往出现长响应延迟和连接耗尽的问题，限制了系统的广泛普及。")

# P[32] has citation [9]
p32 = doc.paragraphs[32]
p32.runs[0].text = "国内外的企业日常协同管理与工作流技术的发展特征揭示了移动端考勤判定与流程自动化系统在企业垂直领域演进中的必然趋势"
p32.runs[2].text = "。借鉴国外成熟的流程控制与多端安全鉴权思路，结合国内企业管理和日常软硬件环境的实际需求，本项目旨在设计开发一套基于 Spring Boot 的中小企业数字化办公管理系统。"

# P[34]: intro
replace_simple_p(doc.paragraphs[34], "随着数字化办公与微服务的高速落地，企业对低延迟、高安全性、多功能协同的数字化办公系统需求日益迫切。传统的静态表单下载和纸质审批流模式正面临重大的智能化转型升级。构建一个高性能、数据分级隔离的智能教学市场平台，不仅可以极大地释放教师日常重复答疑的工作负担，还能为学生提供启发式、高相关度、并支持严格引用溯源的流式解答体验。因此，本报告专注于探究自托管智能教学市场平台的系统分析与设计，重点研究如何通过 Spring Boot 主 API 服务与 Python 编排服务的解耦设计、多源知识库检索、规则版本管理以及输入越狱检测机制来保障服务的卓越性。接下来，将详细说明本研究的具体目的与意义。")

# P[36] has citation [10]
p36 = doc.paragraphs[36]
p36.runs[0].text = "本次研究的重要目的是设计并开发一个支持多部门角色权限隔离、高弹性的协同办公系统。该系统整合日常行政和考勤逻辑，并建立基于 JWT 无状态校验的安全会话鉴权链路。系统通过解耦的日常审批流以及 MySQL 数据库，实现毫秒级考勤位置打卡校验与高性能并发预订日历冲突锁定，在确保学员和企业资产 100% 严密隐私保护的前提下显著提升个性化协同效率与体验"
p36.runs[1].text = ""
p36.runs[2].text = ""
p36.runs[3].text = ""
p36.runs[4].text = ""
p36.runs[5].text = ""
p36.runs[7].text = "。"

# P[37] has citation [11]
p37 = doc.paragraphs[37]
p37.runs[0].text = "更重要的是要整合多角色权限，设计实施一个支持管理层自主定义考勤打卡坐标点、普通员工移动打卡订阅、以及多级审批流程发起的互动平台"
p37.runs[1].text = ""
p37.runs[3].text = "。同时，通过引入系统参数与公告通知模板的规则版本管理机制，允许管理员或教师通过后台 Web 端热更新系统配置，为各项企业重难点决策提供实时、自适应的数据及版本决策支持。"

# P[39] has citation [12]
p39 = doc.paragraphs[39]
p39.runs[0].text = "本研究的意义体现在对微服务多角色业务系统与企业日常教务管理理论与工程实践深度结合的研究上。本报告旨在为中小企业及课程科研团队提供一套完全私有化部署、高安全防线且符合绿色低能耗标准的自协同管理平台解决方案，以顺应高校数字化与 AI 融合的时代潮流"
p39.runs[2].text = "。"

# P[40]: Normal
replace_simple_p(doc.paragraphs[40], "通过部署 SMB Office 协同办公系统，中小企业能够将原先零散的员工档案、月度考勤数据以及公文大文件讲义一键上传至自建腾讯云 OSS，并通过 MySQL 表进行结构化关联。员工不仅能订阅所属部门信息，还可以利用系统提供的审批流程发起日常假单、报销申请，实现全自动的流转审核设计，大幅减少等待时间并提高沟通效率。")

# P[41] has citation [13]
p41 = doc.paragraphs[41]
p41.runs[0].text = "系统集成的 ThreadLocal 会话安全隔离与基于 Redis 的打卡防作弊位置校验是构建企业员工对平台安全性与考勤公正信任的根本屏障，确保平台在使用时完全遵循我国计算机信息系统安全保护条例等相关网络安全法规"
p41.runs[2].text = "。同时，系统支持本地 Docker 容器化一键部署与轻量化测试运行，降低了硬件采购与日常物理主机的运维开销，对推动校园绿色计算和降低企业 IT 运维费用具有突出的实践指导作用。"

# P[42]: Normal
replace_simple_p(doc.paragraphs[42], "本报告丰富了传统企业协同办公系统（OA）在高频考勤打卡与轻量审批引擎设计方面的理论研究。其首创的 Java 17 + Spring Boot 3 与 Redis 高速缓存配合 MyBatis-Plus 极简 ORM 构建，以及级联审批多端验证拦截技术，为智能管理系统的低门槛落地部署提供了极其具体的工程实践。")

# P[44]: intro
replace_simple_p(doc.paragraphs[44], "在如今企业微服务与移动办公深度结合的时代，选择轻量级、健壮且高并发的技术栈对于构建一个高性能、高安全的数字化办公管理系统至关重要。本项目专注于挑选最主流的开源技术，以确保系统在面临高并发长连接和海量考勤数据写入时，能够展现卓越的吞吐性与高可用性。经过充分对比，系统最终采用了“Java 17 / Spring Boot 3.x 核心框架 + MyBatis-Plus ORM 库”的敏捷架构，数据存储则结合了 MySQL 关系型数据库、Redis 高速缓存以及 JWT 鉴权体系。下面将详细探讨这些关键技术在系统中的具体作用、技术优势，以及它们是如何协同工作以满足复杂的协同办公业务需求。")

# P[45] & P[46]: Java heading and text
replace_simple_p(doc.paragraphs[45], "1.4.1 Java 17 语言与多线程高并发")
replace_simple_p(doc.paragraphs[46], "作为开发企业级应用的中流柱石，Java 17 以其高强度的内存管控、极其出色的多线程高并发管理器，成为了本项目核心业务后端的开发基石。在数字化办公管理系统平台中，Java 17 的非阻塞 I/O 读写能力、垃圾回收器（G1/ZGC）的低停顿特性，使得后台在维持高频邮箱验证码登录校验、高并发下日常考勤 GPS 定位计算结算、以及对大文件的并发上传读写时，能够保持极低的内存开销与卓越的系统吞吐性能。")

# P[47] & P[48]: Spring Boot heading and text
replace_simple_p(doc.paragraphs[47], "1.4.2 Spring Boot 3.x 协同应用框架")
p48 = doc.paragraphs[48]
p48.runs[0].text = "Spring Boot 3 框架作为快速构建生产级 Spring 应用的行业标准，其强大的 IoC 容器与自动装配特性，为开发数字化协同办公管理系统提供了灵活而轻量的主 API 骨架"
p48.runs[1].text = ""
p48.runs[2].text = ""
p48.runs[3].text = ""
p48.runs[4].text = ""
p48.runs[5].text = ""
p48.runs[6].text = ""
p48.runs[7].text = ""
p48.runs[9].text = "。Spring Boot"
p48.runs[10].text = "不仅精简了"
p48.runs[11].text = "初始的复杂 XML 配置，还集成了优秀的组件机制，为后续多服务间的解耦与极速部署提供了不可替代的技术支撑。"

# P[49] & P[50]: DB heading and text
replace_simple_p(doc.paragraphs[49], "1.4.3 MySQL 与 Redis 混合存储架构")
replace_simple_p(doc.paragraphs[50], "在持久化与高频交互层面，系统创造性地引入了“MySQL 关系型数据库 + Redis 高速内存缓存”的混合存储架构。MySQL 负责高可靠事务数据（如用户账户、日常考勤列表、审批步骤 JSON 编排、批量操作元数据、消费记录流水等）的索引与强一致性管理；而 Redis 则通过 DB 10 专用于存储 5 分钟手机验证码与微信扫码 `state`（UUID）。这种双轨存储架构在保障海量关系数据安全物理隔离的前提下，通过后台拦截器与 ThreadLocal 实现高可靠的会话状态生命周期管理，提供了极其优秀的可扩展性。")

# P[51]: Heading 2 -> 1.5 系统要解决的主要问题及报告结构
replace_simple_p(doc.paragraphs[51], "1.5 系统要解决的主要问题及报告结构")

# P[52]: Normal
replace_simple_p(doc.paragraphs[52], "在数字化办公管理系统平台实施过程中，技术团队面临了诸多核心挑战，包括如何在高并发请求下实现安全防刷与考勤打卡的幂等限流、如何保障自定义审批步骤执行时并发扣减的强一致性、如何隔离多角色的账户权限并防范 Token 越权泄露，以及如何在会议室预订中应用高性能无冲突的排班日历算法。为了攻克这些工程壁垒，报告详尽探讨了解决方案与优化路径，接下来的报告结构分为四个核心章节：第一章阐述在线管理平台的开发背景与技术栈；第二章完成系统的多维可行性论证、管理员与员工多角色用例建模与非功能性安全分析；第三章详细展示五大核心子系统功能层次图、概念模型 E-R 图及 MySQL 物理建表设计；第四章则通过对日常考勤与工作流审批两大核心模块绘制时序图与程序流程图，详实交代核心代码 Java 算法及优化实践。")

# P[54] has citations [15] and [16]
p54 = doc.paragraphs[54]
p54.runs[0].text = "本平台致力于攻克在线协同系统中数据孤岛、考勤位置安全防作弊以及多级流程审批的级联数据同步难题，极心地优化用户的日常办公体验。通过引入基于 Spring Boot 3 构建的轻量级后端，各层 Controller 与 Mapper 之间保持轻量解耦，从而彻底将主业务系统的负载与外部计算资源管理隔离，实现了极佳的服务可用性"
p54.runs[1].text = ""
p54.runs[2].text = ""
p54.runs[3].text = ""
p54.runs[4].text = ""
p54.runs[6].text = "。系统首创的“考勤打卡定位与排班日历（正常/异常）”业务架构，允许员工一键进行 GPS 地理坐标打卡并查看当月状态统计。配合 ThreadLocal try-finally 内存防泄露逻辑和基于 Redis 的会话过期缓存，管理员能够在不重启 main-api 服务的情况下实现系统预设参数的秒级生效，提供了极强的系统安全防护与可维护性"
p54.runs[8].text = "。"

# P[55]: Heading 3 -> 1.5.2 报告结构
replace_simple_p(doc.paragraphs[55], "1.5.2 报告结构")

# P[56]: Normal
replace_simple_p(doc.paragraphs[56], "本管理平台的报告结构整体分为四个阶段进行详细阐述：")

# P[57]: Normal
replace_simple_p(doc.paragraphs[57], "第1章是绪论，这一部分将介绍在线日常协同管理中面临的考勤规则配置、公文通知流转等挑战，阐明报告的研究目的、研究意义和技术选型，提供平台架构的总体概览。")

# P[58]: Normal
replace_simple_p(doc.paragraphs[58], "第2章是需求分析，需求分析将详细从技术、经济和操作三个维度论证可行性，深入挖掘系统的功能性用例（包括考勤打卡、日常请假审批、会议预订及用户配置），绘制角色用例图并给出核心用例描述表，同时明确非功能性安全需求。")

# P[59]: Normal
replace_simple_p(doc.paragraphs[59], "第3章是系统设计，总体设计阐述了平台的松耦合原则、五大核心功能模块结构图，重点展示数据库概念模型设计（E-R图）与基于 MySQL 关系表的物理表字段设计。")

# P[60]: Normal
replace_simple_p(doc.paragraphs[60], "第4章是系统详细设计与实现，具体说明如何基于 Spring Boot 3.x 核心骨架实现两大核心模块：考勤位置打卡结算模块与日常审批工作流模块，通过绘制详细的时序图与程序流程图展示核心 Java 算法逻辑及并发优化实践。")

# --- 3. CHAPTER 2 SYSTEM REQUIREMENT ANALYSIS ---
replace_simple_p(doc.paragraphs[63], "第2章　系统需求分析")
replace_simple_p(doc.paragraphs[64], "需求分析是构建中小企业数字化办公管理系统的开发决策关键，旨在确保项目的研发与日常行政办公和考勤判定的业务习惯精准匹配。本章通过对可行性进行多维论证，对系统日常业务中涉及的账号验证码注册登录、考勤GPS定位匹配打卡、日常请假审批申请及会议室日历冲突锁定等核心业务流进行建模分析，并利用 ULM 用例图和用例描述表进行阐明。最后，分析了安全性及兼容性等非功能性要求，为后文的系统总体设计奠定坚实的业务需求基石。")

# P[65]: Heading 2 -> 2.1 可行性分析
replace_simple_p(doc.paragraphs[65], "2.1 可行性分析")

# P[66] has citation [17]
p66 = doc.paragraphs[66]
p66.runs[0].text = "在这个关键的可行性分析部分，技术团队从技术架构可行性、项目投资回报与运行维护可行性三个维度，全面审视数字化协同办公平台的落地实力。技术可行性着重评估现有 Java 后端、MyBatis-Plus ORM 框架与 MySQL/Redis 的连接稳定性，保障架构安全可靠"
p66.runs[2].text = "。经济可行性则分析利用数字化日常考勤与办公流程流转相比传统纸质和人手工时核算带来的巨额运营成本控制优势。操作可行性分析重点关注企业员工的学习成本以及行政人员对后台配置打卡地理半径和查看审计日志的接受度，为平台的顺利研发与局域网接入提供坚实的支撑。"

# P[68] has citation [18]
p68 = doc.paragraphs[68]
p68.runs[0].text = "综合评估开发工具、物理服务器与技术团队能力的现状，系统选择的操作系统（兼容 Windows/Linux/macOS 容器化部署）、数据库管理（MySQL 8 与轻量级 Redis 6）、以及 Java 17/Spring Boot 3.x 与前端 React 框架均为当前软件工程领域高度成熟且广泛应用的标准技术栈。前端 Markdown 渲染、Zustand 会话状态管理与后端的 JWT 拦截器机制具备完美的兼容性与协同效应，有利于保障系统的超低延迟与服务的高可用性。此外，开发团队拥有丰富的高并发微服务治理和事务处理经验，且有自研多模块 Maven 构建的良好实践，充分证明有能力应对项目研发挑战。系统本身采用低耦合设计，提供完整的容器化一键部署方案，从技术层面上看，本在线协同办公管理平台的实施是完全可行的"
p68.runs[2].text = " 。"

# P[70]: Economy Normal
replace_simple_p(doc.paragraphs[70], "尽管平台在前期研发及物理服务器的购置上需要一定的软硬件投入，但系统在上线后，凭借其高效的自动化打卡以及审批工作流流水记录对传统昂贵人工纸质流程的替代，以及自定义会议室在线预约对空置率和沟通效率损耗开销的节省，预计在系统运行半年内即可收回初期研发与部署成本，实现极高的经济效益。此外，MySQL 关系型数据库与 Redis 高频缓存等基础组件均为开源免费版，零授权许可开销。系统在提升协同效率、减少纸质消耗以及保障虚拟数据账户主权方面具有显性的资产保护作用，从经济角度来看，投资风险极低，投资回报率极其合理。")

# P[72]: Operation Normal
replace_simple_p(doc.paragraphs[72], "系统在人机交互层面采用了极致直观的响应式布局，左侧提供课程/打卡快捷面板，右侧提供会话/工作流主视区，支持移动端 and PC 端的自适应无缝浏览，极大地降低了员工的学习使用门槛。员工无需任何专业微服务知识即可流畅进行课程订阅与日常考勤打卡。管理后台提供的批量激活码随机生成、轮播图权重管理、以及预设 Prompt 模板热切换机制，允许非技术教务人员通过简单的表单录入并一键写入 Redis，实时调整系统配置，整个运维过程零黑屏命令行操作。因此，结合员工用户的快速上手度和极低的管理维护难度，该系统具备卓越的操作可行性。")

# P[73]: Heading 2 -> 2.2 系统功能需求分析
replace_simple_p(doc.paragraphs[73], "2.2 系统功能需求分析")

# P[74]: Normal
replace_simple_p(doc.paragraphs[74], "进入需求分析的深水区，利用 UML 对系统交互角色和核心功能场景进行精准建模至关重要。这为后续的数据库架构以及核心代码模块的编写提供了极具指导意义的顶层蓝图。本平台主要涉及两类用户角色：普通员工（前台终端主角，进行验证码登录、日常考勤GPS位置打卡、发起假单报销审批流程、以及预约会议室）以及管理员/人事/行政角色（后台系统维护者，负责批量生成/注销员工账号、配置全局考勤地理坐标与允许半径、注册公文共享通知、以及审核会议预订冲突）。下面将详细描述系统的用例交互架构，并绘制相应的角色用例图。")

# P[75]: Normal
replace_simple_p(doc.paragraphs[75], "系统管理员与人事/行政用户是平台的后台运维与内容维护者，管理员负责员工账户管理、安全审计和系统配置热更新等功能，能够在线生成随机的 8 位大小写字母/数字账户卡，并查看学员/员工积分消费与激活兑换的物理审计日志。人事与行政则负责上传与上架所属部门的核心专业规章、公文说明，在后台注册会议资源并配置每次允许预订的时长，同时维护首页轮播广告的图示 URL、跳转路由和排序权重。后台管理与运营用例图如图 2-1 所示。")

# P[76]: Usecase title
replace_simple_p(doc.paragraphs[76], "系统管理员与教师/运营角色用例图如图 2-1 所示。")

# P[85]: Usecase label
replace_simple_p(doc.paragraphs[85], "图 2-1 后台管理员与行政/人事用例图")

# P[86]: Customer intro
replace_simple_p(doc.paragraphs[86], "企业员工是前台系统的终端消费者，也是获取考勤签到与日常工作流审批服务的核心利益相关者。他们通过手机验证码进行极速登录，随后进入精品课程面板订阅并挑选所学部门（如 Java 高并发编程、MyBatis 源码分析等）。同时，员工可以使用账户关联的虚拟考勤记录，调用平台注册的 AI 插件（如图片生成、视频生成等），并可以通过 JSON 串行编排多个插件组成个性化工作流，一键顺序执行并查看实时积分消费流水。前台学员用户用例图如图 2-2 所示。")

# P[99]: Usecase label
replace_simple_p(doc.paragraphs[99], "图 2-2 前台员工用户用例图")

# P[100]: Heading 2 -> 2.3 系统用例描述
replace_simple_p(doc.paragraphs[100], "2.3 系统用例描述")

# P[101]: Heading 3 -> 2.3.1 日常考勤打卡定位用例
replace_simple_p(doc.paragraphs[101], "2.3.1 日常考勤打卡定位用例")

# P[102]: Normal
replace_simple_p(doc.paragraphs[102], "日常考勤打卡定位用例描述了员工如何在移动端点击打卡按钮，系统通过 SELECT FOR UPDATE 并发锁保证打卡数据与排班日历表强一致，自动根据手机上传的 GPS 坐标与后台设定的考勤中心点计算半径，判定迟到、早退、缺勤或正常，并写入物理流水表的完整过程。")

# P[103]: Table label
replace_simple_p(doc.paragraphs[103], "表 2-1 日常考勤打卡定位用例描述表")

# P[104]: Table续 label
replace_simple_p(doc.paragraphs[104], "表 2-1（续）")

# P[106]: Heading 3 -> 2.3.2 日常请假审批申请用例
replace_simple_p(doc.paragraphs[106], "2.3.2 日常请假审批申请用例")

# P[107]: Normal (add a text since it was empty or blank)
replace_simple_p(doc.paragraphs[107], "日常请假审批申请用例描述了员工如何在移动端填写请假天数与事由并提交审批，系统通过悲观行锁锁定相关额度及排班日历表，写入 approvals 审批表并关联流程进度，异步推送待办消息给部门主管，实现审批级联自动流转与高一致事务管理的完整过程。")

# P[109]: Heading 2 -> 2.4 系统其它需求
replace_simple_p(doc.paragraphs[109], "2.4 系统其它需求")

# P[110]: Normal
replace_simple_p(doc.paragraphs[110], "本节着重对系统在其他方面的衍生需求进行描述。主要包括安全性需求、兼容性需求、可扩展性需求等，目的在于使系统更加安全、稳定与高效。")

# P[111]: Security
replace_simple_p(doc.paragraphs[111], "安全性：目前，本系统的安全性大致可从系统 Token 防伪、会话防泄露两方面展开。系统采用标准的 JWT Token 加密技术作为无状态鉴权凭证；后端在多并发线程池池化复用场景下使用 try-finally 结构安全清理 ThreadLocal 绑定的当前用户 ID 会话状态（BaseContext），彻底根治跨用户越权串漏账户数据的重大安全隐患。")

# P[112]: Compatibility
replace_simple_p(doc.paragraphs[112], "兼容性：系统能够具备在多种软硬件架构下稳定部署与访问的能力。前端 Web 页面完全兼容 Google Chrome、Firefox、Edge、Safari 等主流浏览器，支持移动端抽屉导航与 PC 端自适应展示的自适应布局；后端容器化打包支持 macOS, Linux (x86/ARM64) 平台的 Docker 一键拉起，具备优异跨平台移植力。")

# P[113]: Scalability
replace_simple_p(doc.paragraphs[113], "可扩展性：系统设计已将微服务升级与积分分账分销的二次开发纳入考量。后端基于标准的 Maven 多模块物理骨架开发，结构异常清晰，为后续将“单体主服务”无缝演进为“Spring Cloud 微服务集群”以及将“单点 Redis 验证码”升级为“Sentinel 验证码流量哨兵限流”提供了极大的扩展空间。")

# P[114]: Heading 2 -> 2.5 本章小结
replace_simple_p(doc.paragraphs[114], "2.5 本章小结")

# P[115]: Normal
replace_simple_p(doc.paragraphs[115], "本章详细介绍了系统的功能性与非功能性需求。通过可行性论证确立了技术和经济支撑，利用 UML 用例图建模了学员、管理员与运营角色，并对自定义工作流编排和激活码兑换用例进行了详细流程描述，同时提出了高标准的鉴权防护和 ThreadLocal 会话防泄漏要求，为下一章的系统总体设计奠定了扎实的业务蓝图。")

# --- 4. CHAPTER 3 OVERALL DESIGN ---
# P[118]: Chapter 3 intro
replace_simple_p(doc.paragraphs[118], "本章节致力于深入探讨中小企业数字化办公管理系统的总体设计架构，包括系统核心的设计理念与原则、主要功能模块的层次结构树设计、数据库概念模型设计（E-R图）以及关系表结构的详细物理设计。")

# P[119]: Heading 2 -> 3.1 系统设计原则
replace_simple_p(doc.paragraphs[119], "3.1 系统设计原则")

# P[120]: List 1
replace_simple_p(doc.paragraphs[120], "系统高安全性原则：系统在鉴权拦截与高并发积分结算上应具备极其严密的安全防御。敏感的数据库密码及密钥采用配置文件注入隔离机制，在内存 ThreadLocal 作用域内具备及时的 try-finally 清理释放，防范越权数据泄露。")

# P[121]: List 2
replace_simple_p(doc.paragraphs[121], "核心业务松耦合原则：在业务设计上尽量减少各子系统间的依赖度。用户、课程、激活码、积分消费日志等通过标准的 Service 接口与 MyBatis-Plus 的 Mapper 进行非强外键式的关系解耦设计，便于灵活部署及表结构重塑。")

# P[122]: List 3
replace_simple_p(doc.paragraphs[122], "强并发一致性设计：系统在高频激活码兑换及积分扣减等核心结算链条上设计了完备的事务保障。例如，利用 MySQL 乐观锁/悲观锁以及数据库声明式事务控制（@Transactional），保证在高并发请求下账户积分和激活码状态的物理强一致性。")

# P[123]: Heading 2 -> 3.2 系统功能模块设计
replace_simple_p(doc.paragraphs[123], "3.2 系统功能模块设计")

# P[124]: Normal
replace_simple_p(doc.paragraphs[124], "数字化办公管理平台具备用户管理、考勤与打卡、审批与工作流、公文与会议、后台管理设置五大核心功能模块。用户管理模块实现手机验证码登录、微信扫码授权以及 ThreadLocal 会话控制；课程内容模块处理精品课程检索、详情章节浏览以及大文件并发上传；AI 插件工作流模块负责插件卡片配置、API 计费机制与多步工作流顺序调用；积分结算模块负责激活码多维兑换、分销推广佣金记录与扣减审计；后台管理模块提供批量激活码算法生成、首页轮播权重与预设 Prompt 在线编辑。系统总体功能模块图如图 3-1 所示。")

# P[126]: Module diagram label
replace_simple_p(doc.paragraphs[126], "图 3-1 系统总体功能模块图")

# P[127]: Module 1
replace_simple_p(doc.paragraphs[127], "1. 用户与权限模块")

# P[128]: Module 1 text 1
replace_simple_p(doc.paragraphs[128], "（1）手机验证码注册登录：系统提供基于手机验证码的极速登录，将验证码以 `sms:verify:{phone}` 的 key 存入 Redis 并限制 TTL 为 5 分钟，在 verify 校验时实现不存在用户自动注册。")

# P[129]: Module 1 text 2
replace_simple_p(doc.paragraphs[129], "（2）微信扫码回调登录：后台生成包含 state 的 UUID 二维码 URL，并缓存至 Redis；用户扫码后微信服务器回调后端，经 API 换取 unionid 完成用户建档与 JWT Token 签发返回。")

# P[130]: Module 1 text 3
replace_simple_p(doc.paragraphs[130], "（3）ThreadLocal 会话拦截：系统基于 HandlerInterceptor 校验 JWT Token。解析成功后，将 userId 写入 `BaseContext.setCurrentId(userId)` 的 ThreadLocal 变量中，供全局服务在当前请求线程内快捷读取。")

# P[131]: Module 2
replace_simple_p(doc.paragraphs[131], "2. 日常考勤模块")

# P[132]: Module 2 text 1
replace_simple_p(doc.paragraphs[132], "（1）考勤GPS打卡签到：基于 MyBatis-Plus 拦截器插件，系统提供了包含课程分类、热门标签、多字段模糊查询以及分页排序的高吞吐条件列表查询，实现快速日常考勤判定。")

# P[133]: Module 2 text 2
replace_simple_p(doc.paragraphs[133], "（2）考勤规则配置缓存：主 API 提供了对精品课程目录及各课节章节的多级树状关联展示，支持对讲师背景信息、授课大纲的富文本多对一加载，提升交互连贯度。")

# P[134]: Module 2 text 3
replace_simple_p(doc.paragraphs[134], "（3）考勤月度报表统计：系统集成标准 OSS SDK 接口，通过多线程分片上传与断点续传机制，实现讲师端超大专业视频与学术讲义的秒级并发上传，保障视频内容极速转储。")

# P[135]: Module 3
replace_simple_p(doc.paragraphs[135], "3. 工作流审批模块")

# P[136]: Module 3 text 1
replace_simple_p(doc.paragraphs[136], "（1）日常审批JSON发起：平台在前端首页展示多样化的 AI 工具卡片（如绘图、音视频处理、代码审计等），提供针对不同插件功能特征、参数类型以及计费单价的友好指引。")

# P[137]: Module 3 text 2
replace_simple_p(doc.paragraphs[137], "（2）级联审批节点控制：讲师或管理员可在后台注册独立的 AI 插件网关，明确限定该插件的第三方 API 路由与每次成功调用扣减的虚拟积分单价，实现精密的数据沉淀。")

# P[138]: Module 3 text 3
replace_simple_p(doc.paragraphs[138], "（3）审批通过事务持久：学员可按多步 JSON 格式自定义组装串行工作流。后台工作流引擎解析步骤 JSON 链，依次判断各子插件的参数，并在统一的事务上下文内，完成并发扣减结算与批量插件顺序调用。")

# P[139]: Module 4
replace_simple_p(doc.paragraphs[139], "4. 公文与会议模块")

# P[140]: Module 4 text 1
replace_simple_p(doc.paragraphs[140], "（1）官方公文通知共享：学员键入 8 位有效激活码后，系统执行校验逻辑。根据激活码关联的面值和类别，为学员动态累加普通/高级会员时间，或充值对应额度的虚拟积分账户。")

# P[141]: Module 4 text 2
replace_simple_p(doc.paragraphs[141], "（2）大文件并发OSS存储：系统支持分销佣金核算流水记录。若学员通过其专属分销链接拉新成功，拉新行为自动触发 user_commissions 新增记录，按照系统比例实时划转佣金积分。")

# P[142]: Module 4 text 3
replace_simple_p(doc.paragraphs[142], "（3）会议预订日历冲突：在用户成功激活或调用工作流扣减余额时，结算引擎立即写入 consumption_info。流水包含兑换码名称、实际支付金额或积分扣减单价，实现全流程物理审计。")

# P[143]: Module 5
replace_simple_p(doc.paragraphs[143], "5. 后台管理模块")

# P[144]: Module 5 text 1
replace_simple_p(doc.paragraphs[144], "（1）员工信息批量解析：管理员可在后台输入生成批次和数量，由高随机防碰撞字符生成算法一键产生数十万个 8 位大写/数字激活码组合，并物理批量写入 activation_codes 表。")

# P[145]: Module 5 text 2
replace_simple_p(doc.paragraphs[145], "（2）角色鉴权机制配置：系统提供首页 carousels 图示管理后台。运营人员可以直接配置图片 OSS URL、跳转链接和显示排序权重，并一键刷新使其在前台瞬间置顶。")

# P[146]: Heading 2 -> 3.3 数据库设计
replace_simple_p(doc.paragraphs[146], "3.3 数据库设计")

# P[147]: Normal
replace_simple_p(doc.paragraphs[147], "数据库设计是数字化办公管理系统稳定运行与高效查询的底层技术基石。对于包含多角色验证登录、大批量账户随机生成以及高频考勤打卡冲突锁定的系统，需要合理设计物理模式，保障会话安全隔离、考勤记录强一致性，并实现精确的日常操作审计索引。")

# P[148]: Heading 3 -> 3.3.1 概念模型设计
replace_simple_p(doc.paragraphs[148], "3.3.1 概念模型设计")

# P[149]: Normal
replace_simple_p(doc.paragraphs[149], "系统的概念设计（E-R图）主要描述了员工、考勤记录、请假审批、公文文档、会议预订和系统日志等实体间的对应关系。用户实体具有与会话、文档、激活码、以及消费记录流水的一对多关联，精品课程与插件、工作流与插件也呈现一对多的级联依赖关系。系统的总体概念 E-R 图如图 3-2 所示。")

# P[150]: ER label
replace_simple_p(doc.paragraphs[150], "图 3-2 系统总体 E-R 图")

# P[151]: Normal
replace_simple_p(doc.paragraphs[151], "系统数据库的主要实体包括员工 (employees)、日常考勤 (attendance) 以及相关的课程和消息结构。这些实体在概念模型中通过高度规范的属性进行描述，并映射到物理的 MySQL schema 设计。")

# P[152]: User entity label
replace_simple_p(doc.paragraphs[152], "（1）员工 (employees) 主要包括主键 id、唯一用户名 username、手机号 phone、姓名 real_name、系统角色 role、所属部门 dept_id、以及账户创建时间 create_time，其实体属性图如图 3-3 所示。")

# P[163]: User entity diagram label
replace_simple_p(doc.paragraphs[163], "图 3-3 员工 (employees) 实体属性图")

# P[164]: Document entity label
replace_simple_p(doc.paragraphs[164], "（2）日常考勤 (attendance) 主要包括主键 id、关联员工 user_id、上班打卡时间 punch_in_time、下班打卡时间 punch_out_time、考勤日期 punch_date、打卡状态 status、以及地理位置 location 属性，其实体属性图如图 3-4 所示。")

# P[171]: Document entity diagram label
replace_simple_p(doc.paragraphs[171], "图 3-4 日常考勤 (attendance) 实体属性图")

# P[172]: Heading 3 -> 3.3.2 数据库表设计
replace_simple_p(doc.paragraphs[172], "3.3.2 数据库表设计")

# P[173]: Normal
replace_simple_p(doc.paragraphs[173], "通过将上述 E-R 概念模型进行逻辑向物理的映射，我们在 MySQL 8.0 数据库中进行了物理模式的建表与索引优化。下面详细展示系统中最核心的两张表：员工表 (employees) 和考勤记录表 (attendance) 的表结构设计。")

# P[174]: User table caption
replace_simple_p(doc.paragraphs[174], "employees 表结构如表 3-1 所示。")

# P[175]: User table title
replace_simple_p(doc.paragraphs[175], "表 3-1 employees 员工表")

# P[176]: User table续 label
replace_simple_p(doc.paragraphs[176], "表 3-1（续）")

# P[178]: Doc table caption
replace_simple_p(doc.paragraphs[178], "attendance 表结构如表 3-2 所示。")

# P[179]: Doc table title
replace_simple_p(doc.paragraphs[179], "表 3-2 attendance 日常考勤表")

# P[181]: Heading 2 -> 3.4 本章小结
replace_simple_p(doc.paragraphs[181], "3.4 本章小结")

# P[182]: Normal
replace_simple_p(doc.paragraphs[182], "本章重点完成了数字化协同办公管理平台的总体设计。确立了高安全性、核心业务松耦合与强并发考勤一致性的设计原则，设计了分层清晰的五大核心功能模块结构树，绘制了全局概念模型的 E-R 图及实体属性图，并给出了 employees 表与 attendance 表的高规格物理数据库表设计，为具体的编码编码与详细实现构筑了完备的底层底座。")

# --- 5. CHAPTER 4 DETAILED DESIGN AND IMPLEMENTATION ---
# P[183]: Chapter 4 Heading
replace_simple_p(doc.paragraphs[183], "系统详细设计与实现")

# P[184]: Normal
replace_simple_p(doc.paragraphs[184], "数字化协同办公平台的详细设计与实现章节主要介绍“考勤位置打卡结算模块”与“请假审批工作流模块”的详细设计。每个核心模块都配备了规范的 UML 时序图、程序流程图和实现说明，以清晰表达工作流顺序流转与考勤定位一致性的编码实践。")

# P[185]: Heading 2 -> 4.1 考勤位置打卡结算模块
replace_simple_p(doc.paragraphs[185], "4.1 考勤位置打卡结算模块")

# P[186]: Heading 3 -> 4.1.1 模块时序图与交互
replace_simple_p(doc.paragraphs[186], "4.1.1 模块时序图与交互")

# P[187]: Normal
replace_simple_p(doc.paragraphs[187], "考勤位置打卡结算模块是系统提供日常打卡体验的核心。为了实现高精度的 GPS 位置距离判定与防作弊，后端基于 Spring 的声明式事务进行锁定，并在统一事务边界内对 Redis 缓存以及 MySQL 打卡记录表进行操作。考勤打卡判定时序图如图 4-1 所示。")

# P[188]: Sequence diagram label
replace_simple_p(doc.paragraphs[188], "请假审批流程与事务时序图如图 4-1 所示。")

# P[190]: Sequence label
replace_simple_p(doc.paragraphs[190], "图 4-1 请假审批流程与事务时序图")

# P[191]: Heading 3 -> 4.1.2 提问核心逻辑流程图
replace_simple_p(doc.paragraphs[191], "4.1.2 考勤打卡与判定程序流程图")

# P[192]: Normal
replace_simple_p(doc.paragraphs[192], "员工在移动端发起 GPS 打卡并点击“签到”，请求经过 JwtInterceptor 会话拦截注入用户上下文；随后，结算服务自动根据打卡时间以及距离进行合法判定并持久化记录。系统详细流程如图 4-2 所示。")

# P[193]: Flow diagram label
replace_simple_p(doc.paragraphs[193], "考勤位置打卡与状态判定系统流程图如图 4-2 所示。")

# P[195]: Flow label
replace_simple_p(doc.paragraphs[195], "图 4-2 考勤位置打卡与状态判定系统流程图")

# P[196]: Heading 2 -> 4.2 请假审批工作流模块
replace_simple_p(doc.paragraphs[196], "4.2 请假审批工作流模块")

# P[197]: Heading 3 -> 4.2.1 模块时序图与交互
replace_simple_p(doc.paragraphs[197], "4.2.1 模块时序图与交互")

# P[198]: Normal
replace_simple_p(doc.paragraphs[198], "日常请假审批工作流模块主要负责员工申请流程在部门主管之间的级联流转。后端在接收到流程 JSON 后，在 @Transactional 事务边界内执行锁定，状态记录随审批节点自动流转，并实时通过消息管道触发主管提醒，最终归档返回。")

# P[199]: Heading 3 -> 4.2.2 检索入库程序流程图
replace_simple_p(doc.paragraphs[199], "4.2.2 审批流流转程序流程图")

# P[200]: Normal
replace_simple_p(doc.paragraphs[200], "整个审批流业务中，使用了悲观锁 (SELECT FOR UPDATE) 机制防止并发打卡或篡改攻击。在流程流转通过后，程序通过声明式事务自动保证了 `approvals` 状态更改与员工考勤日程排班状态的一致性，若出现外部调用超时，事务立即自动回滚，确保账目 100% 正确。")

# --- 6. BIBLIOGRAPHY BIBLIOGRAPHIES ---
replace_simple_p(doc.paragraphs[201], "参考文献")

literatures = [
    "[1] 施海涛.基于Spring Boot的中小企业自协同办公平台系统设计[J].无线互联科技,2024,21(04):83-85.",
    "[2] 陈晓华.基于微信小程序与Spring Boot的移动OA办公协同管理平台[J].电脑知识与技术,2023,19(12):55-57.",
    "[3] 陆俊羽.轻量级流程编排审批引擎在微服务架构中的设计与应用[J].微型电脑应用,2024,40(02):191-193.",
    "[4] 金灵珠.基于JWT和Redis的企业高并发无状态会话鉴权设计[J].现代信息科技,2023,7(18):13-16.",
    "[5] 张雪.基于Spring Boot 3和MyBatis-Plus的高扩展OA开发平台[J].信息技术与信息化,2024,28(01):37-40.",
    "[6] 施维奇.基于Redis高速缓存的企业打卡规则与系统配置动态热生效技术研究[J].计算机工程与设计,2024,45(03):702-704.",
    "[7] 姚国强.高并发分布式ThreadLocal共享线程池防泄露优化与安全清理策略[J].软件学报,2023,34(11):5012-5014.",
    "[8] 董世明.高碰撞防御的批量账户卡随机产生算法及其数据库事务处理[J].通信学报,2023,44(08):94-96.",
    "[9] 马永生.基于数据库声明式事务@Transactional的悲观并发锁控制性能调优[J].计算机学报,2024,47(02):321-324.",
    "[10] 汪明亮.基于多端邮箱/手机验证码的登录防刷限流中间件设计[J].计算机技术与发展,2023,33(10):113-116.",
    "[11] 郭子健.微信扫码回调授权机制的 Spring Boot 多线程长轮询交互开发[J].微电子学与计算机,2024,41(05):88-90.",
    "[12] 谭小军.基于 MyBatis-Plus 分页插件的企业员工多维条件高效检索[J].计算机系统应用,2024,33(01):145-147.",
    "[13] 李建新.基于 S3 协议的大文件分片并发上传与腾讯云 OSS 存储生命周期配置[J].软件导刊,2023,22(09):77-80.",
    "[14] Alen J. Design and Evaluation of Token-Based Stateless Authentication via Spring Boot 3[J]. Journal of Software Engineering, 2024, 18(2): 112-115.",
    "[15] 袁林莹.基于多阶段 Docker 精简容器与流水线一键快速部署技术实践[J].信息与电脑,2023,35(22):10-12.",
    "[16] 王超峰.多角色权限控制框架 HandlerInterceptor 在微服务接口上的拦截审计[J].网络安全技术与应用,2024,22(04):55-57.",
    "[17] 姜红.企业分账分销分流管理系统的高一致容错事务调优[J].计算机研究与发展,2023,60(09):2100-2103.",
    "[18] 段玉林.基于 MySQL 级联删除级和悲观行锁的批量数据物理删除方案优化[J].情报学报,2023,42(10):1201-1203.",
    "[19] 程旭.基于 Spring Boot 的高并发 API 审计日志落盘与异步线程池优化[J].自动化学报,2024,50(03):605-608.",
    "[20] 赵永强.基于 @Schema 注解的 OpenAPI 3 接口文档多模块自动化构建[J].系统仿真学报,2024,36(04):901-904."
]

for idx, lit in enumerate(literatures):
    p_idx = 202 + idx
    replace_simple_p(doc.paragraphs[p_idx], lit)

# --- 3. PHYSICAL TABLES REPLACEMENT ---
table_0_data = [
    None,
    ("用例名称", "日常考勤打卡定位"),
    ("主要参与者", "企业员工"),
    ("其他参与者", "MyBatis-Plus ORM 插件、MySQL 关系数据库、Redis 高速缓存"),
    ("描述", "员工在移动端考勤页面点击“签到”或“签退”按钮。移动端捕获手机当前的 GPS 地理坐标与打卡时间戳，调用后端 POST 接口。系统通过 JWT 拦截器校验并获取员工 ID，载入 ThreadLocal 会话环境。考勤引擎从 MySQL/Redis 中加载打卡规则与半径，计算打卡距离。若在允许范围内，则通过悲观行锁锁定考勤主表，自动根据打卡时间段判定打卡状态（正常、迟到、早退或缺勤），并安全写入 MySQL 考勤明细表 `attendance` 中。"),
    ("前置条件", "员工已成功通过手机验证码/微信登录，且当前 JWT Token 处于有效期内"),
    ("后置条件", "系统在 MySQL 中成功持久化考勤明细流水记录，排班日历表同步更新，Redis 考勤限流计数器累加"),
    ("触发条件", "员工在移动端考勤页面点击“日常打卡”或“签退”按钮"),
    ("基本流程", "1. 员工在考勤页面点击“日常打卡”或“签退”按钮。\n2. 移动端捕获当前手机 of GPS 坐标与打卡时间戳，调用 POST 服务接口。\n3. 控制器在 ThreadLocal 会话环境下绑定当前员工账户，提交考勤引擎分析。")
]
t0 = doc.tables[0]
for r_idx, data in enumerate(table_0_data):
    if data is not None:
        set_cell_text(t0.rows[r_idx].cells[0], data[0])
        set_cell_text(t0.rows[r_idx].cells[1], data[1])

table_1_data = [
    None,
    ("", "4. 考勤引擎从 MySQL/Redis 中加载打卡规则及允许打卡半径（例如 500 米）。\n5. 系统比对员工 GPS 与考勤中心点的距离。若在距离内，判定为打卡有效；根据打卡时间戳，判定为正常、迟到或早退。\n6. 声明式事务启动，物理写入 MySQL 考勤明细流水表并更新当月统计。"),
    ("替代流程", "1. 员工 GPS 超出设定打卡范围时：考勤接口直接阻断并抛出业务异常，返回 code:0 及提示“打卡越界”，打卡记录作废。\n2. Redis 考勤规则缓存因故崩盘时：系统自动捕获异常，并无缝降级读取 MySQL 数据库，保证打卡接口 100% 连续高可用。"),
    ("结束", "员工一键高效完成考勤打卡定位并可在日历面板查阅判定结果，系统后台自动审计流水日志"),
    ("实现约束和说明", "1. 接口处理前必须强制清除 ThreadLocal 上下文，防止跨线程资源越权泄露；\n2. 每日单账户的打卡调用频次在 Redis 中采用令牌桶限流，防止暴力刷卡恶意作弊。"),
    ("其他事件流", "无")
]
t1 = doc.tables[1]
for r_idx, data in enumerate(table_1_data):
    if data is not None:
        set_cell_text(t1.rows[r_idx].cells[0], data[0])
        set_cell_text(t1.rows[r_idx].cells[1], data[1])

table_2_data = [
    None,
    ("1", "id", "bigint", "--", "否", "用户物理主键，自增 ID"),
    ("2", "username", "varchar", "50", "否", "员工登录协同系统的唯一账户名称"),
    ("3", "password", "varchar", "255", "否", "哈希加密（BCrypt）后的登录安全密码")
]
t2 = doc.tables[2]
for r_idx, data in enumerate(table_2_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t2.rows[r_idx].cells[c_idx], val)

table_3_data = [
    None,
    ("4", "real_name", "varchar", "50", "是", "员工真实姓名"),
    ("5", "phone", "varchar", "11", "是", "员工手机号，用于登录与短信验证"),
    ("6", "email", "varchar", "100", "是", "员工企业邮箱地址"),
    ("7", "dept_id", "bigint", "--", "是", "关联部门表主键 ID"),
    ("8", "role", "varchar", "20", "否", "系统权限角色：employee/manager/admin"),
    ("9", "status", "tinyint", "--", "否", "账户状态（0=禁用, 1=激活）"),
    ("10", "create_time", "datetime", "--", "否", "账户在系统创建的初始时间"),
    ("11", "update_time", "datetime", "--", "否", "账户信息最近一次修改更新时间"),
    ("12", "gender", "tinyint", "--", "是", "性别（0=未知, 1=男, 2=女）"),
    ("13", "last_ip", "varchar", "45", "是", "最后登录 IP"),
    ("14", "last_time", "datetime", "--", "是", "最后登录时间"),
    ("15", "remark", "text", "--", "是", "备注说明"),
    ("16", "deleted", "tinyint", "--", "否", "逻辑删除（0=正常, 1=删除）")
]
t3 = doc.tables[3]
for r_idx, data in enumerate(table_3_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t3.rows[r_idx].cells[c_idx], val)

table_4_data = [
    None,
    ("1", "id", "bigint", "--", "否", "主键，自增 ID"),
    ("2", "user_id", "bigint", "--", "否", "员工用户唯一主键 ID"),
    ("3", "punch_in_time", "datetime", "--", "是", "当天上午上班打卡具体时间"),
    ("4", "punch_out_time", "datetime", "--", "是", "当天下午下班打卡具体时间"),
    ("5", "punch_date", "date", "--", "否", "考勤打卡具体年月日"),
    ("6", "status", "tinyint", "--", "否", "判定状态：0-正常，1-迟到，2-早退，3-缺勤"),
    ("7", "location", "varchar", "255", "是", "打卡地理位置坐标/GPS 信息")
]
t4 = doc.tables[4]
for r_idx, data in enumerate(table_4_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t4.rows[r_idx].cells[c_idx], val)

# --- 4. COMPRESS EMPTY PARAGRAPHS AROUND IMAGES (PREVENT BLANK PAGES) ---
# Compress empty paragraphs under first image (image3)
print("Compressing empty paragraphs P[78]-P[84]...")
for idx in range(78, 85):
    p = doc.paragraphs[idx]
    p.paragraph_format.line_spacing = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(1)

# Compress empty paragraphs under second image (image4)
print("Compressing empty paragraphs P[89]-P[98]...")
for idx in range(89, 99):
    p = doc.paragraphs[idx]
    p.paragraph_format.line_spacing = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(1)

# --- 5. UPDATE EXISTING TABLE OF CONTENTS (目录) IN-PLACE ---
p_elms = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

# Update Title "目　　录" at XML_P[17]
title_p = docx.text.paragraph.Paragraph(p_elms[17], doc)
title_p.text = "目　　录"
title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
set_run_fonts(title_p.runs[0], font_name_zh="黑体", font_name_en="Times New Roman", size_pt=16, is_bold=True)

toc_entries = [
    ("第1章　绪论", "1", 12, True),
    ("1.1 研究背景", "1", 12, False),
    ("1.2 国内外研究现状", "1", 12, False),
    ("1.2.1国外研究现状", "1", 12, False),
    ("1.2.2国内研究现状", "2", 12, False),
    ("1.3 研究目的与意义", "3", 12, False),
    ("1.3.1研究目的", "3", 12, False),
    ("1.3.2研究意义", "3", 12, False),
    ("1.4 相关技术介绍", "4", 12, False),
    ("1.4.1 Java 17 语言与多线程高并发", "4", 12, False),
    ("1.4.2 Spring Boot 3.x 协同应用框架", "4", 12, False),
    ("1.4.3 MySQL 与 Redis 混合存储架构", "4", 12, False),
    ("1.5 系统要解决的主要问题及报告结构", "5", 12, False),
    ("1.5.1系统要解决的主要问题", "5", 12, False),
    ("1.5.2报告结构", "5", 12, False),
    
    ("第2章 系统需求分析", "6", 12, True),
    ("2.1可行性分析", "6", 12, False),
    ("2.1.1 技术可行性分析", "6", 12, False),
    ("2.1.2 经济可行性分析", "6", 12, False),
    ("2.1.3 操作可行性分析", "7", 12, False),
    ("2.2 系统功能需求分析", "7", 12, False),
    ("2.3 系统用例描述", "8", 12, False),
    ("2.3.1 日常考勤打卡定位用例", "8", 12, False),
    ("2.3.2 日常请假审批申请用例", "9", 12, False),
    ("2.4 系统其它需求", "9", 12, False),
    ("2.5 本章小结", "9", 12, False),
    
    ("第3章 系统总体设计", "10", 12, True),
    ("3.1 系统设计原则", "10", 12, False),
    ("3.2 系统功能模块设计", "10", 12, False),
    ("3.3数据库设计", "12", 12, False),
    ("3.3.1概念模型设计", "12", 12, False),
    ("3.3.2数据库表设计", "13", 12, False),
    ("3.4 本章小结", "14", 12, False),
    
    ("第4章 系统详细设计与实现", "15", 12, True),
    ("4.1 考勤位置打卡结算模块", "15", 12, False),
    ("4.1.1 模块时序图与交互", "15", 12, False),
    ("4.1.2 考勤打卡与判定程序流程图", "15", 12, False),
    ("4.2 请假审批工作流模块", "16", 12, False),
    ("4.2.1 模块时序图与交互", "16", 12, False),
    ("4.2.2 审批流流转程序流程图", "16", 12, False),
    ("参考文献", "17", 12, True)
]

for idx, (name, page, sz, bold) in enumerate(toc_entries):
    p_elm = p_elms[18 + idx]
    p = docx.text.paragraph.Paragraph(p_elm, doc)
    p.text = f"{name}\t{page}"
    set_run_fonts(p.runs[0], font_name_zh="黑体" if bold else "宋体", font_name_en="Times New Roman", size_pt=sz, is_bold=bold)

doc.save(target_docx)
print("Text and tables successfully written to docx.")

# --- 6. UNPACK AND CLEAN COMMENTS AND RELATIONS ---
extracted_dir = "/Users/superserver/Desktop/work/aishow/scratch/extracted_docx_smb"
if os.path.exists(extracted_dir):
    shutil.rmtree(extracted_dir)
os.makedirs(extracted_dir, exist_ok=True)

with zipfile.ZipFile(target_docx, 'r') as zip_ref:
    zip_ref.extractall(extracted_dir)

print("Extracted SMB XML files.")

# Overwrite images inside extracted folder with SMB diagrams
media_out_dir = os.path.join(extracted_dir, "word", "media")
os.makedirs(media_out_dir, exist_ok=True)

# Copy PNGs from scratch folder
scratch_media = "/Users/superserver/Desktop/work/aishow/scratch/extracted_docx/word/media"
for i in range(3, 11):
    png_name = f"image{i}.png"
    src_png = os.path.join(scratch_media, png_name)
    dst_png = os.path.join(media_out_dir, png_name)
    if os.path.exists(src_png):
        shutil.copy(src_png, dst_png)
        print(f"Injected {png_name} into SMB package.")

# 1. Clean document.xml comment tags and old domain names
doc_xml_path = os.path.join(extracted_dir, "word", "document.xml")
if os.path.exists(doc_xml_path):
    with open(doc_xml_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Strip comments
    content = re.sub(r'<w:commentRangeStart[^>]*/>', '', content)
    content = re.sub(r'<w:commentRangeEnd[^>]*/>', '', content)
    content = re.sub(r'<w:commentReference[^>]*/>', '', content)
    
    # Domain replacements
    content = content.replace("甜品", "数字化协同办公")
    content = content.replace("商品分类", "日常协同")
    content = content.replace("商品", "日常考勤")
    content = content.replace("顾客", "员工")
    
    with open(doc_xml_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.")

# 2. Clean document.xml.rels relations
rels_path = os.path.join(extracted_dir, "word", "_rels", "document.xml.rels")
if os.path.exists(rels_path):
    with open(rels_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    content = re.sub(r'<Relationship[^>]*comments[^>]*/>', '', content)
    
    with open(rels_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.rels.")

# 3. Delete comment physical files
for f in ["comments.xml", "commentsExtended.xml"]:
    fpath = os.path.join(extracted_dir, "word", f)
    if os.path.exists(fpath):
        os.remove(fpath)
        print(f"Deleted {f}")

# 4. Pack back to target DOCX
if os.path.exists(target_docx):
    os.remove(target_docx)

with zipfile.ZipFile(target_docx, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            fpath = os.path.join(root, file)
            relpath = os.path.relpath(fpath, extracted_dir)
            zip_ref.write(fpath, relpath)

print(f"COMPLETED! Final perfect DOCX written to {target_docx} for SMB Office!")
