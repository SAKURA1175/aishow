import docx

doc_path = "/Users/superserver/Desktop/work/aishow/Study_AI_系统分析与设计课程报告.docx"
doc = docx.Document(doc_path)

def set_cell_text(cell, text):
    # Safe overwrite that preserves runs if possible to inherit font size/alignment
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

# ----------------- REPLACE TABLE 0 -----------------
# Table 0: Usecase 1 Part 1
table_0_data = [
    # Row 0 is header: Keep header ['用例条目', '描述']
    None,
    ("用例名称", "智能学业问答（双路检索 RAG 与联网搜索）"),
    ("主要参与者", "学生用户"),
    ("其他参与者", "AI问答服务端、ChromaDB 向量库、SearXNG 搜索引擎"),
    ("描述", "学生通过 Web 界面向系统发起学科提问，系统提取关键词，进行 BGE-M3 向量检索。若在 ChromaDB 命中 Top-K 资料，则拼接上下文注入 Prompt 并返回大模型流式生成；若未检索到或服务不可用，则自动无缝降级为自托管 SearXNG 联网搜索或通用问答。"),
    ("前置条件", "学生用户已成功登录系统且已建立长连接通道"),
    ("后置条件", "系统成功在 MySQL 记录对话历史与消息，后台异步进行提问主题自动归类及掌握画像日志更新"),
    ("触发条件", "学生在 Chat 双栏对话框中键入提问并点击“发送”按钮或按下回车"),
    ("基本流程", "1. 学生在输入框中键入学科提问并触发发送。\n2. 后端接收请求，首先由 InputSafetyFilter 拦截并校验越狱正则，阻断恶意提示词注入。\n3. 控制器在共享线程池中提交任务，提取学科关键词并启动 RAG 检索链路。")
]

t0 = doc.tables[0]
for r_idx, data in enumerate(table_0_data):
    if data is not None:
        set_cell_text(t0.rows[r_idx].cells[0], data[0])
        set_cell_text(t0.rows[r_idx].cells[1], data[1])

print("Table 0 replaced successfully!")

# ----------------- REPLACE TABLE 1 -----------------
# Table 1: Usecase 1 Part 2
table_1_data = [
    None, # Row 0 is header ['用例条目', '描述']
    ("", "4. 系统使用 BGE-M3 Embedding 模型将问题向量化并调用 ChromaDB API 检索 Top-3 关联切片。\n5. 后端组装 [资料X] 上下文注入 Prompt。若语义关联度差或 ChromaDB 异常，自动无缝降级为 MySQL 全文 LIKE 关键词频匹配与排序，或启动自建 SearXNG 隐私搜索获取实时网页来源。\n6. 大模型流式返回 token 片段，后端 WebFlux SSE 引擎通过 SseEmitter 实时吐给前端，React 异步渲染 Katex 公式与 Markdown。"),
    ("替代流程", "1. ChromaDB 离线时：系统自动切换为 MySQL DocumentChunk 表模糊 LIKE 全文扫描与频数排序兜底。\n2. 大模型超时离线时：系统向前端流式吐出服务不可用提示，并写入错误日志。"),
    ("结束", "学生成功流式收到完整回答且可以点击引用链接卡片跳转源文档，后台提问画像完成自动归类与日志更新"),
    ("实现约束和说明", "1. 为防长连接内存泄漏，SSE 线程的 ThreadLocal 上下文在 try-finally 块中强力释放；\n2. 单次提问长度限制在 4000 字符内，避免浪费上下文 token 资源。"),
    ("其他事件流", "无")
]

t1 = doc.tables[1]
for r_idx, data in enumerate(table_1_data):
    if data is not None:
        set_cell_text(t1.rows[r_idx].cells[0], data[0])
        set_cell_text(t1.rows[r_idx].cells[1], data[1])

print("Table 1 replaced successfully!")

# ----------------- REPLACE TABLE 2 -----------------
# Table 2: User schema Part 1
table_2_data = [
    None, # Row 0 is header ['序号', '字段名', '类型', '长度', '是否为空', '注释']
    ("1", "id", "bigint", "--", "否", "主键，自增ID"),
    ("2", "username", "varchar", "50", "否", "用户名（唯一）"),
    ("3", "password", "varchar", "255", "否", "BCrypt 加密密码")
]

t2 = doc.tables[2]
for r_idx, data in enumerate(table_2_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t2.rows[r_idx].cells[c_idx], val)

print("Table 2 replaced successfully!")

# ----------------- REPLACE TABLE 3 -----------------
# Table 3: User schema Part 2
table_3_data = [
    None, # Row 0 is header
    ("4", "role", "varchar", "20", "否", "角色：student/teacher/admin"),
    ("5", "avatar", "varchar", "500", "是", "头像 URL"),
    ("6", "create_time", "datetime", "--", "否", "账户创建时间"),
    ("7", "update_time", "datetime", "--", "否", "最后一次更新时间"),
    ("8", "nickname", "varchar", "50", "是", "用户昵称"),
    ("9", "email", "varchar", "100", "是", "电子邮箱"),
    ("10", "phone", "varchar", "20", "是", "联系电话"),
    ("11", "status", "int", "--", "否", "账户状态（0=禁用, 1=启用）"),
    ("12", "last_login_ip", "varchar", "45", "是", "最后一次登录 IP"),
    ("13", "last_login_time", "datetime", "--", "是", "最后一次登录时间"),
    ("14", "remark", "text", "--", "是", "备注说明"),
    ("15", "dept_id", "bigint", "--", "是", "关联院系/部门 ID"),
    ("16", "student_no", "varchar", "30", "是", "学生学号或教师工号"),
    ("17", "gender", "tinyint", "--", "是", "性别（0=未知, 1=男, 2=女）"),
    ("18", "deleted", "tinyint", "--", "否", "逻辑删除（0=正常, 1=删除）")
]

t3 = doc.tables[3]
for r_idx, data in enumerate(table_3_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t3.rows[r_idx].cells[c_idx], val)

print("Table 3 replaced successfully!")

# ----------------- REPLACE TABLE 4 -----------------
# Table 4: Document schema
table_4_data = [
    None, # Row 0 is header
    ("1", "id", "bigint", "--", "否", "主键，自增 ID"),
    ("2", "filename", "varchar", "255", "否", "上传课件文件的原始名称"),
    ("3", "file_type", "varchar", "20", "是", "文件类型：pdf/docx/doc/txt/md"),
    ("4", "uploader_id", "bigint", "--", "否", "上传教师用户的主键 ID"),
    ("5", "stored_filename", "varchar", "300", "是", "物理磁盘存储名称"),
    ("6", "char_count", "int", "--", "是", "提取的纯文本总字符数"),
    ("7", "chunk_count", "int", "--", "是", "文档被切片后的片段总数")
]

t4 = doc.tables[4]
for r_idx, data in enumerate(table_4_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t4.rows[r_idx].cells[c_idx], val)

print("Table 4 replaced successfully!")

doc.save(doc_path)
print("All tables updated successfully!")
