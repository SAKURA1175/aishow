import docx

template_path = "/Users/superserver/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_7bfrh4ndaf8l22_bdf7/msg/file/2026-05/23级-系统分析与设计-课程报告模板V1.3.docx"
doc = docx.Document(template_path)

paragraphs_to_check = [20, 21, 25, 26, 27, 30, 32, 36, 37, 39, 41, 48, 54, 66, 68, 200]

for p_idx in paragraphs_to_check:
    if p_idx < len(doc.paragraphs):
        p = doc.paragraphs[p_idx]
        print(f"\nP[{p_idx}] text: '{p.text}'")
        for r_idx, r in enumerate(p.runs):
            print(f"  Run[{r_idx}]: text='{r.text}' | superscript={r.font.superscript}")
