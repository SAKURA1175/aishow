import docx

doc_path = "/private/var/folders/m1/y8wt5fvd3hl_jb3c97rwk3w80000gn/T/MicrosoftEdgeDownloads/638411b0-51e9-4d19-91b8-d7cdff285eae/23级-系统分析与设计-课程报告模板V1.3.docx"
doc = docx.Document(doc_path)

print("--- Last 30 paragraphs ---")
for idx in range(max(0, len(doc.paragraphs)-30), len(doc.paragraphs)):
    p = doc.paragraphs[idx]
    print(f"P[{idx}] | Style: {p.style.name} | Text: {p.text}")
