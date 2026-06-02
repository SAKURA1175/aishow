import docx

doc_path = "/Users/superserver/Desktop/work/aishow/Study_AI_系统分析与设计课程报告.docx"
doc = docx.Document(doc_path)

print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

# Check some paragraphs
print("\n--- Verifying key paragraphs ---")
print(f"P[8] Title: {doc.paragraphs[8].text}")
print(f"P[17] Heading 1: {doc.paragraphs[17].text}")
print(f"P[20] text: {doc.paragraphs[20].text[:80]}...")
print(f"P[20] runs count: {len(doc.paragraphs[20].runs)}")
for idx, r in enumerate(doc.paragraphs[20].runs):
    print(f"  Run[{idx}]: text='{r.text}' | superscript={r.font.superscript}")

print(f"P[201] Bibliography heading: {doc.paragraphs[201].text}")
print(f"P[202] Ref 1: {doc.paragraphs[202].text}")

# Check tables
print("\n--- Verifying tables ---")
print(f"Table 0 Row 1: {[c.text.strip().replace('\n', ' ') for c in doc.tables[0].rows[1].cells]}")
print(f"Table 2 Row 1: {[c.text.strip() for c in doc.tables[2].rows[1].cells]}")
print(f"Table 4 Row 1: {[c.text.strip() for c in doc.tables[4].rows[1].cells]}")
