import docx
from docx.shared import Pt

files = [
    "/Users/superserver/Desktop/work/aishow/OYY_学院_系统分析与设计课程报告.docx",
    "/Users/superserver/Desktop/work/aishow/Study_AI_系统分析与设计课程报告.docx"
]

for doc_path in files:
    print(f"Compressing empty paragraphs for: {doc_path}")
    doc = docx.Document(doc_path)

    # Empty paragraphs under first image (image3)
    for idx in range(78, 85):
        p = doc.paragraphs[idx]
        p.paragraph_format.line_spacing = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(1)

    # Empty paragraphs under second image (image4)
    for idx in range(89, 99):
        p = doc.paragraphs[idx]
        p.paragraph_format.line_spacing = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(1)

    doc.save(doc_path)
    print(f"Successfully compressed {doc_path}!")
