import docx

doc_path = "/Users/superserver/Desktop/work/aishow/OYY_学院_系统分析与设计课程报告.docx"
doc = docx.Document(doc_path)

print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

# 1. Verify key paragraphs
print("\n--- Verifying key paragraphs ---")
print(f"P[8] Title: '{doc.paragraphs[8].text}'")
print(f"P[17] Heading 1: '{doc.paragraphs[17].text}'")
print(f"P[20] text: '{doc.paragraphs[20].text[:80]}...'")
print(f"P[20] runs count: {len(doc.paragraphs[20].runs)}")
for idx, r in enumerate(doc.paragraphs[20].runs):
    print(f"  Run[{idx}]: text='{r.text}' | superscript={r.font.superscript}")

print(f"P[201] Bibliography heading: '{doc.paragraphs[201].text}'")
print(f"P[202] Ref 1: '{doc.paragraphs[202].text}'")
print(f"P[221] Ref 20: '{doc.paragraphs[221].text}'")

# 2. Verify tables
print("\n--- Verifying tables ---")
print(f"Table 0 Row 1: {[c.text.strip().replace('\n', ' ') for c in doc.tables[0].rows[1].cells]}")
print(f"Table 2 Row 1: {[c.text.strip() for c in doc.tables[2].rows[1].cells]}")
print(f"Table 4 Row 1: {[c.text.strip() for c in doc.tables[4].rows[1].cells]}")

# 3. Check for lingering sweet-shop/TutorMarket terms
print("\n--- Scanning for lingering terms ---")
keywords = ['甜品', '商品', '顾客', '购物车', '结账', '送货', '订单', 'TutorMarket', 'Tutor', 'tutor', 'aisteam']
found_occurrences = []

for idx, p in enumerate(doc.paragraphs):
    for kw in keywords:
        if kw in p.text:
            found_occurrences.append(f"P[{idx}] contains '{kw}': '{p.text}'")

for t_idx, t in enumerate(doc.tables):
    for r_idx, r in enumerate(t.rows):
        for c_idx, cell in enumerate(r.cells):
            for kw in keywords:
                if kw in cell.text:
                    found_occurrences.append(f"T[{t_idx}] R[{r_idx}] C[{c_idx}] contains '{kw}': '{cell.text}'")

if found_occurrences:
    print(f"Found {len(found_occurrences)} potential remaining keywords:")
    for occurrence in found_occurrences[:10]:
        print(f"  {occurrence}")
else:
    print("Zero old domain keywords found! OYY Academy document is 100% clean and correct!")
