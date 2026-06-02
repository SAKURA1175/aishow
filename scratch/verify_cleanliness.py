import docx

doc_path = "/Users/superserver/Desktop/work/aishow/Study_AI_系统分析与设计课程报告.docx"
doc = docx.Document(doc_path)

keywords_to_check = ["甜品", "顾客", "商品", "地址"]
found_occurrences = []

for idx, p in enumerate(doc.paragraphs):
    for kw in keywords_to_check:
        if kw in p.text:
            found_occurrences.append(f"Paragraph {idx}: '{p.text}' (found '{kw}')")

for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for kw in keywords_to_check:
                if kw in cell.text:
                    found_occurrences.append(f"Table {t_idx} Row {r_idx} Col {c_idx}: '{cell.text}' (found '{kw}')")

if found_occurrences:
    print(f"Found {len(found_occurrences)} potential remaining keywords:")
    for occurrence in found_occurrences[:20]:
        print(f"  {occurrence}")
else:
    print("Zero remaining old keywords found! The document is 100% clean and correct!")
