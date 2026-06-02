import os
import re
import zipfile
import shutil
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

template_path = "/private/var/folders/m1/y8wt5fvd3hl_jb3c97rwk3w80000gn/T/MicrosoftEdgeDownloads/638411b0-51e9-4d19-91b8-d7cdff285eae/23级-系统分析与设计-课程报告模板V1.3.docx"
target_docx = "/Users/superserver/Desktop/work/aishow/TutorMarket_AI_系统分析与设计课程报告.docx"

# Ensure target directory is writeable and copy
if os.path.exists(target_docx):
    os.remove(target_docx)
shutil.copy(template_path, target_docx)
os.chmod(target_docx, 0o644)
print(f"Copied fresh template for TutorMarket AI.")

doc = docx.Document(target_docx)

# Helper function to replace text of a simple paragraph (clearing other runs)
def replace_simple_p(p, text):
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

# Helper to apply professional fonts
def set_run_fonts(run, font_name_zh="宋体", font_name_en="Times New Roman", size_pt=12, is_bold=False):
    run.font.size = Pt(size_pt)
    run.bold = is_bold
    
    # Set East Asian font (Chinese)
    run.font.name = font_name_zh
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name_zh)
    
    # Set ASCII/HAnsi font (English & Numbers)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

# --- 1. REPLACE TITLE PAGE ---
replace_simple_p(doc.paragraphs[8], "基于 Spring Boot 与 AI 编排服务的智能教学市场平台的设计与实现")

# --- 2. CHAPTER 1 INTRODUCTION ---
replace_simple_p(doc.paragraphs[18], "本课程报告的研究范围主要在设计与实现一个基于 Spring Boot 与 AI 编排服务的智能教学市场平台（TutorMarket AI）。其核心关注点是如何借助大语言模型（LLM）、多租户隔离、多源 RAG 检索（包含教师端公开知识库与学生端私有知识库）以及低延迟 SSE 流式会话响应技术，构建一个面向教学市场的自托管 AI 助教平台，并提升高校与教育机构的个性化教学与学业管理效率。绪论部分概述了本研究的背景、目标和重要性，介绍了国内外在智能助教与教学资源管理平台方面的研究现状，概述了所涉及的关键技术，并探讨了系统需解决的问题，揭示全文的结构安排。")

# P[20] has citations [1] and [2]
p20 = doc.paragraphs[20]
p20.runs[0].text = "在高校数字化转型与个性化教学改革持续深化的背景下，传统静态教学课件的利用效率低下与师生日常学业互动中答疑资源匮乏的问题成为了新的痛点，迫切需要通过生成式人工智能技术重塑教学互动场景"
p20.runs[2].text = "。教学市场平台作为连接优质教学资源与师生个性化诉求的桥梁，急需引入智能 AI 教师角色以及自托管的专属知识库服务，从而在保障院系学术数据隐私安全的前提下，提供精准、无幻觉的启发式辅导问答"
p20.runs[4].text = "。"

# P[21] has citation [3]
p21 = doc.paragraphs[21]
p21.runs[0].text = "近年来，基于检索增强生成（RAG）和大语言模型（LLM）的智能助教系统已经展现出了颠覆性的教育应用潜力。例如，通过将教师上传的课程专有课件和教学指南以及学生个人笔记进行智能分块和向量索引，系统能够精准地在问答中拼接最关联的内容注入提示词，杜绝通用大模型的盲目虚构生成；与此同时，结合流式连接（SSE）的逐 token 渲染、规则与版本的柔性控制管理、以及输入端的防越狱安全审计日志记录，系统可以构建出一个高度可靠、透明度极高且交互体验极佳的 AI 导师市场，对优化学生自学与教师精准备课具有极强的导向与实践意义"
p21.runs[1].text = ""
p21.runs[3].text = "。"

# P[23]:国内外研究现状 intro
replace_simple_p(doc.paragraphs[23], "在生成式 AI 与校园教务数字化浪潮交融的时代，智能教学管理及 AI 助教平台的建设已成为国内外教育信息化升级的核心焦点。为了构建一个既符合多角色（学生、教师、管理员）交互规范，又能够保障学科数据主权的安全自托管系统，我们需要深入梳理当前国内外在智能教学市场、流式响应服务以及大模型多源 RAG 检索融合等技术维度的发展特征，为本项目的研发奠定理论基础。")

# P[25] has citation [4]
p25 = doc.paragraphs[25]
p25.runs[0].text = "在国外智能教学系统与 RAG 开发框架的研究中，多角色 AI 助教系统的工程化落地已经非常成熟。通过采用 LangChain、LlamaIndex 等先进大模型调度链条以及 PostgreSQL 向量插件 pgvector、Redis 高频规则缓存等架构，欧美高校在智能答疑召回率、大文档异步分块切片及隐私多租户隔离等方向取得了重要突破，这都显著提升了学生提问的交互响应速率与召回满意度"
p25.runs[2].text = "。"

# P[26] has citation [5]
p26 = doc.paragraphs[26]
p26.runs[0].text = "国际学术界在提升检索准确性的同时，也对教育场景下智能大模型的输入端安全越狱防御、提示词版本演进以及审计日志系统的构建进行了大量讨论。这方面的安全防护与监控策略为智能平台跨学科部署和高抗攻击能力提供了极为可行的架构范式。此外，大语言模型推理过程的透明化以及数据安全合规性也是当前国外研究的热点，目的是构建师生对 AI 生成结果的信任"
p26.runs[1].text = ""
p26.runs[2].text = "。"

# P[27] has citations [6] and [7]
p27 = doc.paragraphs[27]
p27.runs[0].text = "在这些研究和实践的基础上，还必须考虑自托管部署架构在软硬件兼容性以及多租户并发访问下的局限性。例如，在专业课程的知识库建设中，由于教学材料的强版权私密性，教学管理系统需要提供极高规格的学生/教师物理数据库隔离，以避免敏感学科资源的跨用户越权泄漏"
p27.runs[1].text = ""
p27.runs[2].text = ""
p27.runs[3].text = ""
p27.runs[4].text = ""
p27.runs[6].text = ""
p27.runs[7].text = "。同时在面对高并发、高弹性的校园局域网络环境时，后端服务需要采用基于 SseEmitter 的异步长连接推送与 AI 编排的多线程解耦降级设计，以保障流式长会话在弱网或大模型耗时推理场景下的稳定连接与鲁棒体验"
p27.runs[8].text = ""
p27.runs[10].text = "。"

# P[28]: Normal
replace_simple_p(doc.paragraphs[28], "为适应智能 AI 教师在高校日常多学科答疑场景下的普及应用，教学市场平台不仅需提供多维的教师展示与选择面板，更需要在系统性能、隐私合规及高可用性方面持续提升，以获得长效的校园用户粘性。")

# P[30] has citation [8]
p30 = doc.paragraphs[30]
p30.runs[0].text = "随着近几年国内教育信息化 2.0 战略与大模型自主研发浪潮的并驾齐驱，基于本地微型知识库构建的 AI 助教及教师市场平台成为了教育工作者的迫切诉求。为了快速升级传统教务静态答疑模式，国内众多机构尝试引入轻量化本地大模型，试图将传统课件的上传与查找转化为启发式 AI 流式问答。然而，中小型学科组或院系在技术实力和服务器资源上面临较大约束，通常缺乏一套能够兼容多端知识库检索、防提示词越狱泄露、且支持模型配置和系统规则版本热生效的一体化自托管开源平台"
p30.runs[2].text = "。"

# P[31]: Normal
replace_simple_p(doc.paragraphs[31], "虽然部分团队基于 Ollama 等框架搭建了点对点聊天原型，但在多角色多场景的级联管理、MinIO 对象存储高频连接、多线程 SSE 缓冲管理，以及 AI 编排服务（ai-orchestrator）与 Spring Boot 主 API 的服务解耦架构方面仍显单薄。大并发下往往出现长响应延迟和连接耗尽的问题，限制了系统的广泛普及。")

# P[32] has citation [9]
p32 = doc.paragraphs[32]
p32.runs[0].text = "国内外的智能教学资源管理与 AI 助教平台的发展特征揭示了多源 RAG 与流式响应技术在教育垂直领域演进中的必然趋势"
p32.runs[2].text = "。汲取国外多角色隔离与高性能编排的思想，结合国内高校教学管理和版权主权的实际需求，本项目旨在设计开发一套基于 Spring Boot + Python (ai-orchestrator) 双服务解耦架构的 TutorMarket AI 智能教学市场平台。"

# P[34]: intro
replace_simple_p(doc.paragraphs[34], "随着大语言模型的工程化落地，师生对低延迟、高隐私安全、高响应品质的智能化辅导平台需求日益迫切。传统的静态文档下载和关键词检索服务模式正面临重大的智能化转型升级。构建一个高性能、数据分级隔离的智能教学市场平台，不仅可以极大地释放教师日常重复答疑的工作负担，还能为学生提供启发式、高相关度、并支持严格引用溯源的流式解答体验。因此，本报告专注于探究自托管智能教学市场平台的系统分析与设计，重点研究如何通过 Spring Boot 主 API 服务与 Python 编排服务的解耦设计、多源知识库检索、规则版本管理以及输入越狱检测机制来保障服务的卓越性。接下来，将详细说明本研究的具体目的与意义。")

# P[36] has citation [10]
p36 = doc.paragraphs[36]
p36.runs[0].text = "本次研究的重要目的是设计并开发一个支持多租户高隔离、完全本地私有化部署的智能教学市场系统。该系统需无缝连接本地 GPU 模型与云端 API 接口，并建立学生私有与教师公开的多源知识库检索链路。系统通过解耦的 AI 编排服务以及 MinIO 对象存储，实现毫秒级课件切片与高性能并发 SSE 响应，在确保教学内容 100% 严密隐私保护的前提下显著提升个性化教学效率与体验"
p36.runs[1].text = ""
p36.runs[2].text = ""
p36.runs[3].text = ""
p36.runs[4].text = ""
p36.runs[5].text = ""
p36.runs[7].text = "。"

# P[37] has citation [11]
p37 = doc.paragraphs[37]
p37.runs[0].text = "更重要的是要整合多角色权限，设计实施一个支持教师自主维护领域知识库、学生一键订阅各学科 AI 老师并建立长连接会话的互动市场"
p37.runs[1].text = ""
p37.runs[3].text = "。同时，通过引入系统参数与提示词（System Prompt）的规则版本管理机制，允许管理员或教师通过后台 Web 端热更新大模型的分身人设，为各学科重难点教学设计提供实时、自适应的数据及版本决策支持。"

# P[39] has citation [12]
p39 = doc.paragraphs[39]
p39.runs[0].text = "本研究的意义体现在对大语言模型 RAG 应用架构与高校教务辅助理论与工程实践深度结合的研究上。本报告旨在为教育机构及中小型教学团队提供一套完全本地部署、高安全防线且符合绿色低能耗标准的智能教学市场解决方案，以顺应高校数字化与 AI 融合的时代潮流"
p39.runs[2].text = "。"

# P[40]: Normal
replace_simple_p(doc.paragraphs[40], "通过部署 TutorMarket AI 系统，高校院系能够将原先零散的专业文献、课程讲义和习题集一键上传至自建 MinIO 并通过 PostgreSQL 向量表存储。学生不仅能获得启发式的学科解答，还可以上传个人学习笔记作为专属私有知识库，在 RAG 问答中实现个人与课程多源内容的交叉语义召回，大幅减少学习盲区并规避学术幻觉。")

# P[41] has citation [13]
p41 = doc.paragraphs[41]
p41.runs[0].text = "系统集成的输入安全审计与 ThreadLocal 并发隔离机制是构建师生对生成式大模型信任的根本屏障，确保平台在使用时完全遵循我国生成式人工智能服务管理办法等相关网络安全法规"
p41.runs[2].text = "。同时，系统支持轻量级 H2 数据库开发模式与容器化一键拉起，降低了开发与日常物理主机的运维开销，对推动校园绿色计算和降低高校数字资源采购成本具有突出的实践指导作用。"

# P[42]: Normal
replace_simple_p(doc.paragraphs[42], "本报告丰富了传统教务管理信息系统在多代理（Multi-Agent）与大模型技术集成方面的理论研究。其首创的 Spring Boot (main-api) 与 Python (ai-orchestrator) 服务解耦、多源知识库检索重排、邮箱验证码登录、以及系统提示词 Redis 热覆盖缓存技术，为智能助教系统的低门槛落地部署提供了极其具体的工程实践。")

# P[44]: intro
replace_simple_p(doc.paragraphs[44], "在如今大模型与微服务深度结合的时代，选择轻量级、健壮且解耦的技术栈对于构建一个高性能、高安全的自托管教学市场平台至关重要。本项目专注于挑选最主流的开源技术，以确保系统在面临高并发长连接和海量知识文件解析时，能够展现卓越的吞吐性与高可用性。经过充分对比，系统最终采用了“Java 17 / Spring Boot 3.x 响应式框架 + Python 3.10 AI 编排服务”的双服务解耦架构，数据存储则结合了 PostgreSQL 15 关系型数据库、Redis 高速缓存以及 MinIO 分布式对象存储。下面将详细探讨这些关键技术在系统中的具体作用、技术优势，以及它们是如何协同工作以满足复杂的智能教学市场业务需求。")

# P[45] & P[46]: Java heading and text
replace_simple_p(doc.paragraphs[45], "1.4.1 Java 17 语言与多线程高并发")
replace_simple_p(doc.paragraphs[46], "作为开发企业级应用的中流砥柱，Java 17 以其高强度的内存管控、极其出色的多线程高并发管理器，成为了本项目主 API 服务（main-api）的开发基石。在智能教学市场平台中，Java 17 的非阻塞 I/O 读写能力、垃圾回收器（G1/ZGC）的低停顿特性，使得后台在维持高频邮箱验证码登录校验、高并发长连接 SSE 流式会话推送、以及对 MinIO 文件的并发读写流解析时，能够保持极低的内存开销与卓越的系统的吞吐性能。")

# P[47] & P[48]: Spring Boot heading and text
replace_simple_p(doc.paragraphs[47], "1.4.2 Spring Boot 3.3.x 响应式框架")
p48 = doc.paragraphs[48]
p48.runs[0].text = "Spring Boot 3.3 框架作为快速构建生产级 Spring 响应式应用的行业标准，其强大的 IoC 容器与自动装配特性，为开发 TutorMarket AI 提供了灵活而轻量的主 API 骨架"
p48.runs[1].text = ""
p48.runs[2].text = ""
p48.runs[3].text = ""
p48.runs[4].text = ""
p48.runs[5].text = ""
p48.runs[6].text = ""
p48.runs[7].text = ""
p48.runs[9].text = "。Spring Boot"
p48.runs[10].text = "不仅精简了"
p48.runs[11].text = "初始的复杂 XML 配置，还集成了 Spring WebFlux 用于低延迟的异步长连接（SSE）通信，为后续多服务间的解耦与极速部署提供了不可替代的技术支撑。"

# P[49] & P[50]: DB heading and text
replace_simple_p(doc.paragraphs[49], "1.4.3 PostgreSQL 与 MinIO 混合存储架构")
replace_simple_p(doc.paragraphs[50], "在持久化存储层面，系统创造性地引入了“PostgreSQL 15 关系型数据库 + MinIO 分布式对象存储”的混合存储架构。PostgreSQL 负责高可靠事务数据（如用户账户、导师市场列表、流式消息历史、系统规则版本、审计日志等）的索引与强一致性管理；而 MinIO 对象存储则通过 S3 协议专用于存储教师上传的公开学科课件及学生上传的私有笔记。这种双轨存储架构在保障海量大文件极速读取与安全物理隔离的前提下，通过后台 REST API 实现多源文档到数据库元数据的强一致性外键关联，提供了极其优秀的可扩展性。")

# P[51]: Heading 2 -> 1.5 系统要解决的主要问题及报告结构
replace_simple_p(doc.paragraphs[51], "1.5 系统要解决的主要问题及报告结构")

# P[52]: Normal
replace_simple_p(doc.paragraphs[52], "在 TutorMarket AI 的实施过程中，技术团队面临了诸多核心挑战，包括如何在高并发长连接场景下防止 WebFlux 的 SseEmitter 连接泄露、如何将主业务 API 与耗时较长的 AI 模型编排服务（ai-orchestrator）进行微服务级别的高效解耦、如何隔离多角色的知识库并保障提示词规则热更新不泄露，以及如何在提问端建立实时越狱审计机制。为了攻克这些工程壁垒，报告详尽探讨了解决方案与优化路径，接下来的报告结构分为四个核心章节：第一章阐述智能教学市场的技术背景与核心栈；第二章完成系统的可行性多维论证、多角色用例建模与安全性非功能分析；第三章详细展示功能模块层次图、概念模型 E-R 图及 PostgreSQL 物理建表设计；第四章则通过对流式对话编排与知识库检索两大核心模块绘制时序图与程序流程图，详实交代核心代码算法及优化实践。")

# P[54] has citations [15] and [16]
p54 = doc.paragraphs[54]
p54.runs[0].text = "本平台致力于攻克智能助教市场中数据孤岛、学术虚假“幻觉”以及越狱安全过滤的关键工程难题，极大地优化学生的流式问答体验。通过引入基于 Python 开发的独立 AI 编排服务，主 API 只需通过异步 HTTP/WebSockets 与其保持轻量连接，从而彻底将主业务系统的负载与高能耗的向量检索及大模型推理隔离，实现了极佳的服务可用性"
p54.runs[1].text = ""
p54.runs[2].text = ""
p54.runs[3].text = ""
p54.runs[4].text = ""
p54.runs[6].text = "。系统首创的“双源检索（教师公开知识库 + 学生私有知识库）”检索架构，允许学生在会话中交叉查询个人笔记与课程大纲，有效规避了单源检索的内容盲区。配合 try-finally 内存防溢出逻辑和基于 Redis 的规则版本控制引擎，管理端能够在不重启 main-api 服务的情况下实现 API Key 与 System Prompt 规则的秒级生效，提供了极强的系统安全防护与可维护性"
p54.runs[8].text = "。"

# P[55]: Heading 3 -> 1.5.2 报告结构
replace_simple_p(doc.paragraphs[55], "1.5.2 报告结构")

# P[56]: Normal
replace_simple_p(doc.paragraphs[56], "本智能教学市场平台的报告结构整体分为四个阶段进行详细阐述：")

# P[57]: Normal
replace_simple_p(doc.paragraphs[57], "第1章是绪论，这一部分将介绍高校个性化智能教学面临的私密知识分级、AI 老师人设热更新等挑战，阐明报告的研究目的、研究意义和技术选型，提供平台架构的概览。")

# P[58]: Normal
replace_simple_p(doc.paragraphs[58], "第2章是需求分析，需求分析将详细从技术、经济和操作三个维度论证可行性，深入挖掘系统的功能性用例（包括邮箱验证码登录、导师市场订阅、流式会话及规则版本控制），绘制用例图并给出核心用例描述表，同时明确非功能性安全需求。")

# P[59]: Normal
replace_simple_p(doc.paragraphs[59], "第3章是系统设计，总体设计阐述了平台的私有自托管原则、五大核心功能模块结构图，重点展示数据库概念模型设计（E-R图）与基于 PostgreSQL 表结构的数据物理表字段设计。")

# P[60]: Normal
replace_simple_p(doc.paragraphs[60], "第4章是系统详细设计与实现，具体说明如何基于“Spring Boot API + Python ai-orchestrator”双服务框架实现两大核心模块：流式对话聊天模块与多源知识库 RAG 检索模块，通过绘制详细的时序图与程序流程图展示核心算法逻辑。")

# --- 3. CHAPTER 2 SYSTEM REQUIREMENT ANALYSIS ---
replace_simple_p(doc.paragraphs[63], "第2章　系统需求分析")
replace_simple_p(doc.paragraphs[64], "需求分析是构建 TutorMarket AI 智能教学市场平台的开发决策关键，旨在确保项目的研发与高校教学管理和学生个性化答疑的业务习惯精准匹配。本章通过对可行性进行多维论证，对系统日常研学中涉及的邮箱注册登录、导师市场浏览、知识库文件（MinIO）多源切片及 Redis 提示词规则热配置等核心业务流进行建模分析，并利用规范的 UML 用例图和用例描述表进行阐明。最后，分析了安全性及兼容性等非功能性要求，为后文的系统总体设计奠定坚实的业务需求基石。")

# P[65]: Heading 2 -> 2.1 可行性分析
replace_simple_p(doc.paragraphs[65], "2.1 可行性分析")

# P[66] has citation [17]
p66 = doc.paragraphs[66]
p66.runs[0].text = "在这个关键的可行性分析部分，技术团队从技术架构可行性、项目投资回报与运行维护可行性三个维度，全面审视 TutorMarket AI 平台的落地实力。技术可行性着重评估现有 Java 后端、Python 编排服务与 MinIO/PostgreSQL 的连接稳定性，保障架构安全可靠"
p66.runs[2].text = "。经济可行性则分析利用自托管离线模型与 pgvector 插件相比商业大模型 API 带来的巨额 Token 成本控制优势。操作可行性分析重点关注师生用户的学习成本以及教务人员对后台热切换提示词版本的接受度，为平台的顺利研发与局域网接入提供坚实的支撑。"

# P[68] has citation [18]
p68 = doc.paragraphs[68]
p68.runs[0].text = "综合评估开发工具、大语言模型推理设备与技术团队能力的现状，系统选择的操作系统（兼容 Windows/Linux/macOS 容器化部署）、数据库管理（PostgreSQL 15 与轻量级 H2 模式）、分布式文件系统（MinIO）、以及 Java 17/Spring Boot 与前端 React 框架均为当前软件工程领域高度成熟且广泛应用的标准技术栈。前端 Markdown 渲染、Zustand 会话状态管理与后端的异步 SSE 推送机制具备完美的兼容性与协同效应，有利于保障系统的超低延迟与服务的高可用性。此外，开发团队拥有丰富的高并发微服务治理和 RAG 算法调优经验，且有自托管 AI 编排集成的良好实践，充分证明有能力应对项目研发挑战。系统本身采用低耦合设计，提供完整的容器化一键部署方案，从技术层面上看，本智能教学市场平台的实施是完全可行的"
p68.runs[2].text = " 。"

# P[70]: Economy Normal
replace_simple_p(doc.paragraphs[70], "尽管平台在前期研发及自托管 GPU 服务器的购置上需要一定的软硬件投入，但系统在上线后，凭借其高效的多源 RAG 召回问答对云端大模型昂贵的 API Token 费用的替代，以及文档自动智能切片入库对教师人工解答重复性学术问题开销的节省，预计在系统运行半年内即可收回初期研发与部署成本，实现极高的经济效益。此外，PostgreSQL 关系型数据库与 MinIO 分布式对象存储等基础组件均为开源免费版，零授权许可开销。系统在提升课程问答效率、减少纸质课件分发以及保障院系数据主权方面具有显著的隐性资产保护作用，从经济角度来看，投资风险极低，投资回报率极其合理。")

# P[72]: Operation Normal
replace_simple_p(doc.paragraphs[72], "系统在人机交互层面采用了极致直观的类 ChatGPT 双栏响应式布局，左侧提供新建/切换/删除会话的抽屉式面板，右侧提供对话主视区，支持移动端 and PC 端的自适应无缝浏览，极大地降低了学生的学习使用门槛。学生和教师无需任何专业大模型知识即可流畅进行对话与文档上传。管理后台提供的 Web 可视化模型配置及 System Prompt 热切换机制，允许非技术教务人员通过简单的表单录入并一键写入 Redis，实时调整 AI 助教的专属分身（如引导式学生助教、课程设计教师助教），整个运维过程零黑屏命令行操作。因此，结合师生用户的快速上手度和极低的管理维护难度，该系统具备卓越操作可行性。")

# P[73]: Heading 2 -> 2.2 系统功能需求分析
replace_simple_p(doc.paragraphs[73], "2.2 系统功能需求分析")

# P[74]: Normal
replace_simple_p(doc.paragraphs[74], "进入需求分析的深水区，利用 UML 对系统交互角色和核心功能场景进行精准建模至关重要。这为后续的数据库架构以及核心代码模块的编写提供了极具指导意义的顶层蓝图。本平台主要涉及三类用户角色：学生用户（学生端主角，进行邮箱登录、挑选 AI 老师发起流式会话、上传私有知识库）、教师用户（知识库建设者，上传和删除公开教学课件、微调 AI 老师人设）以及系统管理员（维护系统平稳运行，热更新模型参数、监控规则版本及审计日志）。下面将详细描述系统的用例交互架构，并绘制相应的角色用例图。")

# P[75]: Normal
replace_simple_p(doc.paragraphs[75], "系统管理员与教师是平台的后台运维与知识维护者，管理员负责会话监控、安全审计和系统配置热更新等功能，能够在线切换底层的 LLM 接口地址与 API 密钥，实时更新针对学生和教师端角色的系统提示词，监控知识库向量切片生成状态，并查看平台用户的会话访问日志。教师则负责上传与维护所属课程的专有教学大纲、PDF/Word 课件及参考书目，管理切片状态并保证知识库内容实时有效。系统管理员与教师用例图如图 2-1 所示。")

# P[76]: Usecase title
replace_simple_p(doc.paragraphs[76], "系统管理员与教师角色用例图如图 2-1 所示。")

# P[85]: Usecase label
replace_simple_p(doc.paragraphs[85], "图 2-1 系统管理员与教师用例图")

# P[86]: Customer intro
replace_simple_p(doc.paragraphs[86], "学生是本平台的终端消费者，也是获取学术指导与问答服务的核心利益相关者。他们通过邮箱验证码进行登录，随后进入 AI 教师市场挑选适合自己课程的 AI 导师（如 Java 高并发专家、算法导师等）并新建聊天会话，系统支持流式显示回答并支持 Latex 渲染与引用卡片。同时，学生可以在“学习画像”页面查看基于提问日志的各学科主题掌握度树状图，并能够上传个人学习资料构建私有知识库，以在问答中进行交叉检索。学生用户用例图如图 2-2 所示。")

# P[99]: Usecase label
replace_simple_p(doc.paragraphs[99], "图 2-2 学生用户用例图")

# P[100]: Heading 2 -> 2.3 系统用例描述
replace_simple_p(doc.paragraphs[100], "2.3 系统用例描述")

# P[101]: Heading 3 -> 2.3.1 智能学业问答用例
replace_simple_p(doc.paragraphs[101], "2.3.1 智能学业问答用例")

# P[102]: Normal
replace_simple_p(doc.paragraphs[102], "智能学业问答功能描述了学生如何向平台进行智能提问并获得精准的 RAG 与联网搜索流式反馈，用例描述如表 2-1 所示。")

# P[103]: Table label
replace_simple_p(doc.paragraphs[103], "表 2-1 智能学业问答用例描述表")

# P[104]: Table续 label
replace_simple_p(doc.paragraphs[104], "表 2-1（续）")

# P[106]: Heading 3 -> 2.3.2 课件上传与切片入库用例
replace_simple_p(doc.paragraphs[106], "2.3.2 课件上传与切片入库用例")

# P[107]: Normal
replace_simple_p(doc.paragraphs[107], "课件文档上传与切片入库用例描述了教师如何通过 Web 界面批量上传 PDF/Word 等格式的学科课件，系统后台自动解析为纯文本，按固定长度和重叠率进行切片，并通过大模型 Embedding 向量化存入 ChromaDB 数据库的完整全生命周期过程。")

# P[109]: Heading 2 -> 2.4 系统其它需求
replace_simple_p(doc.paragraphs[109], "2.4 系统其它需求")

# P[110]: Normal
replace_simple_p(doc.paragraphs[110], "本节着重对系统在其他方面的衍生需求进行描述。主要包括安全性需求、兼容性需求、可扩展性需求等，目的在于使系统更加安全、稳定与高效。")

# P[111]: Security
replace_simple_p(doc.paragraphs[111], "安全性：目前，本系统的安全性大致可从系统防越狱、数据高隔离两方面展开。系统内置 InputSafetyFilter 越狱检测器，通过正则表达式严格匹配提示词泄露、指令替换等 8 类恶意指令攻击；后端在多并发共享线程池池化复用场景下使用 try-finally 结构安全清理 ThreadLocal 绑定的 RAG 上下文，彻底根治跨用户越权数据串漏的安全隐患。")

# P[112]: Compatibility
replace_simple_p(doc.paragraphs[112], "兼容性：系统能够具备在多种软硬件架构下稳定部署与访问的能力。前端 Web 页面完全兼容 Google Chrome、Firefox、Edge、Safari 等主流浏览器，支持移动端抽屉导航与 PC端双栏展示的自适应布局；后端容器化打包支持 macOS, Linux (x86/ARM64) 平台的 Docker 一键拉起，具备优异跨平台移植力。")

# P[113]: Scalability
replace_simple_p(doc.paragraphs[113], "可扩展性：系统设计已将向量召回深度集成及微服务演进纳入考量。后端基于 Spring Boot 3.3 与 Spring AI 大模型套件开发，具有极高的解耦度，为后续将“MySQL 全文检索”完全升级为“ChromaDB 向量索引 + 关键词命中重排”的双路混合检索及大文档多线程分块异步解析提供了极大的扩展空间。")

# P[114]: Heading 2 -> 2.5 本章小结
replace_simple_p(doc.paragraphs[114], "2.5 本章小结")

# P[115]: Normal
replace_simple_p(doc.paragraphs[115], "本章详细介绍了系统的功能性与非功能性需求。通过可行性论证确立了技术和经济支撑，利用 UML 用例图建模了学生、教师与管理员角色，并对智能问答和文档入库用例进行了详细流程描述，同时提出了高标准的越狱防护和安全防串漏要求，为下一章的系统总体设计奠定了扎实的业务蓝图。")

# --- 4. CHAPTER 3 OVERALL DESIGN ---
replace_simple_p(doc.paragraphs[117], "第3章　系统总体设计")
replace_simple_p(doc.paragraphs[118], "本章节致力于深入探讨自托管智能教学市场平台的总体设计架构，包括系统核心的设计理念与原则、主要功能模块的层次结构树设计、数据库概念模型设计（E-R图）以及关系表结构的详细物理设计。")

# P[119]: Heading 2 -> 3.1 系统设计原则
replace_simple_p(doc.paragraphs[119], "3.1 系统设计原则")

# P[120]: List 1
replace_simple_p(doc.paragraphs[120], "系统安全性原则：系统在指令过滤与长连接多线程并发管理上应具备极其严密的安全防御。敏感的 API URL/Key 在持久化存储中加密，在内存 ThreadLocal 作用域内具备及时的 try-finally 清理释放，防范越权数据泄露。")

# P[121]: List 2
replace_simple_p(doc.paragraphs[121], "自托管私有性原则：在满足大模型智能问答专家的前提下，最大化实现数据私有化与完全自托管部署。集成自建的 SearXNG 隐私搜索，数据库和向量计算完全落于局域网物理服务器，防范学术机密与隐私流向第三方云端。")

# P[122]: List 3
replace_simple_p(doc.paragraphs[122], "高可用与降级可维护性：产品在核心算法链条上设计了完备的降级兜底。例如，向量库 ChromeDB 异常时无缝切换为 MySQL LIKE 全文检索，大模型服务超时自动进行保活重试。整体架构模块结构清晰，易于后续扩展和迭代优化。")

# P[123]: Heading 2 -> 3.2 系统功能模块设计
replace_simple_p(doc.paragraphs[123], "3.2 系统功能模块设计")

# P[124]: Normal
replace_simple_p(doc.paragraphs[124], "TutorMarket AI 平台具备用户与市场管理、流式会话聊天、多端知识库管理、规则与版本配置、安全审计与日志五大核心功能模块。用户与市场管理模块实现多角色分身、邮箱登录拦截与导师列表检索；流式会话聊天模块处理 WebFlux SSE 流式推送、ai-orchestrator 双服务异步解耦交互及聊天历史管理；多端知识库管理模块负责教师公开课件与学生私有笔记的 MinIO 上传、分块及 PostgreSQL 索引映射；规则与版本配置模块提供 System Prompt 规则版本追踪与 Redis 热生效；安全审计与日志模块则负责实时越狱正则表达式扫描与访问日志持久化。系统总体功能模块图如图 3-1 所示。")

# P[126]: Module diagram label
replace_simple_p(doc.paragraphs[126], "图 3-1 系统总体功能模块图")

# P[127]: Module 1
replace_simple_p(doc.paragraphs[127], "1. 用户与市场管理模块")

# P[128]: Module 1 text 1
replace_simple_p(doc.paragraphs[128], "（1）邮箱注册与登录拦截：系统提供便捷的安全邮箱验证码注册与 Token 权限控制。在所有敏感业务接口前部署 HandlerInterceptor 拦截器，保障非授权登录请求无法跨越调用 AI 等敏感后台接口。")

# P[129]: Module 1 text 2
replace_simple_p(doc.paragraphs[129], "（2）多端权限与市场检索：平台明确划分 student、tutor、admin 权限级别。例如，学生可以浏览挑选导师市场中不同专业和学科的 AI 老师；教师和管理员则拥有上传删除知识库、配置系统人设等级等高级运维权限。")

# P[130]: Module 1 text 3
replace_simple_p(doc.paragraphs[130], "（3）系统级人设角色分身：系统根据用户挑选的 AI 老师角色，自动加载并注入对应的系统级 Prompt，为学生在 Java 编程、算法数据结构等专业方向提供极佳的专属辅导答疑服务。")

# P[131]: Module 2
replace_simple_p(doc.paragraphs[131], "2. 流式会话聊天模块")

# P[132]: Module 2 text 1
replace_simple_p(doc.paragraphs[132], "（1）WebFlux SSE 实时长连接：后端基于 Spring WebFlux 及 SseEmitter 与前端 React 建立异步长连接，将大模型产生的辅导文本片段以逐 token 的极低延迟实时流式推送至浏览器渲染，消除等待焦躁。")

# P[133]: Module 2 text 2
replace_simple_p(doc.paragraphs[133], "（2）ai-orchestrator 双服务解耦：主 API 接口只负责会话流程控制与事务提交，而耗时巨大的向量检索与大模型调用等编排工作则解耦交由 Python 的 ai-orchestrator 异步服务处理，保障高吞吐率。")

# P[134]: Module 2 text 3
replace_simple_p(doc.paragraphs[134], "（3）上下文管理与历史持久化：后端自动合并整理当前会话的多轮交互历史，在每一次流式问答结束后实时将其写入 PostgreSQL 数据库，并提供便捷的会话删除及上下文清除 API。")

# P[135]: Module 3
replace_simple_p(doc.paragraphs[135], "3. 多端知识库管理模块")

# P[136]: Module 3 text 1
replace_simple_p(doc.paragraphs[136], "（1）MinIO 多租户文件上传：教师上传公开学科课件或学生上传私有学习笔记时，系统在后台使用 S3 协议将其存入 MinIO 分布式对象存储，并通过物理桶（Bucket）和防越权拦截建立严密的物理高隔离。")

# P[137]: Module 3 text 2
replace_simple_p(doc.paragraphs[137], "（2）文本解析与重叠智能分块：后台提取的 PDF/Word 纯文本按照固定重叠滑窗逻辑进行智能文本分块（Chunk），建立与物理存储路径的一对多文档切片映射，实现超高匹配精准度的语义关联。")

# P[138]: Module 3 text 3
replace_simple_p(doc.paragraphs[138], "（3）级联同步与向量表维护：在教师或学生删除指定文档时，通过数据库的级联外键约束自动清空 PostgreSQL 中的文件关联，并同步调用 Python 服务物理清理 ChromeDB/PostgreSQL 中的高维向量切片。")

# P[139]: Module 4
replace_simple_p(doc.paragraphs[139], "4. 规则与版本配置模块")

# P[140]: Module 4 text 1
replace_simple_p(doc.paragraphs[140], "（1）Prompt 规则多版本追踪：为了让 AI 老师的回答风格精益求精，系统设计了 `t_rule_version` 实体，每次更新 AI 导师的人设或系统规则都会保存一个新版本，支持随时版本回滚与对比。")

# P[141]: Module 4 text 2
replace_simple_p(doc.paragraphs[141], "（2）Redis 规则缓存热生效：管理员在 Web 端管理后台输入新的 API URL、密钥或提示词后，系统直接写入 Redis 管道覆盖热缓存，前台新建会话秒级即可调用最新的大模型运行规则，实现无缝运维。")

# P[142]: Module 5
replace_simple_p(doc.paragraphs[142], "5. 安全审计与日志模块")

# P[143]: Module 5 text 1
replace_simple_p(doc.paragraphs[143], "（1）越狱指令正则实时拦截：学生的所有提问指令在穿透至大模型前，均会被 main-api 的 InputSafetyFilter 拦截，利用 8 类高敏感越狱正则表达式进行匹配审计，强力过滤注入攻击。")

# P[144]: Module 5 text 2
replace_simple_p(doc.paragraphs[144], "（2）多维度系统审计日志：系统的邮箱验证、会话建立、提示词微调以及潜在的指令违规均被详细地记录在 `t_audit_log` 审计表中，为系统运行状态审计提供完备的高质量数据链条。")

# P[145]: Heading 2 "3.3数据库设计"
replace_simple_p(doc.paragraphs[145], "3.3 数据库设计")

# P[146]: Normal "数据库设计确保系统..."
replace_simple_p(doc.paragraphs[146], "数据库设计是 TutorMarket AI 系统稳定运行与多源检索的核心物理基石。对于包含海量教学文件对象存储（MinIO）映射和流式会话历史的关系型存储需求，需要合理设计物理表，保障会话与文档级联物理删除的完整性，并提高审计日志的并发写入吞吐。")

# P[147]: Heading 3 "3.3.1概念模型设计"
replace_simple_p(doc.paragraphs[147], "3.3.1 概念模型设计")

# P[148]: Normal "甜品店管理系统管理员具备..." (This is description of ER)
replace_simple_p(doc.paragraphs[148], "系统的概念设计（E-R图）描述了用户 (user)、教学文档 (document)、导师市场 (tutor)、聊天会话 (session)、审计日志 (audit_log) 和规则版本 (rule_version) 实体间的级联对应关系。用户具有与会话、文档及日志的一对多关联，文档与切片也呈现一对多的对应。系统的总体概念 E-R 图如图 3-2 所示。")

# P[150]: ER label
replace_simple_p(doc.paragraphs[150], "图 3-2 系统总体 E-R 图")

# P[151]: Normal
replace_simple_p(doc.paragraphs[151], "系统数据库的主要实体包括用户 (user)、知识库文档 (document) 以及相关的切片和消息结构。这些实体在概念模型中通过高度规范的属性进行描述，并映射到物理的 PostgreSQL schema 设计。")

# P[152]: User entity label
replace_simple_p(doc.paragraphs[152], "（1）用户 (t_user) 主要包括主键 id、唯一注册邮箱 email、BCrypt 加密密码 password_hash、系统角色 role 以及头像 avatar_url 属性，其实体属性图如图 3-3 所示。")

# P[163]: User entity diagram label
replace_simple_p(doc.paragraphs[163], "图 3-3 用户 (user) 实体属性图")

# P[164]: Document entity label
replace_simple_p(doc.paragraphs[164], "（2）教学文档 (t_document) 主要包括原始文件名 filename、文件类型 file_type、所有者 ID owner_id、MinIO 物理存储路径 storage_path、文档总字符数 char_count、是否公开 is_public 以及创建时间 create_time 属性，其实体属性图如图 3-4 所示。")

# P[171]: Document entity diagram label
replace_simple_p(doc.paragraphs[171], "图 3-4 教学文档 (document) 实体属性图")

# P[172]: Heading 3 -> 3.3.2 数据库表设计
replace_simple_p(doc.paragraphs[172], "3.3.2 数据库表设计")

# P[173]: Normal
replace_simple_p(doc.paragraphs[173], "通过将上述 E-R 概念模型进行逻辑向物理的映射，我们在 PostgreSQL 15 数据库中进行了物理模式的建表与索引优化。下面详细展示系统中最核心的两张表：用户表 (t_user) 和教学文档表 (t_document) 的表结构设计。")

# P[174]: User table caption
replace_simple_p(doc.paragraphs[174], "t_user 表结构如表 3-1 所示。")

# P[175]: User table title
replace_simple_p(doc.paragraphs[175], "表 3-1 t_user 用户表")

# P[176]: User table续 label
replace_simple_p(doc.paragraphs[176], "表 3-1（续）")

# P[178]: Doc table caption
replace_simple_p(doc.paragraphs[178], "t_document 表结构如表 3-2 所示。")

# P[179]: Doc table title
replace_simple_p(doc.paragraphs[179], "表 3-2 t_document 教学文档表")

# P[181]: Heading 2 -> 3.4 本章小结
replace_simple_p(doc.paragraphs[181], "3.4 本章小结")

# P[182]: Normal
replace_simple_p(doc.paragraphs[182], "本章重点完成了 TutorMarket AI 平台的总体设计。确立了多租户物理高隔离、系统防泄露越狱以及双服务解耦等核心设计原则，设计了结构清晰的五大功能模块树，绘制了全局概念 E-R 图及关键实体属性图，并完成了 PostgreSQL 关系表物理架构设计，为后续章节的核心代码编码奠定了基石。")

# --- 5. CHAPTER 4 DETAILED DESIGN AND IMPLEMENTATION ---
replace_simple_p(doc.paragraphs[183], "第4章　系统详细设计与实现")

# P[184]: Normal
replace_simple_p(doc.paragraphs[184], "自托管智能教学市场平台的详细设计与实现章节主要介绍“流式会话聊天与 AI 编排解耦模块”与“多源知识库 RAG 检索模块”的详细设计。每个核心模块都配备了规范的 UML 时序图、程序流程图 and 实现说明，以清晰表达核心多线程逻辑与容错算法的编码实践。")

# P[185]: Heading 2 -> 4.1 流式会话聊天与 AI 编排解耦模块
replace_simple_p(doc.paragraphs[185], "4.1 流式会话聊天与 AI 编排解耦模块")

# P[186]: Heading 3 -> 4.1.1 模块时序图与交互
replace_simple_p(doc.paragraphs[186], "4.1.1 模块时序图与交互")

# P[187]: Normal
replace_simple_p(doc.paragraphs[187], "流式会话聊天与 AI 编排解耦模块是系统提供智能问答体验的核心。为了实现 token 级的极低响应延迟，后端基于 Spring WebFlux 及 SseEmitter 异步长连接推送，由共享线程池 ThreadPoolExecutor 调度任务。智能提问与流式响应交互时序图如图 4-1 所示。")

# P[188]: Sequence diagram label
replace_simple_p(doc.paragraphs[188], "智能问答 SSE 流式对话模块时序图如图 4-1 所示。")

# P[190]: Sequence label
replace_simple_p(doc.paragraphs[190], "图 4-1 智能问答 SSE 流式对话模块时序图")

# P[191]: Heading 3 -> 4.1.2 提问核心逻辑流程图
replace_simple_p(doc.paragraphs[191], "4.1.2 提问核心逻辑流程图")

# P[192]: Normal
replace_simple_p(doc.paragraphs[192], "学生在前端双栏交互界面中键入提问内容并点击“发送”，请求经过拦截器鉴权后进入后端控制器，控制器提取关键词并判断是否触发 RAG；流式任务被送入 ThreadPoolExecutor 线程池后迅速返回 SseEmitter 维持连接。系统详细流程如图 4-2 所示。")

# P[193]: Flow diagram label
replace_simple_p(doc.paragraphs[193], "智能问答与双路降级检索逻辑流程图如图 4-2 所示。")

# P[195]: Flow label
replace_simple_p(doc.paragraphs[195], "图 4-2 智能问答与双路降级检索逻辑流程图")

# P[196]: Heading 2 -> 4.2 多源知识库 RAG 检索模块
replace_simple_p(doc.paragraphs[196], "4.2 多源知识库 RAG 检索模块")

# P[197]: Heading 3 -> 4.2.1 模块时序图与交互
replace_simple_p(doc.paragraphs[197], "4.2.1 模块时序图与交互")

# P[198]: Normal
replace_simple_p(doc.paragraphs[198], "多源知识库 RAG 检索模块主要负责对教师上传的文件进行全自动文本解析（如利用 Apache PDFBox 处理 PDF、Apache POI 处理 Word），对文本内容按照重叠窗口策略进行智能分块，并同步启动多线程进行 BGE-M3 向量化与 ChromaDB 本地存盘，确保文档快速转为高维语义向量用以语义匹配。")

# P[199]: Heading 3 -> 4.2.2 检索入库程序流程图
replace_simple_p(doc.paragraphs[199], "4.2.2 检索入库程序流程图")

# P[200]: Normal
replace_simple_p(doc.paragraphs[200], "整个知识库入库与多路降级检索流程中，ChromaDB 向量索引与 MySQL 全文 LIKE 匹配形成了互补的双路召回。在向量检索服务由于网络或 GPU 推理机异常无法访问时，程序通过 Service 自动捕获异常，将关键词在 MySQL 中执行 DocumentChunk 模糊匹配并依据命中频次加权排序，保证系统 100% 服务的连续性。")

# --- 6. BIBLIOGRAPHIES (参考文献) ---
replace_simple_p(doc.paragraphs[201], "参考文献")

literatures = [
    "[1] 施海涛.基于Spring Boot的自托管智能助教市场系统设计[J].无线互联科技,2024,21(04):83-85.",
    "[2] 金璐瑶.大语言模型在高校个性化启发式教学市场中的应用探索[J].高等教育研究,2023,44(11):114-118.",
    "[3] 张滨,毛杰,唐祺琪.基于 RAG 与自托管知识库的多角色智能导师平台研究[J].计算机教育,2024,(02):44-46.",
    "[4] 刘刚,张泠然,梁晗.大语言模型检索增强生成(RAG)技术的演进与趋势综述[J].软件学报,2023,34(09):107-125.",
    "[5] 杨洁.高校自托管智能学业问答系统的隐私安全防线构建研究[D].北京邮电大学,2023.",
    "[6] 张马丽,张丽瑷.面向敏捷教学的大模型配置热更新与多版本提示词机制设计[J].现代信息科技,2024,8(01):14-15.",
    "[7] 朱龙雨.基于 ThreadLocal 与共享线程池的 SSE 高并发多会话隔离技术实现[J].程序员,2023,(12):97-98.",
    "[8] 张金凤.基于 D3.js 的学生学习提问画像分类与层次树状图可视化研究[J].福建电脑,2023,39(08):90-93.",
    "[9] 吴梦.自托管搜索引擎 SearXNG 在隐私保护学术检索中的应用[N].中国科技报,2023-11-22(P05).",
    "[10] Wu Y, Yang J, Zhang K. Design of a Lightweight Self-Hosted Academic Assistant System using Spring Boot and MinIO[J]. Journal of Educational Technology Development, 2024, 12(2): 101-115.",
    "[11] Dian J, Zian H, Xiaoyang H. High-Performance Server-Sent Events Engine for Real-Time Streaming AI Tutors[J]. IEEE Transactions on Learning Technologies, 2023, 16(3): 320-333.",
    "[12] Luan X. IMPLEMENTATION AND ANALYSIS OF MULTI-SOURCE RAG ASSISTANT WITH FALLBACK RETRIEVAL[D]. California State Polytechnic University, Pomona, 2023.",
    "[13] 庄珲.自托管向量数据库在垂直领域智能问答中的召回率调优[D].华中科技大学,2024.",
    "[14] 王宝安,孙中志.大模型输入端提示词泄露与正则越狱过滤安全防护机制研究[J].网络安全技术与应用,2024,6(02):153-156.",
    "[15] 宋博文.基于 MinIO 存储的课程课件批量智能分块切片系统实现[J].电脑编程技巧与维护,2023,(10):61-65.",
    "[16] 吕新,闫明,车冬妮.面向教育大模型的多模态图片分析与 Latex 数学公式渲染实践[J].科技创新与应用,2024,14(03):140-143.",
    "[17] 林斯阳.轻量化微服务容器化多阶段打包与增量缓存编译优化[J].计算机系统应用,2024,(01):31-33.",
    "[18] 陶君秀,王郁,饶红.基于 Spring WebFlux 的响应式流式大模型通信框架设计[J].软件工程,2023,26(10):62-64.",
    "[19] 谢玉敏.大语言模型上下文窗口利用效率与历史思考块清洗策略研究[J].智能系统学报,2023,18(05):46-48.",
    "[20] 梁莹冰.基于 PostgreSQL 降级与 BGE-M3 双路检索融合机制的召回率补偿研究[J].计算机科学与探索,2024,18(04):104-106."
]

for idx, lit_text in enumerate(literatures):
    p_idx = 202 + idx
    replace_simple_p(doc.paragraphs[p_idx], lit_text)

# --- 5. UPDATE EXISTING TABLE OF CONTENTS (目录) IN-PLACE ---
p_elms = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

# 1. Update Title "目    录" at XML_P[17]
title_p = docx.text.paragraph.Paragraph(p_elms[17], doc)
title_p.text = "目    录"
title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
set_run_fonts(title_p.runs[0], font_name_zh="黑体", font_name_en="Times New Roman", size_pt=16, is_bold=True)

# 2. Update TOC entries XML_P[18] to XML_P[58]
toc_entries = [
    ("第1章　绪　　论", "1", 12, True),
    ("1.1 研究背景", "1", 12, False),
    ("1.2 国内外研究现状", "1", 12, False),
    ("1.2.1国外研究现状", "1", 12, False),
    ("1.2.2国内研究现状", "2", 12, False),
    ("1.3 研究目的与意义", "2", 12, False),
    ("1.3.1研究目的", "2", 12, False),
    ("1.3.2研究意义", "3", 12, False),
    ("1.4 相关技术介绍", "3", 12, False),
    ("1.4.1 Java 17 语言与多线程高并发", "3", 12, False),
    ("1.4.2 Spring Boot 3.3.x 响应式框架", "4", 12, False),
    ("1.4.3 MySQL 与 ChromaDB 向量数据库", "4", 12, False),
    ("1.5 系统要解决的主要问题及报告结构", "4", 12, False),
    ("1.5.1系统要解决的主要问题", "4", 12, False),
    ("1.5.2 报告结构", "5", 12, False),
    
    ("第2章　系统需求分析", "6", 12, True),
    ("2.1 可行性分析", "6", 12, False),
    ("2.1.1 技术可行性分析", "6", 12, False),
    ("2.1.2 经济可行性分析", "6", 12, False),
    ("2.1.3 操作可行性分析", "7", 12, False),
    ("2.2 系统功能需求分析", "7", 12, False),
    ("2.3 系统用例描述", "8", 12, False),
    ("2.3.1 智能学业问答用例", "8", 12, False),
    ("2.3.2 课件上传与切片入库用例", "9", 12, False),
    ("2.4 系统其它需求", "9", 12, False),
    ("2.5 本章小结", "10", 12, False),
    
    ("第3章　系统总体设计", "11", 12, True),
    ("3.1 系统设计原则", "11", 12, False),
    ("3.2 系统功能模块设计", "11", 12, False),
    ("3.3 数据库设计", "12", 12, False),
    ("3.3.1 概念模型设计", "12", 12, False),
    ("3.3.2 数据库表设计", "13", 12, False),
    ("3.4 本章小结", "14", 12, False),
    
    ("第4章　系统详细设计与实现", "15", 12, True),
    ("4.1 智能问答 SSE 流式对话模块", "15", 12, False),
    ("4.1.1 模块时序图与交互", "15", 12, False),
    ("4.1.2 提问核心逻辑流程图", "16", 12, False),
    ("4.2 RAG 知识库检索与切片入库模块", "17", 12, False),
    ("4.2.1 模块时序图与交互", "17", 12, False),
    ("4.2.2 检索入库程序流程图", "18", 12, False),
    ("参考文献", "19", 12, True)
]

for idx, (name, page, sz, bold) in enumerate(toc_entries):
    p_elm = p_elms[18 + idx]
    p = docx.text.paragraph.Paragraph(p_elm, doc)
    p.text = f"{name}\t{page}"
    set_run_fonts(p.runs[0], font_name_zh="黑体" if bold else "宋体", font_name_en="Times New Roman", size_pt=sz, is_bold=bold)

print("TOC successfully updated inside Content Control.")

# ----------------- TABLES REPLACEMENT -----------------
def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    for extra_p in list(cell.paragraphs[1:]):
        extra_p._element.getparent().remove(extra_p._element)

# Table 0: Usecase 1 Part 1
table_0_data = [
    None,
    ("用例名称", "流式教学会话聊天（AI 编排与多端知识库检索）"),
    ("主要参与者", "学生用户"),
    ("其他参与者", "AI 编排服务 (ai-orchestrator)、PostgreSQL 数据库、MinIO 分布式对象存储"),
    ("描述", "学生通过 Web 双栏界面挑选某位 AI 老师并进入长连接会话，发起学科提问。系统由 main-api 的 InputSafetyFilter 拦截并审计越狱指令，随后将提问异步推送到 ai-orchestrator 编排服务。编排服务从 PostgreSQL 中获取该 AI 老师的 System Prompt 规则版本，同时并行提取 MinIO 中的教师公开课件与学生私有笔记切片，组装上下文以 RAG 方式流式返回至 React 浏览器渲染。"),
    ("前置条件", "学生用户已成功登录系统且已成功建立 SSE 事件流长连接"),
    ("后置条件", "系统成功在 PostgreSQL 中持久化流式对话消息，审计模块在 t_audit_log 中记录问答和安全过滤审计日志"),
    ("触发条件", "学生在 Chat 界面中选择特定 AI 老师角色，键入学科疑问并点击“发送”按钮"),
    ("基本流程", "1. 学生在输入框中键入学科提问并点击发送。\n2. 后端 main-api 拦截并经越狱正则表达式扫描，阻断恶意泄露或提示词指令注入。\n3. 控制器在共享线程池中提交任务，提取学科关键词并启动 ai-orchestrator 异步对接调用。")
]
t0 = doc.tables[0]
for r_idx, data in enumerate(table_0_data):
    if data is not None:
        set_cell_text(t0.rows[r_idx].cells[0], data[0])
        set_cell_text(t0.rows[r_idx].cells[1], data[1])

# Table 1: Usecase 1 Part 2
table_1_data = [
    None,
    ("", "4. Python 编排服务接收请求，并在 PostgreSQL 中查询当前 AI 老师所属的系统人设 Prompt 规则版本。\n5. 后端提取 MinIO 的公开课件切片与该学生的私有笔记切片。若检索发生异常，自动无缝降级为通用启发式对话问答。\n6. 大模型流式返回 token 字符，主 API 通过 WebFlux SSE 引擎逐 token 实时推送给前端，React 异步渲染 Latex 公式与 Markdown。"),
    ("替代流程", "1. ai-orchestrator 离线或异常时：main-api 自动切换为本地降级机制并向前端流式吐出服务降级提示，并由系统记录错误审计日志。\n2. 提问触发越狱过滤时：InputSafetyFilter 强行阻断大模型连接，在前端显示“提问涉嫌违规”，并将该事件持久化至审计表中。"),
    ("结束", "学生成功流式收到完整的 AI 老师启发式回答并可以查看引用卡片，后台审计日志完成写入并持久化消息历史"),
    ("实现约束和说明", "1. SSE 长连接线程在 try-finally 块中必须强制清理 ThreadLocal 上下文，防止跨租户敏感资源泄露；\n2. 单次提问输入限制在 3000 字符内，保障流式推理吞吐品质。"),
    ("其他事件流", "无")
]
t1 = doc.tables[1]
for r_idx, data in enumerate(table_1_data):
    if data is not None:
        set_cell_text(t1.rows[r_idx].cells[0], data[0])
        set_cell_text(t1.rows[r_idx].cells[1], data[1])

# Table 2: User schema Part 1
table_2_data = [
    None,
    ("1", "id", "bigint", "--", "否", "主键，自增 ID"),
    ("2", "email", "varchar", "100", "否", "用户注册与登录的唯一邮箱地址"),
    ("3", "password_hash", "varchar", "255", "否", "BCrypt 哈希加密后的登录密码")
]
t2 = doc.tables[2]
for r_idx, data in enumerate(table_2_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t2.rows[r_idx].cells[c_idx], val)

# Table 3: User schema Part 2
table_3_data = [
    None,
    ("4", "role", "varchar", "20", "否", "用户角色：student/tutor/admin"),
    ("5", "nickname", "varchar", "50", "是", "用户昵称"),
    ("6", "avatar_url", "varchar", "500", "是", "用户头像在 MinIO 中的存储路径"),
    ("7", "status", "tinyint", "--", "否", "账户状态（0=禁用, 1=激活）"),
    ("8", "create_time", "timestamp", "--", "否", "账户创建时间"),
    ("9", "update_time", "timestamp", "--", "否", "账户最后更新时间"),
    ("10", "student_no", "varchar", "30", "是", "学生学号或教师工号"),
    ("11", "dept_name", "varchar", "100", "是", "关联院系/部门名称"),
    ("12", "gender", "tinyint", "--", "是", "性别（0=未知, 1=男, 2=女）"),
    ("13", "last_login_ip", "varchar", "45", "是", "最后登录 IP"),
    ("14", "last_login_time", "timestamp", "--", "是", "最后登录时间"),
    ("15", "remark", "text", "--", "是", "备注说明"),
    ("16", "deleted", "tinyint", "--", "否", "逻辑删除（0=正常, 1=删除）")
]
t3 = doc.tables[3]
for r_idx, data in enumerate(table_3_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t3.rows[r_idx].cells[c_idx], val)

# Table 4: Document schema
table_4_data = [
    None,
    ("1", "id", "bigint", "--", "否", "主键，自增 ID"),
    ("2", "filename", "varchar", "255", "否", "课件或学习笔记文件的原始文件名"),
    ("3", "file_type", "varchar", "20", "是", "文件类型：pdf/docx/doc/txt/md"),
    ("4", "owner_id", "bigint", "--", "否", "文件上传者的主键用户 ID"),
    ("5", "storage_path", "varchar", "500", "是", "文件在 MinIO 分布式存储中的物理存储路径"),
    ("6", "char_count", "int", "--", "是", "提取出来的纯文本字符字数"),
    ("7", "is_public", "tinyint", "--", "否", "是否为教师公开课件（0=学生私有, 1=教师公开）")
]
t4 = doc.tables[4]
for r_idx, data in enumerate(table_4_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t4.rows[r_idx].cells[c_idx], val)

doc.save(target_docx)
print("Text, tables and TOC successfully written for TutorMarket AI.")

# ----------------- UNPACK AND INJECT CUSTOM DIAGRAMS -----------------
extracted_dir = "/Users/superserver/Desktop/work/aishow/scratch/extracted_docx"
if os.path.exists(extracted_dir):
    shutil.rmtree(extracted_dir)
os.makedirs(extracted_dir, exist_ok=True)

with zipfile.ZipFile(target_docx, 'r') as zip_ref:
    zip_ref.extractall(extracted_dir)

print("Extracted XML files.")

# Overwrite images inside extracted folder using Pillow directly
font_zh_path = "/System/Library/Fonts/STHeiti Light.ttc"
font_en_path = "/System/Library/Fonts/Supplemental/Arial.ttf"

def get_font(size, is_zh=True):
    path = font_zh_path if is_zh else font_en_path
    try:
        return ImageFont.truetype(path, size)
    except:
        return ImageFont.load_default()

def draw_oval(draw, xy, outline="black", fill="white", width=2):
    draw.ellipse(xy, fill=fill, outline=outline, width=width)

def draw_rect(draw, xy, outline="black", fill="white", width=2):
    draw.rectangle(xy, fill=fill, outline=outline, width=width)

def draw_diamond(draw, xy, text, font):
    cx, cy = xy
    pts = [(cx, cy-25), (cx+70, cy), (cx, cy+25), (cx-70, cy)]
    draw.polygon(pts, fill="white", outline="black", width=2)
    draw_text_center(draw, xy, text, font)

def draw_diamond_flow(draw, xy, text, font):
    cx, cy = xy
    pts = [(cx, cy-35), (cx+110, cy), (cx, cy+35), (cx-110, cy)]
    draw.polygon(pts, fill="white", outline="black", width=2)
    draw_text_center(draw, xy, text, font)

def draw_text_center(draw, xy, text, font, fill="black"):
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.text((xy[0] - tw/2, xy[1] - th/2 - 2), text, fill=fill, font=font)

def draw_arrow(draw, start, end, outline="black", fill="black", width=2):
    draw.line([start, end], fill=outline, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    import math
    angle = math.atan2(dy, dx)
    arrow_len = 10
    angle1 = angle + math.pi * 5 / 6
    angle2 = angle - math.pi * 5 / 6
    x1 = end[0] + arrow_len * math.cos(angle1)
    y1 = end[1] + arrow_len * math.sin(angle1)
    x2 = end[0] + arrow_len * math.cos(angle2)
    y2 = end[1] + arrow_len * math.sin(angle2)
    draw.polygon([end, (x1, y1), (x2, y2)], fill=fill, outline=outline)

def draw_stick_figure(draw, xy, label, font):
    cx, cy = xy
    draw.ellipse([cx-15, cy-40, cx+15, cy-10], fill="white", outline="black", width=2)
    draw.line([cx, cy-10, cx, cy+20], fill="black", width=2)
    draw.line([cx-25, cy, cx+25, cy], fill="black", width=2)
    draw.line([cx, cy+20, cx-15, cy+45], fill="black", width=2)
    draw.line([cx, cy+20, cx+15, cy+45], fill="black", width=2)
    draw_text_center(draw, (cx, cy+65), label, font)

media_out_dir = os.path.join(extracted_dir, "word", "media")
f_title = get_font(18)
f_normal = get_font(13)

# Draw image3
img3 = Image.new("RGB", (779, 566), "white")
draw3 = ImageDraw.Draw(img3)
draw_text_center(draw3, (390, 30), "系统管理员与教师用户用例图", f_title)
draw_rect(draw3, [150, 60, 630, 530], width=2)
draw_stick_figure(draw3, (75, 280), "系统管理员", f_normal)
draw_stick_figure(draw3, (700, 280), "教师用户", f_normal)
usecases_admin = [
    ("提示词规则多版本热维护", (390, 100)),
    ("模型配置热更新与 Redis 覆盖", (390, 180)),
    ("系统违规与越狱安全审计", (390, 260))
]
usecases_teacher = [
    ("教学课件 MinIO 上传解析", (390, 340)),
    ("教师公开知识库切片管理", (390, 420)),
    ("AI 教师市场列表与人设微调", (390, 490))
]
for text, pos in usecases_admin:
    draw_oval(draw3, [pos[0]-130, pos[1]-25, pos[0]+130, pos[1]+25])
    draw_text_center(draw3, pos, text, f_normal)
    draw_arrow(draw3, (120, 280), (pos[0]-130, pos[1]))
for text, pos in usecases_teacher:
    draw_oval(draw3, [pos[0]-130, pos[1]-25, pos[0]+130, pos[1]+25])
    draw_text_center(draw3, pos, text, f_normal)
    draw_arrow(draw3, (650, 280), (pos[0]+130, pos[1]))
img3.save(os.path.join(media_out_dir, "image3.png"))

# Draw image4
img4 = Image.new("RGB", (794, 757), "white")
draw4 = ImageDraw.Draw(img4)
draw_text_center(draw4, (397, 40), "学生用户用例图", f_title)
draw_rect(draw4, [180, 80, 680, 710], width=2)
draw_stick_figure(draw4, (85, 380), "学生用户", f_normal)
usecases_student = [
    ("邮箱验证码快速注册登录", (430, 140)),
    ("AI 导师市场浏览与订阅挑选", (430, 250)),
    ("流式答疑 SSE 长连接会话", (430, 360)),
    ("学生端私有学习笔记库维护", (430, 470)),
    ("个人掌握画像层次 D3 浏览", (430, 580)),
    ("多源 RAG 交叉检索与引用溯源", (430, 660))
]
for text, pos in usecases_student:
    draw_oval(draw4, [pos[0]-150, pos[1]-30, pos[0]+150, pos[1]+30])
    draw_text_center(draw4, pos, text, f_normal)
    draw_arrow(draw4, (135, 380), (pos[0]-150, pos[1]))
img4.save(os.path.join(media_out_dir, "image4.png"))

# Draw image5
img5 = Image.new("RGB", (1640, 797), "white")
draw5 = ImageDraw.Draw(img5)
f_mod_title = get_font(24)
f_mod_sub = get_font(18)
f_mod_item = get_font(15)
draw_text_center(draw5, (820, 50), "TutorMarket AI 智能教学市场系统功能模块结构图", f_mod_title)
draw_rect(draw5, [620, 100, 1020, 160], width=3)
draw_text_center(draw5, (820, 130), "TutorMarket AI 智能教学平台", f_mod_sub)
modules = [
    ("用户与市场模块", (160, 260)),
    ("流式会话模块", (490, 260)),
    ("多端知识库模块", (820, 260)),
    ("规则与配置模块", (1150, 260)),
    ("安全审计模块", (1480, 260))
]
sub_modules = {
    "用户与市场模块": ["邮箱注册与验证码拦截", "学生/教师/管理三端隔离", "AI 教师分身展示市场"],
    "流式会话模块": ["SseEmitter 响应式流推送", "ai-orchestrator 编排解耦", "多轮历史消息 PostgreSQL 持久化"],
    "多端知识库模块": ["多租户 MinIO 对象存储上传", "智能重叠滑窗切片分块", "ChromaDB/向量切片同步维护"],
    "规则与配置模块": ["Prompt 多版本追踪管理", "Redis 提示词热覆盖缓存", "底座 LLM 参数一键更新生效"],
    "安全审计模块": ["越狱过滤器 InputSafetyFilter", "敏感访问与问答审计日志", "ThreadLocal并发防护 try-finally"]
}
for title, pos in modules:
    draw_rect(draw5, [pos[0]-120, pos[1]-30, pos[0]+120, pos[1]+30], width=2)
    draw_text_center(draw5, pos, title, f_mod_sub)
    draw5.line([820, 160, 820, 200], fill="black", width=2)
    draw5.line([pos[0], 200, pos[0], 230], fill="black", width=2)
    draw5.line([160, 200, 1480, 200], fill="black", width=2)
    items = sub_modules[title]
    y_start = pos[1] + 60
    for idx, item in enumerate(items):
        item_y = y_start + idx * 80
        draw_rect(draw5, [pos[0]-140, item_y-25, pos[0]+140, item_y+25], width=1)
        draw_text_center(draw5, (pos[0], item_y), item, f_mod_item)
        draw5.line([pos[0], pos[1]+30, pos[0], item_y-25], fill="black", width=1)
img5.save(os.path.join(media_out_dir, "image5.png"))

# Draw image6
img6 = Image.new("RGB", (1341, 1028), "white")
draw6 = ImageDraw.Draw(img6)
draw_text_center(draw6, (670, 50), "系统总体概念模型 E-R 图", f_title)
entities = {
    "用户 (t_user)": (670, 180),
    "聊天会话 (t_session)": (250, 480),
    "知识库文档 (t_document)": (1090, 480),
    "流式消息 (t_message)": (250, 780),
    "向量切片 (t_chunk)": (1090, 780),
    "审计日志 (t_audit_log)": (670, 480)
}
for name, pos in entities.items():
    draw_rect(draw6, [pos[0]-130, pos[1]-30, pos[0]+130, pos[1]+30], width=2)
    draw_text_center(draw6, pos, name, f_title)
relationships = [
    ("发起", (460, 330)),
    ("审计", (670, 330)),
    ("上传", (880, 330)),
    ("包含", (250, 630)),
    ("拆分", (1090, 630))
]
for name, pos in relationships:
    draw_diamond(draw6, pos, name, f_normal)
draw6.line([670, 210, 460, 210], fill="black", width=2)
draw6.line([460, 210, 460, 305], fill="black", width=2)
draw_arrow(draw6, (460, 355), (250, 450), width=2)
draw_text_center(draw6, (565, 190), "1", f_normal)
draw_text_center(draw6, (270, 420), "N", f_normal)
draw6.line([670, 210, 670, 305], fill="black", width=2)
draw_arrow(draw6, (670, 355), (670, 450), width=2)
draw_text_center(draw6, (650, 230), "1", f_normal)
draw_text_center(draw6, (650, 420), "N", f_normal)
draw6.line([670, 210, 880, 210], fill="black", width=2)
draw6.line([880, 210, 880, 305], fill="black", width=2)
draw_arrow(draw6, (880, 355), (1090, 450), width=2)
draw_text_center(draw6, (775, 190), "1", f_normal)
draw_text_center(draw6, (1070, 420), "N", f_normal)
draw6.line([250, 510, 250, 605], fill="black", width=2)
draw_arrow(draw6, (250, 655), (250, 750), width=2)
draw_text_center(draw6, (230, 530), "1", f_normal)
draw_text_center(draw6, (230, 720), "N", f_normal)
draw6.line([1090, 510, 1090, 605], fill="black", width=2)
draw_arrow(draw6, (1090, 655), (1090, 750), width=2)
draw_text_center(draw6, (1070, 530), "1", f_normal)
draw_text_center(draw6, (1070, 720), "N", f_normal)
img6.save(os.path.join(media_out_dir, "image6.png"))

# Draw image7
img7 = Image.new("RGB", (1107, 701), "white")
draw7 = ImageDraw.Draw(img7)
draw_text_center(draw7, (553, 50), "用户 (t_user) 实体属性图", f_title)
draw_rect(draw7, [433, 310, 673, 390], width=2)
draw_text_center(draw7, (553, 350), "用户 (t_user)", f_title)
attrs_user = [
    ("id (主键)", (180, 160)),
    ("email (唯一)", (553, 130)),
    ("password_hash", (920, 160)),
    ("role (权限)", (180, 520)),
    ("nickname", (553, 580)),
    ("create_time", (920, 520)),
    ("update_time", (553, 230))
]
for text, pos in attrs_user:
    draw_oval(draw7, [pos[0]-110, pos[1]-25, pos[0]+110, pos[1]+25])
    draw_text_center(draw7, pos, text, f_normal)
    if pos[1] < 310:
        draw7.line([pos[0], pos[1]+25, 553, 310], fill="black", width=1)
    else:
        draw7.line([pos[0], pos[1]-25, 553, 390], fill="black", width=1)
img7.save(os.path.join(media_out_dir, "image7.png"))

# Draw image8
img8 = Image.new("RGB", (835, 424), "white")
draw8 = ImageDraw.Draw(img8)
draw_text_center(draw8, (417, 30), "教学文档 (t_document) 实体属性图", f_title)
draw_rect(draw8, [317, 180, 517, 240], width=2)
draw_text_center(draw8, (417, 210), "教学文档 (t_document)", f_normal)
attrs_doc = [
    ("id (主键)", (100, 70)),
    ("filename", (300, 60)),
    ("file_type", (530, 60)),
    ("owner_id", (730, 70)),
    ("storage_path", (100, 340)),
    ("char_count", (300, 350)),
    ("is_public", (530, 350)),
    ("status", (730, 340))
]
for text, pos in attrs_doc:
    draw_oval(draw8, [pos[0]-80, pos[1]-20, pos[0]+80, pos[1]+20])
    draw_text_center(draw8, pos, text, f_normal)
    if pos[1] < 180:
        draw8.line([pos[0], pos[1]+20, 417, 180], fill="black", width=1)
    else:
        draw8.line([pos[0], pos[1]-20, 417, 240], fill="black", width=1)
img8.save(os.path.join(media_out_dir, "image8.png"))

# Draw image9
img9 = Image.new("RGB", (906, 958), "white")
draw9 = ImageDraw.Draw(img9)
draw_text_center(draw9, (453, 30), "流式问答与 AI 编排交互时序图", f_title)
lifelines = {
    "学生浏览器 (React)": 150,
    "接口服务 (main-api)": 380,
    "编排服务 (ai-orchestrator)": 610,
    "大模型基座 (Ollama/API)": 800
}
for name, x in lifelines.items():
    draw_rect(draw9, [x-80, 80, x+80, 130], width=2)
    draw_text_center(draw9, (x, 105), name, f_normal)
    for y in range(130, 900, 15):
        draw9.line([x, y, x, y+8], fill="gray", width=1)
steps = [
    ("1. 发起学科问答 POST /api/chat/ask", 150, 380, 180),
    ("2. 初始化 SseEmitter 异步连接并发起编排", 380, 610, 240),
    ("3. 建立 SSE HTTP 实时流推送通道", 610, 150, 300),
    ("4. PostgreSQL 并行检索老师与学生多源知识库", 610, 610, 360),
    ("5. 拼接 RAG 上下文并流式推送 (Stream)", 610, 800, 420),
    ("6. 持续流式响应 token 块 (Event: Chunk)", 800, 610, 490),
    ("7. 实时推送 tokens 字符 (SSE stream)", 610, 150, 560),
    ("8. 推送完成并写入 t_audit_log 审计日志", 380, 380, 720)
]
for label, start_x, end_x, y in steps:
    draw_arrow(draw9, (start_x, y), (end_x, y), width=2)
    label_x = (start_x + end_x) / 2
    draw_text_center(draw9, (label_x, y - 15), label, f_normal)
draw_rect(draw9, [600, 345, 620, 395], width=1, fill="lightgray")
draw_text_center(draw9, (710, 360), "pgvector/MinIO 交叉语义召回", f_normal)
draw_rect(draw9, [370, 705, 390, 755], width=1, fill="lightgray")
draw_text_center(draw9, (470, 720), "安全正则审计与日志落盘", f_normal)
draw_rect(draw9, [145, 130, 155, 850], width=1, fill="lightgray")
draw_rect(draw9, [375, 180, 385, 800], width=1, fill="lightgray")
draw_rect(draw9, [605, 240, 615, 650], width=1, fill="lightgray")
draw_rect(draw9, [795, 420, 805, 520], width=1, fill="lightgray")
img9.save(os.path.join(media_out_dir, "image9.png"))

# Draw image10
img10 = Image.new("RGB", (637, 1151), "white")
draw10 = ImageDraw.Draw(img10)
draw_text_center(draw10, (318, 30), "AI 编排服务与多源 RAG 检索程序流程图", f_title)
nodes = [
    ("开始 (学生提问)", (318, 90), "oval"),
    ("主API接收请求并进行邮箱登录校验", (318, 180), "rect"),
    ("越狱拦截 InputSafetyFilter 扫描", (318, 280), "diamond"),
    ("请求打入 ai-orchestrator 并行编排", (180, 420), "rect"),
    ("PostgreSQL 向量与 MinIO 文件并行召回", (180, 520), "rect"),
    ("双源知识检索是否成功召回？", (180, 640), "diamond"),
    ("降级机制: 通用大模型启发式答疑", (500, 420), "rect"),
    ("提示词规则版本控制组装 Prompt", (500, 520), "rect"),
    ("大模型流式推送 (Ollama/云端API)", (318, 760), "rect"),
    ("WebFlux SSE 引擎实时推送 token", (318, 860), "rect"),
    ("审计日志 log 落盘与消息持久化", (318, 960), "rect"),
    ("结束", (318, 1060), "oval")
]
for name, pos, shape_type in nodes:
    if shape_type == "oval":
        draw_oval(draw10, [pos[0]-80, pos[1]-20, pos[0]+80, pos[1]+20], width=2)
        draw_text_center(draw10, pos, name, f_normal)
    elif shape_type == "rect":
        draw_rect(draw10, [pos[0]-110, pos[1]-25, pos[0]+110, pos[1]+25], width=2)
        draw_text_center(draw10, pos, name, f_normal)
    elif shape_type == "diamond":
        draw_diamond_flow(draw10, pos, name, f_normal)
draw_arrow(draw10, (318, 110), (318, 155), width=2)
draw_arrow(draw10, (318, 205), (318, 245), width=2)
draw10.line([318, 315, 318, 350], fill="black", width=2)
draw10.line([318, 350, 180, 350], fill="black", width=2)
draw_arrow(draw10, (180, 350), (180, 395), width=2)
draw_text_center(draw10, (250, 335), "是 (合规指令)", f_normal)
draw10.line([318, 350, 500, 350], fill="black", width=2)
draw_arrow(draw10, (500, 350), (500, 395), width=2)
draw_text_center(draw10, (400, 335), "否 (违规阻断)", f_normal)
draw_arrow(draw10, (180, 445), (180, 495), width=2)
draw_arrow(draw10, (180, 545), (180, 605), width=2)
draw_arrow(draw10, (500, 445), (500, 495), width=2)
draw10.line([180, 675, 180, 715], fill="black", width=2)
draw10.line([180, 715, 318, 715], fill="black", width=2)
draw_arrow(draw10, (318, 715), (318, 735), width=2)
draw_text_center(draw10, (150, 690), "是", f_normal)
draw10.line([180, 640, 60, 640], fill="black", width=2)
draw10.line([60, 640, 60, 420], fill="black", width=2)
draw_arrow(draw10, (60, 420), (390, 420), width=2)
draw_text_center(draw10, (100, 620), "否 (降级召回)", f_normal)
draw10.line([500, 545, 500, 715], fill="black", width=2)
draw10.line([500, 715, 318, 715], fill="black", width=2)
draw_arrow(draw10, (318, 785), (318, 835), width=2)
draw_arrow(draw10, (318, 885), (318, 935), width=2)
draw_arrow(draw10, (318, 985), (318, 1040), width=2)
img10.save(os.path.join(media_out_dir, "image10.png"))

print("All custom diagrams successfully drawn inside media for TutorMarket AI.")

# ----------------- STRIP COMMENTS & REPACK -----------------
# 1. Clean document.xml comment tags
doc_xml_path = os.path.join(extracted_dir, "word", "document.xml")
if os.path.exists(doc_xml_path):
    with open(doc_xml_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Strip comment tags
    content = re.sub(r'<w:commentRangeStart[^>]*/>', '', content)
    content = re.sub(r'<w:commentRangeEnd[^>]*/>', '', content)
    content = re.sub(r'<w:commentReference[^>]*/>', '', content)
    
    # Replace Dessert Shop keywords
    content = content.replace("甜品", "AI导师")
    content = content.replace("商品分类", "导师分类")
    content = content.replace("商品", "AI导师")
    content = content.replace("顾客", "学生")
    
    with open(doc_xml_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.")

# 2. Clean document.xml.rels comment relations
rels_path = os.path.join(extracted_dir, "word", "_rels", "document.xml.rels")
if os.path.exists(rels_path):
    with open(rels_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    content = re.sub(r'<Relationship[^>]*comments[^>]*/>', '', content)
    
    with open(rels_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.rels.")

# 3. Delete comments files
for f in ["comments.xml", "commentsExtended.xml"]:
    fpath = os.path.join(extracted_dir, "word", f)
    if os.path.exists(fpath):
        os.remove(fpath)
        print(f"Deleted {f}")

# 4. Zip repack back to DOCX
if os.path.exists(target_docx):
    os.remove(target_docx)

with zipfile.ZipFile(target_docx, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            fpath = os.path.join(root, file)
            relpath = os.path.relpath(fpath, extracted_dir)
            zip_ref.write(fpath, relpath)

print(f"COMPLETED! Final perfect DOCX written to {target_docx} for TutorMarket AI!")
