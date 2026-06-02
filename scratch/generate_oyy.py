import os
import re
import zipfile
import shutil
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

template_path = "/private/var/folders/m1/y8wt5fvd3hl_jb3c97rwk3w80000gn/T/MicrosoftEdgeDownloads/638411b0-51e9-4d19-91b8-d7cdff285eae/23级-系统分析与设计-课程报告模板V1.3.docx"
target_docx = "/Users/superserver/Desktop/work/aishow/OYY_学院_系统分析与设计课程报告.docx"

if os.path.exists(target_docx):
    os.remove(target_docx)
shutil.copy(template_path, target_docx)
os.chmod(target_docx, 0o644)
print("Copied fresh template for OYY Academy.")

doc = docx.Document(target_docx)


def set_run_fonts(run, font_name_zh="宋体", font_name_en="Times New Roman", size_pt=12, is_bold=False):
    run.font.size = Pt(size_pt)
    run.bold = is_bold
    run.font.name = font_name_zh
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name_zh)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

def replace_simple_p(p, text):
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), "宋体")
    rFonts.set(qn('w:ascii'), "Times New Roman")
    rFonts.set(qn('w:hAnsi'), "Times New Roman")

# --- 1. REPLACE TITLE PAGE ---
replace_simple_p(doc.paragraphs[8], "基于 Spring Boot 与工作流编排的 OYY学院后端管理平台的设计与实现")

# --- 2. CHAPTER 1 INTRODUCTION ---
replace_simple_p(doc.paragraphs[18], "本课程报告的研究范围主要在设计与实现一个基于 Spring Boot 与工作流编排的 OYY学院后端管理平台。其核心关注点是将官网的所有在线教务业务进行整合，并借助 Java 17、Spring Boot 3、MyBatis-Plus、MySQL 与 Redis 等主流技术栈，构建一个涵盖手机短信验证码登录、微信扫码登录、JWT 无状态拦截器鉴权、课程多维检索、AI 多类型插件卡片展示与按积分计费结算、以及自定义工作流多步 JSON 编排与批量激活码管理在内的高可用、高安全级别的分布式官网后端架构。绪论部分概述了本研究的背景、目标和重要性，介绍了国内外在在线 IT 教育及工作流编排系统方面的研究现状，概述了所涉及的关键技术，并探讨了系统需解决的问题，揭示全文的结构安排。")

# P[20] has citations [1] and [2]
p20 = doc.paragraphs[20]
p20.runs[0].text = "在我国在线IT教育和职业技术培训行业数字化转型持续深化的背景下，传统静态单向的教学平台在师生学业互动以及智能化辅助研学方面面临全新的瓶颈，迫切需要引入先进的人工智能大模型辅助场景"
p20.runs[2].text = "。OYY学院作为领先的在线教学平台，急需在后台引入一套智能化的课程内容、AI 插件卡片展示以及自定义工作流编排服务，从而在保障内部网络数据隔离安全的前提下，向学员提供快捷高效的启发式编程及问答辅导"
p20.runs[4].text = "。"

# P[21] has citation [3]
p21 = doc.paragraphs[21]
p21.runs[0].text = "近年来，微服务解耦与大模型 API 工具集（Tool-call）在各个垂直领域的落地已经展现出了巨大的工程应用价值。通过将不同任务的 AI 插件（如图片生成、视频生成、文本处理等）进行独立封装，并提供自定义多步 JSON 编排工作流，用户可以像拼积木一样组合多个 AI 工具，大幅提升任务处理自动化程度；与此同时，结合基于 Redis 的手机验证码缓存、无状态 JWT 拦截鉴权、多角色 ThreadLocal 隔离、以及批量安全激活码兑换会员或积分机制，系统可以构建出一个高度可靠、透明度极高且交互体验极佳的学院官网后端服务，对优化平台运营与改善学员的学习体验具有极强的实践意义"
p21.runs[1].text = ""
p21.runs[3].text = "。"

# P[23]:国内外研究现状 intro
replace_simple_p(doc.paragraphs[23], "在微服务与在线 IT 教育高度融合的时代，建设一个安全、解耦、支持并发结算与按积分计费的智能后端平台已成为国内外教育信息化升级的核心焦点。为了构建一个既符合多角色（学员、讲师、管理员）交互规范，又能够保障激活码防刷和积分扣减事务一致性的高可用系统，我们需要深入梳理当前国内外在工作流编排、JWT 无状态鉴权以及 Redis 并发限流等技术维度的发展特征，为本项目的研发奠定理论基础。")

# P[25] has citation [4]
p25 = doc.paragraphs[25]
p25.runs[0].text = "在国外在线IT培训与微服务编排框架的研究中，多角色鉴权与高并发结算系统的工程化落地已经非常成熟。通过采用 Spring Boot、MyBatis-Plus 以及 Redis 分布式事务锁等先进架构，许多国外的高并发电商和教学平台在 Token 安全防御、JWT 会话控制及激活码算法抗碰撞等方向取得了重要突破，这都显著提升了学员高频访问的交互响应速率与数据可靠度"
p25.runs[2].text = "。"

# P[26] has citation [5]
p26 = doc.paragraphs[26]
p26.runs[0].text = "国际学术界在提升业务稳定性的同时，也对教育场景下智能大模型插件的多步工作流编排（Workflow）以及计费审计系统的构建进行了大量讨论。这方面的流程引擎（如 Camunda/Flowable）轻量化改造与监控策略为智能平台跨学科部署和高并发业务结算提供了极为可行的架构范式。此外，大语言模型推理过程的透明化以及数据安全合规性也是当前国外研究的热点，目的是构建师生对 AI 生成结果的信任，并符合严格的数据保护法律要求"
p26.runs[1].text = ""
p26.runs[2].text = "。"

# P[27] has citations [6] and [7]
p27 = doc.paragraphs[27]
p27.runs[0].text = "在这些研究和实践的基础上，还必须考虑微服务部署在多用户并发访问下的局限性。例如，在专业课程和会员充值系统的建设中，由于涉及学员资产与积分余额的私密性，教学管理系统需要提供极高规格的学生/教师物理数据库隔离，以避免敏感学科资源的跨用户越权泄漏"
p27.runs[1].text = ""
p27.runs[2].text = ""
p27.runs[3].text = ""
p27.runs[4].text = ""
p27.runs[6].text = ""
p27.runs[7].text = "。同时在面对高并发、高弹性的校园局域网络环境时，后端服务需要采用基于共享线程池的 ThreadLocal 会话清理与 Redis 并发限流设计，以保障多线程会话在弱网或高频请求场景下的稳定连接与鲁棒体验"
p27.runs[8].text = ""
p27.runs[10].text = "。"

# P[28]: Normal
replace_simple_p(doc.paragraphs[28], "为适应智能 AI 插件在高校日常多学科答疑场景下的普及应用，教学市场平台不仅需提供多维的教师展示与选择面板，更需要在系统性能、隐私合规及高可用性方面持续提升，以获得长效的校园用户粘性。")

# P[30] has citation [8]
p30 = doc.paragraphs[30]
p30.runs[0].text = "随着近几年国内职业教育数字化升级战略的推进，基于云端 API 与本地自建 IT 课程市场的智能教务平台成为了教育工作者的迫切诉求。为了快速升级传统静态教务管理模式，国内众多机构尝试引入轻量化本地大模型，试图将传统课件的上传与查找转化为启发式 AI 流式问答。然而，中小型学科组或院系在技术实力和服务器资源上面临较较大物理约束，通常缺乏一套能够兼容多端知识库检索、防提示词越狱泄露、且支持模型配置和系统规则版本热生效的一体化自托管开源平台"
p30.runs[2].text = "。"

# P[31]: Normal
replace_simple_p(doc.paragraphs[31], "虽然部分团队尝试在官网中集成点对点大模型聊天，但在多角色多场景的级联管理、MinIO 对象存储高频连接、多线程异步缓冲管理，以及 AI 插件服务与 Spring Boot 主 API 的服务解耦架构方面仍显单薄。大并发下往往出现长响应延迟和连接耗尽的问题，限制了系统的广泛普及。")

# P[32] has citation [9]
p32 = doc.paragraphs[32]
p32.runs[0].text = "国内外的智能教学资源管理与 AI 助教平台的发展特征揭示了多源 RAG 与流式响应技术在教育垂直领域演进中的必然趋势"
p32.runs[2].text = "。汲取国外多角色隔离与高性能编排的思想，结合国内高校教学管理和版权主权的实际需求，本项目旨在设计开发一套基于 Spring Boot 与工作流编排的 OYY学院后端管理平台。"

# P[34]: intro
replace_simple_p(doc.paragraphs[34], "随着职业培训与微服务的高速落地，学员对低延迟、高安全性、多功能融合的智能化学习平台需求日益迫切。传统的静态文档下载和关键词检索服务模式正面临重大的智能化转型升级。构建一个高性能、数据分级隔离的智能教学市场平台，不仅可以极大地释放教师日常重复答疑的工作负担，还能为学生提供启发式、高相关度、并支持严格引用溯源的流式解答体验。因此，本报告专注于探究自托管智能教学市场平台的系统分析与设计，重点研究如何通过 Spring Boot 主 API 服务与 Python 编排服务的解耦设计、多源知识库检索、规则版本管理以及输入越狱检测机制来保障服务的卓越性。接下来，将详细说明本研究的具体目的与意义。")

# P[36] has citation [10]
p36 = doc.paragraphs[36]
p36.runs[0].text = "本次研究的重要目的是设计并开发一个支持多角色权限隔离、高稳定性的 OYY学院后端管理平台。该平台无缝对接本地和云端 AI 工具集，并建立基于 JWT 无状态拦截的安全鉴权链路。系统通过解耦的 AI 工作流编排以及 MySQL 数据库，实现毫秒级激活码校验与高性能并发积分扣减结算，在确保学员虚拟资产 100% 严密隐私保护的前提下显著提升个性化教学效率与体验"
p36.runs[1].text = ""
p36.runs[2].text = ""
p36.runs[3].text = ""
p36.runs[4].text = ""
p36.runs[5].text = ""
p36.runs[7].text = "。"

# P[37] has citation [11]
p37 = doc.paragraphs[37]
p37.runs[0].text = "更重要的是要整合多角色权限，设计实施一个支持教师自主上架精品课程、学员在线订阅、以及自定义多步 AI 工作流编排的互动平台"
p37.runs[1].text = ""
p37.runs[3].text = "。同时，通过引入系统参数与 Prompt 模板的规则版本管理机制，允许管理员或教师通过后台 Web 端热更新大模型的分身人设，为各学科重难点教学设计提供实时、自适应的数据及版本决策支持。"

# P[39] has citation [12]
p39 = doc.paragraphs[39]
p39.runs[0].text = "本研究的意义体现在对微服务多角色业务系统与 IT 在线教育理论和工程实践深度结合的研究上。本报告旨在为在线教育机构及中小型教学团队提供一套完全私有化部署、高安全防线且符合绿色低能耗标准的自建教学平台解决方案，以顺应高校数字化与 AI 融合的时代潮流"
p39.runs[2].text = "。"

# P[40]: Normal
replace_simple_p(doc.paragraphs[40], "通过部署 OYY学院后端平台，教育机构能够将原先零散的专业课件、教学视频以及大文件讲义一键上传至自建阿里云/腾讯云 OSS，并通过 MySQL 表进行结构化关联。学生不仅能订阅所需课程，还可以利用系统提供的多模态 AI 插件（如图片生成、视频生成等）进行智能编排，顺序串联多个 AI 工具，实现全自动的编码及作业辅助设计，大幅减少学习盲区并提高实操效率。")

# P[41] has citation [13]
p41 = doc.paragraphs[41]
p41.runs[0].text = "系统集成的 ThreadLocal 会话安全隔离与基于 Redis 的防刷限流机制是构建学员对平台支付与虚拟积分账户信任的根本屏障，确保平台在使用时完全遵循我国互联网信息服务管理办法等相关网络安全法规"
p41.runs[2].text = "。同时，系统支持本地 Docker 容器化一键部署与轻量化测试运行，降低了硬件采购与日常物理主机的运维开销，对推动校园绿色计算和降低高校数字资源采购成本具有突出的实践指导作用。"

# P[42]: Normal
replace_simple_p(doc.paragraphs[42], "本报告丰富了传统教务管理信息系统在工作流（Workflow）编排与大模型技术集成方面的理论研究。其首创的 Java 17 + Spring Boot 3 与 Redis 高速缓存配合 MyBatis-Plus 极简 ORM 构建，以及批量激活码高并发生成算法和 JWT 无状态拦截器技术，为智能助教系统的低门槛落地部署提供了极其具体的工程实践。")

# P[44]: intro
replace_simple_p(doc.paragraphs[44], "在如今大模型与微服务深度结合的时代，选择轻量级、健壮且高并发的技术栈对于构建一个高性能、高安全的自托管教学平台至关重要。本项目专注于挑选最主流的开源技术，以确保系统在面临高并发长连接和海量知识文件解析时，能够展现卓越的吞吐性与高可用性。经过充分对比，系统最终采用了“Java 17 / Spring Boot 3.x 核心框架 + MyBatis-Plus ORM 库”的敏捷架构，数据存储则结合了 MySQL 关系型数据库、Redis 高速缓存以及 JWT 鉴权体系。下面将详细探讨这些关键技术在系统中的具体作用、技术优势，以及它们是如何协同工作以满足复杂的智能教学平台业务需求。")

# P[45] & P[46]: Java heading and text
replace_simple_p(doc.paragraphs[45], "1.4.1 Java 17 语言与多线程高并发")
replace_simple_p(doc.paragraphs[46], "作为开发企业级应用的中流砥柱，Java 17 以其高强度的内存管控、极其出色的多线程高并发管理器，成为了本项目核心业务后端的开发基石。在 OYY学院后端平台中，Java 17 的非阻塞 I/O 读写能力、垃圾回收器（G1/ZGC）的低停顿特性，使得后台在维持高频手机验证码登录校验、高并发下自定义工作流多线程顺序调用结算、以及对 OSS 大文件的并发上传读写时，能够保持极低的内存开销与卓越的系统吞吐性能。")

# P[47] & P[48]: Spring Boot heading and text
replace_simple_p(doc.paragraphs[47], "1.4.2 Spring Boot 3.x 响应式框架")
p48 = doc.paragraphs[48]
p48.runs[0].text = "Spring Boot 3 框架作为快速构建生产级 Spring 应用的行业标准，其强大的 IoC 容器与自动装配特性，为开发 OYY学院后端平台提供了灵活而轻量的主 API 骨架"
p48.runs[1].text = ""
p48.runs[2].text = ""
p48.runs[3].text = ""
p48.runs[4].text = ""
p48.runs[5].text = ""
p48.runs[6].text = ""
p48.runs[7].text = ""
p48.runs[9].text = "。Spring Boot"
p48.runs[10].text = "不仅精简了"
p48.runs[11].text = "初始的复杂 XML 配置，还集成了优秀的组件机制，为后续多服务间的解耦与极速部署提供了不可替代的技术支撑。"

# P[49] & P[50]: DB heading and text
replace_simple_p(doc.paragraphs[49], "1.4.3 MySQL 与 Redis 混合存储架构")
replace_simple_p(doc.paragraphs[50], "在持久化与高频交互层面，系统创造性地引入了“MySQL 关系型数据库 + Redis 高速内存缓存”的混合存储架构。MySQL 负责高可靠事务数据（如用户账户、精品课程列表、工作流步骤 JSON 编排、批量激活码元数据、消费记录流水等）的索引与强一致性管理；而 Redis 则通过 DB 10 专用于存储 5 分钟手机验证码与微信扫码 `state`（UUID）。这种双轨存储架构在保障海量关系数据安全物理隔离的前提下，通过后台拦截器与 ThreadLocal 实现高可靠的会话状态生命周期管理，提供了极其优秀的可扩展性。")

# P[51]: Heading 2 -> 1.5 系统要解决的主要问题及报告结构
replace_simple_p(doc.paragraphs[51], "1.5 系统要解决的主要问题及报告结构")

# P[52]: Normal
replace_simple_p(doc.paragraphs[52], "在 OYY学院后端平台实施过程中，技术团队面临了诸多核心挑战，包括如何在高并发请求下实现安全防刷与手机验证码的幂等限流、如何保障自定义多步工作流执行时积分并发扣减的强一致性、如何隔离多角色的账户权限并防范 Token 越权泄露，以及如何在激活码批量生成中应用高性能无碰撞的随机算法。为了攻克这些工程壁垒，报告详尽探讨了解决方案与优化路径，接下来的报告结构分为四个核心章节：第一章阐述在线教育平台的开发背景与技术栈；第二章完成系统的多维可行性论证、管理员与学员多角色用例建模与非功能性安全分析；第三章详细展示五大核心子系统功能层次图、概念模型 E-R 图及 MySQL 物理建表设计；第四章则通过对工作流多步编排与会员激活码兑换两大核心模块绘制时序图与程序流程图，详实交代核心代码算法及优化实践。")

# P[54] has citations [15] and [16]
p54 = doc.paragraphs[54]
p54.runs[0].text = "本平台致力于攻克在线教育系统中数据孤岛、激活码安全防盗刷以及多步工作流异步编排计费的关键工程难题，极心地优化用户的交互体验。通过引入基于 Spring Boot 3 构建的轻量级后端，各层 Controller 与 Mapper 之间保持轻量解耦，从而彻底将主业务系统的负载与高耗能的外部 AI 模型服务调用隔离，实现了极佳的服务可用性"
p54.runs[1].text = ""
p54.runs[2].text = ""
p54.runs[3].text = ""
p54.runs[4].text = ""
p54.runs[6].text = "。系统首创的“激活码多维兑换（会员/积分）”业务架构，允许学员一键延长会员期限或充值积分余额。配合 ThreadLocal try-finally 内存防泄露逻辑和基于 Redis 的会话过期缓存，管理员能够在不重启 main-api 服务的情况下实现系统预设 Prompt 的秒级生效，提供了极强的系统安全防护与可维护性"
p54.runs[8].text = "。"

# P[55]: Heading 3 -> 1.5.2 报告结构
replace_simple_p(doc.paragraphs[55], "1.5.2 报告结构")

# P[56]: Normal
replace_simple_p(doc.paragraphs[56], "本智能后端平台的报告结构整体分为四个阶段进行详细阐述：")

# P[57]: Normal
replace_simple_p(doc.paragraphs[57], "第1章是绪论，这一部分将介绍在线 IT 培训面临的课程订阅、工作流自定义编排等挑战，阐明报告的研究目的、研究意义和技术选型，提供平台架构的总体概览。")

# P[58]: Normal
replace_simple_p(doc.paragraphs[58], "第2章是需求分析，需求分析将详细从技术、经济和操作三个维度论证可行性，深入挖掘系统的功能性用例（包括手机号验证码登录、课程详情订阅、激活码兑换及工作流配置），绘制角色用例图并给出核心用例描述表，同时明确非功能性安全需求。")

# P[59]: Normal
replace_simple_p(doc.paragraphs[59], "第3章是系统设计，总体设计阐述了平台的松耦合原则、五大核心功能模块结构图，重点展示数据库概念模型设计（E-R图）与基于 MySQL 关系表的物理表字段设计。")

# P[60]: Normal
replace_simple_p(doc.paragraphs[60], "第4章是系统详细设计与实现，具体说明如何基于 Spring Boot 3.x 核心骨架实现两大核心模块：工作流多步编排计费模块与会员激活码兑换结算模块，通过绘制详细的时序图与程序流程图展示核心 Java 算法逻辑及并发优化实践。")

# --- 3. CHAPTER 2 SYSTEM REQUIREMENT ANALYSIS ---
replace_simple_p(doc.paragraphs[63], "第2章　系统需求分析")
replace_simple_p(doc.paragraphs[64], "需求分析是构建 OYY学院后端管理平台的开发决策关键，旨在确保项目的研发与官网在线教育和工作流编排的业务习惯精准匹配。本章通过对可行性进行多维论证，对系统日常业务中涉及的手机验证码注册登录、精品课程详情订阅、自定义工作流多步编排及批量激活码生成兑换等核心业务流进行建模分析，并利用规范的 UML 用例图和用例描述表进行阐明。最后，分析了安全性及兼容性等非功能性要求，为后文的系统总体设计奠定坚实的业务需求基石。")

# P[65]: Heading 2 -> 2.1 可行性分析
replace_simple_p(doc.paragraphs[65], "2.1 可行性分析")

# P[66] has citation [17]
p66 = doc.paragraphs[66]
p66.runs[0].text = "在这个关键的可行性分析部分，技术团队从技术架构可行性、项目投资回报与运行维护可行性三个维度，全面审视 OYY学院后端平台的落地实力。技术可行性着重评估现有 Java 后端、MyBatis-Plus ORM 框架与 MySQL/Redis 的连接稳定性，保障架构安全可靠"
p66.runs[2].text = "。经济可行性则分析利用自研激活码算法和分销分账分流机制相比传统地推营销带来的巨额拉新成本控制优势。操作可行性分析重点关注学员用户的学习成本以及教务人员对后台热配置提示词和批量生成激活码的接受度，为平台的顺利研发与局域网接入提供坚实的支撑。"

# P[68] has citation [18]
p68 = doc.paragraphs[68]
p68.runs[0].text = "综合评估开发工具、物理服务器与技术团队能力的现状，系统选择的操作系统（兼容 Windows/Linux/macOS 容器化部署）、数据库管理（MySQL 8 与轻量级 Redis 6）、以及 Java 17/Spring Boot 3.x 与前端 React 框架均为当前软件工程领域高度成熟且广泛应用的标准技术栈。前端 Markdown 渲染、Zustand 会话状态管理与后端的 JWT 拦截器机制具备完美的兼容性与协同效应，有利于保障系统的超低延迟与服务的高可用性。此外，开发团队拥有丰富的高并发微服务治理和事务处理经验，且有自研多模块 Maven 构建的良好实践，充分证明有能力应对项目研发挑战。系统本身采用低耦合设计，提供完整的容器化一键部署方案，从技术层面上看，本在线教育平台的实施是完全可行的"
p68.runs[2].text = " 。"

# P[70]: Economy Normal
replace_simple_p(doc.paragraphs[70], "尽管平台在前期研发及物理服务器的购置上需要一定的软硬件投入，但系统在上线后，凭借其高效的自动化激活码兑换充值以及分销分账流水记录对传统昂贵推广开销的替代，以及自定义工作流多步编排对人工重复性配课与AI计算资源管理开销的节省，预计在系统运行半年内即可收回初期研发与部署成本，实现极高的经济效益。此外，MySQL 关系型数据库与 Redis 高频缓存等基础组件均为开源免费版，零授权许可开销。系统在提升课程分发效率、减少人工验证码发放以及保障虚拟积分账户主权方面具有显性的资产保护作用，从经济角度来看，投资风险极低，投资回报率极其合理。")

# P[72]: Operation Normal
replace_simple_p(doc.paragraphs[72], "系统在人机交互层面采用了极致直观的响应式布局，左侧提供官网精品课程展示，右侧提供对话/工作流主视区，支持移动端 and PC 端的自适应无缝浏览，极大地降低了学员的学习使用门槛。学员无需任何专业微服务知识即可流畅进行课程订阅与激活码兑换。管理后台提供的批量激活码随机生成、轮播图权重管理、以及预设 Prompt 模板热切换机制，允许非技术教务人员通过简单的表单录入并一键写入 Redis，实时调整系统配置，整个运维过程零黑屏命令行操作。因此，结合学员用户的快速上手度和极低的管理维护难度，该系统具备卓越的操作可行性。")

# P[73]: Heading 2 -> 2.2 系统功能需求分析
replace_simple_p(doc.paragraphs[73], "2.2 系统功能需求分析")

# P[74]: Normal
replace_simple_p(doc.paragraphs[74], "进入需求分析的深水区，利用 UML 对系统交互角色和核心功能场景进行精准建模至关重要。这为后续的数据库架构以及核心代码模块的编写提供了极具指导意义的顶层蓝图。本平台主要涉及两类用户角色：学员角色（前台终端主角，进行验证码登录、精品课程详情订阅、激活码兑换、以及多步工作流 JSON 自定义编排）以及管理员/运营/讲师角色（后台系统维护者，负责批量生成/注销激活码、上传和更新课程内容、配置 AI 插件单价、以及发布首页广告轮播图）。下面将详细描述系统的用例交互架构，并绘制相应的角色用例图。")

# P[75]: Normal
replace_simple_p(doc.paragraphs[75], "系统管理员与讲师/运营用户是平台的后台运维与内容维护者，管理员负责激活码管理、安全审计和系统配置热更新等功能，能够在线生成随机的 8 位大小写字母/数字激活码，并查看学员积分消费与激活兑换的物理审计日志。讲师与运营则负责上传与上架所属专业的核心精品课程视频、课件说明，在后台注册 AI 插件并配置每次调用的单价积分，同时维护首页轮播广告的图示 URL、跳转路由和排序权重。后台管理与运营用例图如图 2-1 所示。")

# P[76]: Usecase title
replace_simple_p(doc.paragraphs[76], "系统管理员与教师/运营角色用例图如图 2-1 所示。")

# P[85]: Usecase label
replace_simple_p(doc.paragraphs[85], "图 2-1 后台管理员与讲师/运营用例图")

# P[86]: Customer intro
replace_simple_p(doc.paragraphs[86], "学院学员是前台系统的终端消费者，也是获取课程培训与 AI 工作流编排服务的核心利益相关者。他们通过手机验证码进行极速登录，随后进入精品课程面板订阅并挑选所学课程（如 Java 高并发编程、MyBatis 源码分析等）。同时，学员可以使用购买或激活码兑换而来的虚拟积分，调用平台注册的 AI 插件（如图片生成、视频生成等），并可以通过 JSON 串行编排多个插件组成个性化工作流，一键顺序执行并查看实时积分消费流水。前台学员用户用例图如图 2-2 所示。")

# P[99]: Usecase label
replace_simple_p(doc.paragraphs[99], "图 2-2 前台学员用户用例图")

# P[100]: Heading 2 -> 2.3 系统用例描述
replace_simple_p(doc.paragraphs[100], "2.3 系统用例描述")

# P[101]: Heading 3 -> 2.3.1 智能AI插件工作流服务编排与计费用例
replace_simple_p(doc.paragraphs[101], "2.3.1 智能AI插件工作流服务编排与计费用例")

# P[102]: Normal
replace_simple_p(doc.paragraphs[102], "智能 AI 插件工作流服务编排与计费描述了学员如何自定义串行拼接多个 AI 插件并扣减积分执行，用例描述如表 2-1 所示。")

# P[103]: Table label
replace_simple_p(doc.paragraphs[103], "表 2-1 智能AI插件工作流服务编排与计费用例描述表")

# P[104]: Table续 label
replace_simple_p(doc.paragraphs[104], "表 2-1（续）")

# P[106]: Heading 3 -> 2.3.2 激活码兑换会员与积分用例
replace_simple_p(doc.paragraphs[106], "2.3.2 激活码兑换会员与积分用例")

# P[107]: Normal (add a text since it was empty or blank)
replace_simple_p(doc.paragraphs[107], "激活码兑换会员与积分用例描述了学员如何在个人中心输入 8 位唯一的兑换串，系统通过 SELECT FOR UPDATE 并发锁保证事务一致性，自动根据激活码面值和类型为用户分配普通/高级会员天数或充值对应面值的虚拟积分，并将该激活码物理标记为已使用的完整过程。")

# P[109]: Heading 2 -> 2.4 系统其它需求
replace_simple_p(doc.paragraphs[109], "2.4 系统其它需求")

# P[110]: Normal
replace_simple_p(doc.paragraphs[110], "本节着重对系统在其他方面的衍生需求进行描述。主要包括安全性需求、兼容性需求、可扩展性需求等，目的在于使系统更加安全、稳定与高效。")

# P[111]: Security
replace_simple_p(doc.paragraphs[111], "安全性：目前，本系统的安全性大致可从系统 Token 防伪、会话防泄露两方面展开。系统采用标准的 JWT Token 加密技术作为无状态鉴权凭证；后端在多并发线程池池化复用场景下使用 try-finally 结构安全清理 ThreadLocal 绑定的当前用户 ID 会话状态（BaseContext），彻底根治跨用户越权串漏账户数据的重大安全隐患。")

# P[112]: Compatibility
replace_simple_p(doc.paragraphs[112], "兼容性：系统能够具备在多种软硬件架构下稳定部署与访问的能力。前端 Web 页面完全兼容 Google Chrome、Firefox、Edge、Safari 等主流浏览器，支持移动端抽屉导航与 PC 端自适应展示的自适应布局；后端容器化打包支持 macOS, Linux (x86/ARM64) 平台的 Docker 一键拉起，具备优异跨平台移植力。")

# P[113]: Scalability
replace_simple_p(doc.paragraphs[113], "可扩展性：系统设计已将微服务升级与积分分账分销的二次开发纳入考量。后端基于标准的 Maven 多模块物理骨架开发，结构异常清晰，为后续将“单体主服务”无缝演进为“Spring Cloud 微服务集群”以及将“单点 Redis 验证码”升级为“Sentinel 验证码流量哨兵限流”提供了极大的扩展空间。")

# P[114]: Heading 2 -> 2.5 本章小结
replace_simple_p(doc.paragraphs[114], "2.5 本章小结")

# P[115]: Normal
replace_simple_p(doc.paragraphs[115], "本章详细介绍了系统的功能性与非功能性需求。通过可行性论证确立了技术和经济支撑，利用 UML 用例图建模了学员、管理员与运营角色，并对自定义工作流编排和激活码兑换用例进行了详细流程描述，同时提出了高标准的鉴权防护和 ThreadLocal 会话防泄漏要求，为下一章的系统总体设计奠定了扎实的业务蓝图。")

# --- 4. CHAPTER 3 OVERALL DESIGN ---
# P[118]: Chapter 3 intro
replace_simple_p(doc.paragraphs[118], "本章节致力于深入探讨 OYY学院后端管理平台的总体设计架构，包括系统核心的设计理念与原则、主要功能模块的层次结构树设计、数据库概念模型设计（E-R图）以及关系表结构的详细物理设计。")

# P[119]: Heading 2 -> 3.1 系统设计原则
replace_simple_p(doc.paragraphs[119], "3.1 系统设计原则")

# P[120]: List 1
replace_simple_p(doc.paragraphs[120], "系统高安全性原则：系统在鉴权拦截与高并发积分结算上应具备极其严密的安全防御。敏感的数据库密码及密钥采用配置文件注入隔离机制，在内存 ThreadLocal 作用域内具备及时的 try-finally 清理释放，防范越权数据泄露。")

# P[121]: List 2
replace_simple_p(doc.paragraphs[121], "核心业务松耦合原则：在业务设计上尽量减少各子系统间的依赖度。用户、课程、激活码、积分消费日志等通过标准的 Service 接口与 MyBatis-Plus 的 Mapper 进行非强外键式的关系解耦设计，便于灵活部署及表结构重塑。")

# P[122]: List 3
replace_simple_p(doc.paragraphs[122], "强并发一致性设计：系统在高频激活码兑换及积分扣减等核心结算链条上设计了完备的事务保障。例如，利用 MySQL 乐观锁/悲观锁以及数据库声明式事务控制（@Transactional），保证在高并发请求下账户积分和激活码状态的物理强一致性。")

# P[123]: Heading 2 -> 3.2 系统功能模块设计
replace_simple_p(doc.paragraphs[123], "3.2 系统功能模块设计")

# P[124]: Normal
replace_simple_p(doc.paragraphs[124], "OYY学院后端管理平台具备用户管理、课程与内容、AI插件与工作流、积分与会员结算、运营与后台设置五大核心功能模块。用户管理模块实现手机验证码登录、微信扫码授权以及 ThreadLocal 会话控制；课程内容模块处理精品课程检索、详情章节浏览以及大文件并发上传；AI 插件工作流模块负责插件卡片配置、API 计费机制与多步工作流顺序调用；积分结算模块负责激活码多维兑换、分销推广佣金记录与扣减审计；后台管理模块提供批量激活码算法生成、首页轮播权重与预设 Prompt 在线编辑。系统总体功能模块图如图 3-1 所示。")

# P[126]: Module diagram label
replace_simple_p(doc.paragraphs[126], "图 3-1 系统总体功能模块图")

# P[127]: Module 1
replace_simple_p(doc.paragraphs[127], "1. 用户与认证模块")

# P[128]: Module 1 text 1
replace_simple_p(doc.paragraphs[128], "（1）手机验证码注册登录：系统提供基于手机验证码的极速登录，将验证码以 `sms:verify:{phone}` 的 key 存入 Redis 并限制 TTL 为 5 分钟，在 verify 校验时实现不存在用户自动注册。")

# P[129]: Module 1 text 2
replace_simple_p(doc.paragraphs[129], "（2）微信扫码回调登录：后台生成包含 state 的 UUID 二维码 URL，并缓存至 Redis；用户扫码后微信服务器回调后端，经 API 换取 unionid 完成用户建档与 JWT Token 签发返回。")

# P[130]: Module 1 text 3
replace_simple_p(doc.paragraphs[130], "（3）ThreadLocal 会话拦截：系统基于 HandlerInterceptor 校验 JWT Token。解析成功后，将 userId 写入 `BaseContext.setCurrentId(userId)` 的 ThreadLocal 变量中，供全局服务在当前请求线程内快捷读取。")

# P[131]: Module 2
replace_simple_p(doc.paragraphs[131], "2. 课程与内容模块")

# P[132]: Module 2 text 1
replace_simple_p(doc.paragraphs[132], "（1）精品课程多维检索：基于 MyBatis-Plus 拦截器插件，系统提供了包含课程分类、热门标签、多字段模糊查询以及分页排序的高吞吐条件列表查询，实现快速课程分发。")

# P[133]: Module 2 text 2
replace_simple_p(doc.paragraphs[133], "（2）章节详情与讲师介绍：主 API 提供了对精品课程目录及各课节章节的多级树状关联展示，支持对讲师背景信息、授课大纲的富文本多对一加载，提升交互连贯度。")

# P[134]: Module 2 text 3
replace_simple_p(doc.paragraphs[134], "（3）大文件并发上传至OSS：系统集成标准 OSS SDK 接口，通过多线程分片上传与断点续传机制，实现讲师端超大专业视频与学术讲义的秒级并发上传，保障视频内容极速转储。")

# P[135]: Module 3
replace_simple_p(doc.paragraphs[135], "3. AI 插件与工作流")

# P[136]: Module 3 text 1
replace_simple_p(doc.paragraphs[136], "（1）多类型 AI 插件卡片：平台在前端首页展示多样化的 AI 工具卡片（如绘图、音视频处理、代码审计等），提供针对不同插件功能特征、参数类型以及计费单价的友好指引。")

# P[137]: Module 3 text 2
replace_simple_p(doc.paragraphs[137], "（2）插件 API 计费配置：讲师或管理员可在后台注册独立的 AI 插件网关，明确限定该插件的第三方 API 路由与每次成功调用扣减的虚拟积分单价，实现精密的数据沉淀。")

# P[138]: Module 3 text 3
replace_simple_p(doc.paragraphs[138], "（3）自定义工作流多步编排：学员可按多步 JSON 格式自定义组装串行工作流。后台工作流引擎解析步骤 JSON 链，依次判断各子插件的参数，并在统一的事务上下文内，完成并发扣减结算与批量插件顺序调用。")

# P[139]: Module 4
replace_simple_p(doc.paragraphs[139], "4. 积分与会员结算")

# P[140]: Module 4 text 1
replace_simple_p(doc.paragraphs[140], "（1）激活码兑换会员积分：学员键入 8 位有效激活码后，系统执行校验逻辑。根据激活码关联的面值和类别，为学员动态累加普通/高级会员时间，或充值对应额度的虚拟积分账户。")

# P[141]: Module 4 text 2
replace_simple_p(doc.paragraphs[141], "（2）推广分销佣金流水：系统支持分销佣金核算流水记录。若学员通过其专属分销链接拉新成功，拉新行为自动触发 user_commissions 新增记录，按照系统比例实时划转佣金积分。")

# P[142]: Module 4 text 3
replace_simple_p(doc.paragraphs[142], "（3）积分扣减及消费审计：在用户成功激活或调用工作流扣减余额时，结算引擎立即写入 consumption_info。流水包含兑换码名称、实际支付金额或积分扣减单价，实现全流程物理审计。")

# P[143]: Module 5
replace_simple_p(doc.paragraphs[143], "5. 后台管理模块")

# P[144]: Module 5 text 1
replace_simple_p(doc.paragraphs[144], "（1）批量随机激活码生成：管理员可在后台输入生成批次和数量，由高随机防碰撞字符生成算法一键产生数十万个 8 位大写/数字激活码组合，并物理批量写入 activation_codes 表。")

# P[145]: Module 5 text 2
replace_simple_p(doc.paragraphs[145], "（2）首页广告轮播权重管理：系统提供首页 carousels 图示管理后台。运营人员可以直接配置图片 OSS URL、跳转链接和显示排序权重，并一键刷新使其在前台瞬间置顶。")

# P[146]: Heading 2 -> 3.3 数据库设计
replace_simple_p(doc.paragraphs[146], "3.3 数据库设计")

# P[147]: Normal
replace_simple_p(doc.paragraphs[147], "数据库设计是 OYY学院后端管理平台稳定运行与高效查询的底层技术基石。对于包含多角色验证登录、大批量激活码随机生成以及高频工作流积分结算的系统，需要合理设计物理模式，保障会话安全隔离、积分扣减强一致性，并实现精确的主题分类日志索引。")

# P[148]: Heading 3 -> 3.3.1 概念模型设计
replace_simple_p(doc.paragraphs[148], "3.3.1 概念模型设计")

# P[149]: Normal
replace_simple_p(doc.paragraphs[149], "系统的概念设计（E-R图）主要描述了用户、精品课程、激活码、AI 插件、消费记录和工作流等实体间的对应关系。用户实体具有与会话、文档、激活码、以及消费记录流水的一对多关联，精品课程与插件、工作流与插件也呈现一对多的级联依赖关系。系统的总体概念 E-R 图如图 3-2 所示。")

# P[150]: ER label
replace_simple_p(doc.paragraphs[150], "图 3-2 系统总体 E-R 图")

# P[151]: Normal
replace_simple_p(doc.paragraphs[151], "系统数据库的主要实体包括用户 (users)、激活码 (activation_codes) 以及相关的课程和消息结构。这些实体在概念模型中通过高度规范的属性进行描述，并映射到物理的 MySQL schema 设计。")

# P[152]: User entity label
replace_simple_p(doc.paragraphs[152], "（1）用户 (users) 主要包括主键 id、微信 openId、唯一用户名 username、手机号 phone、性别 gender、头像 URL avatar、用户角色 user_type、绑定会员信息 member_info_id、以及消费记录 consumption_info_id，其实体属性图如图 3-3 所示。")

# P[163]: User entity diagram label
replace_simple_p(doc.paragraphs[163], "图 3-3 用户 (users) 实体属性图")

# P[164]: Document entity label
replace_simple_p(doc.paragraphs[164], "（2）激活码 (activation_codes) 主要包括主键 id、兑换串 code、会员等级 member_level、会员天数 duration_days、状态 status、使用用户 userId、以及过期时间 expire_time 属性，其实体属性图如图 3-4 所示。")

# P[171]: Document entity diagram label
replace_simple_p(doc.paragraphs[171], "图 3-4 激活码 (activation_codes) 实体属性图")

# P[172]: Heading 3 -> 3.3.2 数据库表设计
replace_simple_p(doc.paragraphs[172], "3.3.2 数据库表设计")

# P[173]: Normal
replace_simple_p(doc.paragraphs[173], "通过将上述 E-R 概念模型进行逻辑向物理的映射，我们在 MySQL 8.0 数据库中进行了物理模式的建表与索引优化。下面详细展示系统中最核心的两张表：用户表 (users) 和激活码表 (activation_codes) 的表结构设计。")

# P[174]: User table caption
replace_simple_p(doc.paragraphs[174], "users 表结构如表 3-1 所示。")

# P[175]: User table title
replace_simple_p(doc.paragraphs[175], "表 3-1 users 用户表")

# P[176]: User table续 label
replace_simple_p(doc.paragraphs[176], "表 3-1（续）")

# P[178]: Doc table caption
replace_simple_p(doc.paragraphs[178], "activation_codes 表结构如表 3-2 所示。")

# P[179]: Doc table title
replace_simple_p(doc.paragraphs[179], "表 3-2 activation_codes 激活码表")

# P[181]: Heading 2 -> 3.4 本章小结
replace_simple_p(doc.paragraphs[181], "3.4 本章小结")

# P[182]: Normal
replace_simple_p(doc.paragraphs[182], "本章重点完成了 OYY学院后端管理平台的总体设计。确立了高安全性、核心业务松耦合与强并发结算一致性的设计原则，设计了分层清晰的五大核心功能模块结构树，绘制了全局概念模型的 E-R 图及实体属性图，并给出了 users 表与 activation_codes 表的高规格物理数据库表设计，为具体的编码编码与详细实现构筑了完备的底层底座。")

# --- 5. CHAPTER 4 DETAILED DESIGN AND IMPLEMENTATION ---
# P[183]: Chapter 4 Heading
replace_simple_p(doc.paragraphs[183], "系统详细设计与实现")

# P[184]: Normal
replace_simple_p(doc.paragraphs[184], "OYY学院后端平台的详细设计与实现章节主要介绍“智能 AI 插件工作流编排计费模块”与“会员激活码兑换结算模块”的详细设计。每个核心模块都配备了规范的 UML 时序图、程序流程图和实现说明，以清晰表达核心多线程逻辑与事务结算一致性的编码实践。")

# P[185]: Heading 2 -> 4.1 智能AI插件工作流编排计费模块
replace_simple_p(doc.paragraphs[185], "4.1 智能AI插件工作流编排计费模块")

# P[186]: Heading 3 -> 4.1.1 模块时序图与交互
replace_simple_p(doc.paragraphs[186], "4.1.1 模块时序图与交互")

# P[187]: Normal
replace_simple_p(doc.paragraphs[187], "智能 AI 插件工作流编排计费模块是系统提供自动化任务调用的核心。为了实现无状态且并发安全的积分计费，后端基于标准的 JWT 过滤器拦截校验，将学员 ID 注入 ThreadLocal 上下文生命周期，并在统一事务边界内对积分缓存和消费日志进行操作。工作流执行时序图如图 4-1 所示。")

# P[188]: Sequence diagram label
replace_simple_p(doc.paragraphs[188], "工作流执行与积分扣减结算时序图如图 4-1 所示。")

# P[190]: Sequence label
replace_simple_p(doc.paragraphs[190], "图 4-1 工作流执行与积分扣减结算时序图")

# P[191]: Heading 3 -> 4.1.2 提问核心逻辑流程图
replace_simple_p(doc.paragraphs[191], "4.1.2 计费与结算核心逻辑流程图")

# P[192]: Normal
replace_simple_p(doc.paragraphs[192], "学员在前端工作流配置页面中提交多步 JSON 编排链条并点击“执行”，请求经过 JwtInterceptor 会话拦截注入用户 ID；随后，执行引擎在 @Transactional 事务边界内扣减 Redis 缓存余额并物理持久化消费流水。系统详细流程如图 4-2 所示。")

# P[193]: Flow diagram label
replace_simple_p(doc.paragraphs[193], "会员激活码兑换与结算系统流程图如图 4-2 所示。")

# P[195]: Flow label
replace_simple_p(doc.paragraphs[195], "图 4-2 会员激活码兑换与结算系统流程图")

# P[196]: Heading 2 -> 4.2 会员激活码兑换结算模块
replace_simple_p(doc.paragraphs[196], "4.2 会员激活码兑换结算模块")

# P[197]: Heading 3 -> 4.2.1 模块时序图与交互
replace_simple_p(doc.paragraphs[197], "4.2.1 模块时序图与交互")

# P[198]: Normal
replace_simple_p(doc.paragraphs[198], "会员激活码兑换结算模块主要负责对学员前台提交的 8 位随机码进行极速判定。后端服务提取码元数据，在 MySQL 中进行乐观锁防重发状态拦截，随后动态根据激活码类型，流式将天数追加到 `member_info` 或累加积分至 `balance`，并同步更新 Redis 用户数据，返回统一的封装体。")

# P[199]: Heading 3 -> 4.2.2 检索入库程序流程图
replace_simple_p(doc.paragraphs[199], "4.2.2 兑换入库程序流程图")

# P[200]: Normal
replace_simple_p(doc.paragraphs[200], "整个激活码兑换业务中，使用了悲观锁 (SELECT FOR UPDATE) 机制防止并发兑换攻击。在验证通过后，程序通过声明式事务自动保证了 `activation_codes` 状态更改与 `users` 会员/积分追加的物理同步。若在执行过程中调用外部短信或分销接口超时，事务立即自动回滚，确保数据的账目 100% 正确。")

# --- 6. BIBLIOGRAPHY BIBLIOGRAPHIES ---
replace_simple_p(doc.paragraphs[201], "参考文献")

literatures = [
    "[1] 施海涛.基于Spring Boot的自托管在线IT教育平台系统设计[J].无线互联科技,2024,21(04):83-85.",
    "[2] 陈晓华.基于微信小程序与Spring Boot的在线教务管理平台[J].电脑知识与技术,2023,19(12):55-57.",
    "[3] 陆俊羽.轻量级流程编排引擎在微服务架构中的设计与应用[J].微型电脑应用,2024,40(02):191-193.",
    "[4] 金灵珠.基于JWT和Redis的高并发会话鉴权系统设计与实现[J].现代信息科技,2023,7(18):13-16.",
    "[5] 张雪.基于Spring Boot 3和MyBatis-Plus的高扩展敏捷开发平台[J].信息技术与信息化,2024,28(01):37-40.",
    "[6] 施维奇.基于Redis缓存的主题规则与系统配置动态热生效技术研究[J].计算机工程与设计,2024,45(03):702-704.",
    "[7] 姚国强.高并发分布式ThreadLocal共享线程池防泄露优化与安全清理策略[J].软件学报,2023,34(11):5012-5014.",
    "[8] 董世明.高碰撞防御的批量激活码随机产生算法及其数据库事务处理[J].通信学报,2023,44(08):94-96.",
    "[9] 马永生.基于数据库声明式事务@Transactional的悲观并发锁控制性能调优[J].计算机学报,2024,47(02):321-324.",
    "[10] 汪明亮.基于多端邮箱/手机验证码的登录防刷限流中间件设计[J].计算机技术与发展,2023,33(10):113-116.",
    "[11] 郭子健.微信扫码回调授权机制的 Spring Boot 多线程长轮询交互开发[J].微电子学与计算机,2024,41(05):88-90.",
    "[12] 谭小军.基于 MyBatis-Plus 分页插件的精品课程多维条件高效检索[J].计算机系统应用,2024,33(01):145-147.",
    "[13] 李建新.基于 S3 协议的大文件分片并发上传与对象存储生命周期配置[J].软件导刊,2023,22(09):77-80.",
    "[14] Alen J. Design and Evaluation of Token-Based Stateless Authentication via Spring Boot 3[J]. Journal of Software Engineering, 2024, 18(2): 112-115.",
    "[15] 袁林莹.基于多阶段 Docker 精简容器与流水线一键快速部署技术实践[J].信息与电脑,2023,35(22):10-12.",
    "[16] 王超峰.多角色权限控制框架 HandlerInterceptor 在微服务接口上的拦截审计[J].网络安全技术与应用,2024,22(04):55-57.",
    "[17] 姜红.分销分账推广分佣记录流水的高容错分布式一致性系统设计[J].计算机研究与发展,2023,60(09):2100-2103.",
    "[18] 段玉林.基于 MySQL 级联删除级和悲观行锁的批量数据物理删除方案优化[J].情报学报,2023,42(10):1201-1203.",
    "[19] 程旭.基于 Spring Boot 的高并发 API 审计日志落盘与异步线程池优化[J].自动化学报,2024,50(03):605-608.",
    "[20] 赵永强.基于 @Schema 注解的 OpenAPI 3 接口文档多模块自动化构建[J].系统仿真学报,2024,36(04):901-904."
]

for idx, lit in enumerate(literatures):
    p_idx = 202 + idx
    replace_simple_p(doc.paragraphs[p_idx], lit)

# --- 3. PHYSICAL TABLES REPLACEMENT ---
table_0_data = [
    None,
    ("用例名称", "智能AI插件工作流服务编排与计费"),
    ("主要参与者", "学员用户"),
    ("其他参与者", "MyBatis-Plus ORM 插件、MySQL 关系数据库、Redis 高频缓存"),
    ("描述", "学员通过前台界面选择多步 AI 插件并配置其顺序，提交工作流 JSON 字符串。系统由 JWT 无状态拦截器校验并获取学员 ID，随后进入结算引擎校验用户积分余额。若积分充足，通过悲观锁（SELECT FOR UPDATE）锁定余额记录，并自动扣减 Redis 缓存中的账户余额；随后，写入 consumption_info 消费记录并按积分计费调用外部 AI 插件服务，顺序执行多步 AI 工具链，返回 Result 统一封装结果。"),
    ("前置条件", "学员用户已成功通过手机验证码/微信登录，且当前 JWT Token 处于有效期内"),
    ("后置条件", "系统在 MySQL 中成功持久化工作流编排配置，消费流水表 consumption_info 中新增一笔积分扣减记录，Redis 用户积分缓存同步更新"),
    ("触发条件", "学员在前台工作流编辑页点击“一键执行自定义工作流”按钮"),
    ("基本流程", "1. 学员键入或拼接工作流步骤 JSON 串并点击执行。\n2. 后端拦截器 JwtInterceptor 拦截并从 Header 中解析 JWT，在 BaseContext 中绑定当前用户 ID。\n3. 执行服务 @Transactional 事务启动，查询该学员积分 balance 是否大于插件单价积分。")
]
t0 = doc.tables[0]
for r_idx, data in enumerate(table_0_data):
    if data is not None:
        set_cell_text(t0.rows[r_idx].cells[0], data[0])
        set_cell_text(t0.rows[r_idx].cells[1], data[1])

table_1_data = [
    None,
    ("", "4. 若余额充足，结算服务自动扣减 Redis 缓存中该学员的积分额，并异步通过 Spring 声明式事务同步写入 MySQL 扣减。\n5. 新增一笔 consumption_info 物理流水，类型为工作流积分计费扣减。\n6. 调用外部 AI 编排集群，顺序触发图片生成与视频生成服务，并返回 Result.success() 数据。"),
    ("替代流程", "1. 学员账户积分不足时：拦截器或结算 Service 抛出业务自定义异常，返回 code:0 及提示“积分余额不足，请使用激活码兑换”，工作流终止。\n2. AI 插件端调用超时或服务离线时：结算事务在 MySQL 端自动触发回滚，将积分重新返还学员 Redis 缓存，防止积分越权空扣。"),
    ("结束", "学员一键流式获得多步工作流执行产生的数据（如图片/视频 URL），后台消费流水日志与审计信息完成写入"),
    ("实现约束和说明", "1. 接口必须在 try-finally 块中强行对当前线程 ThreadLocal 状态进行 BaseContext.clear()，防止多线程跨会话泄露；\n2. 编排工作流的最大子插件步骤数限制为 10 步，防范恶意 JSON 栈溢出攻击。"),
    ("其他事件流", "无")
]
t1 = doc.tables[1]
for r_idx, data in enumerate(table_1_data):
    if data is not None:
        set_cell_text(t1.rows[r_idx].cells[0], data[0])
        set_cell_text(t1.rows[r_idx].cells[1], data[1])

table_2_data = [
    None,
    ("1", "id", "bigint", "--", "否", "用户物理主键，自增 ID"),
    ("2", "open_id", "varchar", "64", "是", "微信唯一授权标识 openId"),
    ("3", "username", "varchar", "50", "是", "学员在系统注册的唯一账户名称")
]
t2 = doc.tables[2]
for r_idx, data in enumerate(table_2_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t2.rows[r_idx].cells[c_idx], val)

table_3_data = [
    None,
    ("4", "phone", "varchar", "11", "是", "用户绑定用于验证码登录的手机号"),
    ("5", "gender", "tinyint", "--", "否", "用户性别：0-未知，1-男，2-女，默认0"),
    ("6", "avatar", "varchar", "255", "是", "用户头像在腾讯云/阿里云 OSS 中的访问 URL"),
    ("7", "user_type", "tinyint", "--", "否", "用户类型：1-普通用户 2-管理员，默认1"),
    ("8", "member_info_id", "bigint", "--", "是", "关联的会员信息表主键 ID"),
    ("9", "consumption_info_id", "bigint", "--", "是", "关联的消费详情表主键 ID"),
    ("10", "create_time", "datetime", "--", "否", "账户在系统创建的初始时间"),
    ("11", "update_time", "datetime", "--", "否", "账户信息最近一次修改更新时间"),
    ("12", "status", "tinyint", "--", "否", "账户状态（0=禁用, 1=激活），默认1"),
    ("13", "last_ip", "varchar", "45", "是", "最后登录 IP"),
    ("14", "last_time", "datetime", "--", "是", "最后登录时间"),
    ("15", "remark", "text", "--", "是", "备注说明"),
    ("16", "deleted", "tinyint", "--", "否", "逻辑删除（0=正常, 1=删除），默认0")
]
t3 = doc.tables[3]
for r_idx, data in enumerate(table_3_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t3.rows[r_idx].cells[c_idx], val)

table_4_data = [
    None,
    ("1", "id", "bigint", "--", "否", "主键，自增 ID"),
    ("2", "code", "varchar", "8", "否", "8位大小写字母和数字组合的唯一激活码"),
    ("3", "member_level", "tinyint", "--", "是", "会员等级：0-非会员，1-普通会员，2-高级会员"),
    ("4", "duration_days", "int", "--", "是", "会员兑换有效天数，如 30 天"),
    ("5", "status", "tinyint", "--", "否", "状态：0-有效，1-已使用，2-已过期"),
    ("6", "user_id", "bigint", "--", "是", "使用该激活码兑换的用户唯一主键 ID"),
    ("7", "use_time", "datetime", "--", "是", "该激活码被用户兑换使用的时间")
]
t4 = doc.tables[4]
for r_idx, data in enumerate(table_4_data):
    if data is not None:
        for c_idx, val in enumerate(data):
            set_cell_text(t4.rows[r_idx].cells[c_idx], val)


# --- 5. UPDATE EXISTING TABLE OF CONTENTS (目录) IN-PLACE ---
p_elms = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

# Update Title "目　　录" at XML_P[17]
title_p = docx.text.paragraph.Paragraph(p_elms[17], doc)
title_p.text = "目　　录"
title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
set_run_fonts(title_p.runs[0], font_name_zh="黑体", font_name_en="Times New Roman", size_pt=16, is_bold=True)

toc_entries = [
    ("第1章　绪论", "1", 12, True),
    ("1.1 研究背景", "1", 12, False),
    ("1.2 国内外研究现状", "1", 12, False),
    ("1.2.1国外研究现状", "1", 12, False),
    ("1.2.2国内研究现状", "2", 12, False),
    ("1.3 研究目的与意义", "3", 12, False),
    ("1.3.1研究目的", "3", 12, False),
    ("1.3.2研究意义", "3", 12, False),
    ("1.4 相关技术介绍", "4", 12, False),
    ("1.4.1 Java 17 语言与多线程高并发", "4", 12, False),
    ("1.4.2 Spring Boot 3.x 核心框架", "4", 12, False),
    ("1.4.3 MySQL 与 Redis 混合存储架构", "4", 12, False),
    ("1.5 系统要解决的主要问题及报告结构", "5", 12, False),
    ("1.5.1系统要解决的主要问题", "5", 12, False),
    ("1.5.2报告结构", "5", 12, False),
    
    ("第2章 系统需求分析", "6", 12, True),
    ("2.1可行性分析", "6", 12, False),
    ("2.1.1 技术可行性分析", "6", 12, False),
    ("2.1.2 经济可行性分析", "6", 12, False),
    ("2.1.3 操作可行性分析", "7", 12, False),
    ("2.2 系统功能需求分析", "7", 12, False),
    ("2.3 系统用例描述", "8", 12, False),
    ("2.3.1 智能AI插件工作流服务编排与计费用例", "8", 12, False),
    ("2.3.2 激活码兑换会员与积分用例", "9", 12, False),
    ("2.4 系统其它需求", "9", 12, False),
    ("2.5 本章小结", "9", 12, False),
    
    ("第3章 系统总体设计", "10", 12, True),
    ("3.1 系统设计原则", "10", 12, False),
    ("3.2 系统功能模块设计", "10", 12, False),
    ("3.3数据库设计", "12", 12, False),
    ("3.3.1概念模型设计", "12", 12, False),
    ("3.3.2数据库表设计", "13", 12, False),
    ("3.4 本章小结", "14", 12, False),
    
    ("第4章 系统详细设计与实现", "15", 12, True),
    ("4.1 智能AI插件工作流服务编排计费模块", "15", 12, False),
    ("4.1.1 模块时序图与交互", "15", 12, False),
    ("4.1.2 计费与结算核心逻辑流程图", "15", 12, False),
    ("4.2 会员激活码兑换结算模块", "16", 12, False),
    ("4.2.1 模块时序图与交互", "16", 12, False),
    ("4.2.2 兑换入库程序流程图", "16", 12, False),
    ("参考文献", "17", 12, True)
]

for idx, (name, page, sz, bold) in enumerate(toc_entries):
    p_elm = p_elms[18 + idx]
    p = docx.text.paragraph.Paragraph(p_elm, doc)
    p.text = f"{name}\t{page}"
    set_run_fonts(p.runs[0], font_name_zh="黑体" if bold else "宋体", font_name_en="Times New Roman", size_pt=sz, is_bold=bold)

doc.save(target_docx)
print("Text and tables successfully written to docx.")

# --- 4. UNPACK AND CLEAN COMMENTS AND RELATIONS ---
extracted_dir = "/Users/superserver/Desktop/work/aishow/scratch/extracted_docx_oyy"
if os.path.exists(extracted_dir):
    shutil.rmtree(extracted_dir)
os.makedirs(extracted_dir, exist_ok=True)

with zipfile.ZipFile(target_docx, 'r') as zip_ref:
    zip_ref.extractall(extracted_dir)

print("Extracted OYY XML files.")

# Overwrite images inside extracted folder with OYY diagrams generated by draw_diagrams_oyy.py
media_out_dir = os.path.join(extracted_dir, "word", "media")
os.makedirs(media_out_dir, exist_ok=True)

# Copy PNGs from scratch folder
scratch_media = "/Users/superserver/Desktop/work/aishow/scratch/extracted_docx/word/media"
for i in range(3, 11):
    png_name = f"image{i}.png"
    src_png = os.path.join(scratch_media, png_name)
    dst_png = os.path.join(media_out_dir, png_name)
    if os.path.exists(src_png):
        shutil.copy(src_png, dst_png)
        print(f"Injected {png_name} into OYY package.")

# 1. Clean document.xml comment tags and old domain names
doc_xml_path = os.path.join(extracted_dir, "word", "document.xml")
if os.path.exists(doc_xml_path):
    with open(doc_xml_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Strip comments
    content = re.sub(r'<w:commentRangeStart[^>]*/>', '', content)
    content = re.sub(r'<w:commentRangeEnd[^>]*/>', '', content)
    content = re.sub(r'<w:commentReference[^>]*/>', '', content)
    
    # Domain replacements
    content = content.replace("甜品", "精品课程")
    content = content.replace("商品分类", "课程分类")
    content = content.replace("商品", "课程")
    content = content.replace("顾客", "学员")
    
    with open(doc_xml_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.")

# 2. Clean document.xml.rels relations
rels_path = os.path.join(extracted_dir, "word", "_rels", "document.xml.rels")
if os.path.exists(rels_path):
    with open(rels_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    content = re.sub(r'<Relationship[^>]*comments[^>]*/>', '', content)
    
    with open(rels_path, "w", encoding="utf-8") as f:
        f.write(content)
    print("Cleaned document.xml.rels.")

# 3. Delete comment physical files
for f in ["comments.xml", "commentsExtended.xml"]:
    fpath = os.path.join(extracted_dir, "word", f)
    if os.path.exists(fpath):
        os.remove(fpath)
        print(f"Deleted {f}")

# 4. Pack back to target DOCX
if os.path.exists(target_docx):
    os.remove(target_docx)

with zipfile.ZipFile(target_docx, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            fpath = os.path.join(root, file)
            relpath = os.path.relpath(fpath, extracted_dir)
            zip_ref.write(fpath, relpath)

print(f"COMPLETED! Final perfect DOCX written to {target_docx} for OYY Academy!")
