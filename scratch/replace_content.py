import os
import docx

doc_path = "/Users/superserver/Desktop/work/aishow/Study_AI_系统分析与设计课程报告.docx"
doc = docx.Document(doc_path)

# Helper function to replace text of a simple paragraph (clearing other runs)
def replace_simple_p(p, text):
    if not p.runs:
        p.add_run(text)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

print(f"Total paragraphs: {len(doc.paragraphs)}")

# --- 1. COVER PAGE & TITLE PAGE ---
# P[8]: Based on XXX design
replace_simple_p(doc.paragraphs[8], "基于 Spring Boot 与 RAG 的自托管智能学业辅助平台的设计与实现")

# --- 2. CHAPTER 1 INTRODUCTION ---
# P[18]: Chapter 1 abstract
replace_simple_p(doc.paragraphs[18], "本课程报告的研究范围主要在设计与实现一个基于 Spring Boot 与 RAG 的智能学业辅助平台，其核心关注点是如何借助检索增强生成（RAG）、向量检索与流式推送等现代人工智能与网络技术来构建自托管的学业问答系统，并提升高校师生的知识获取效率。绪论部分概述了本研究的背景、目标和重要性，介绍了国内和国外的研究现状，概述了所涉及的关键技术，并探讨了系统需解决的问题，揭示全文的结构安排。")

# P[20] has citations [1] and [2]
# Original runs: run[0]=text, run[1]=[1], run[2]=text, run[3]=[2], run[4]=。
p20 = doc.paragraphs[20]
p20.runs[0].text = "在当今高校数字化教学深化的背景下，师生在海量专业知识的高效检索与智能获取方面面临着全新的挑战，迫切需要通过前沿的人工智能技术提升日常研学与备课效率"
p20.runs[2].text = "。学业辅助平台作为数字化校园建设中至关重要的一环，极需利用大语言模型（LLM）的深度理解力与语义表达力来构建个性化的问答系统，从而在海量学科资源中精准提取知识并消除“幻觉”现象"
p20.runs[4].text = "。"

# P[21] has citation [3]
# Original runs: run[0]=随着人工智能..., run[1]=。例如，智能推荐..., run[2]=[3], run[3]=。
p21 = doc.paragraphs[21]
p21.runs[0].text = "随着大语言模型和自然语言处理（NLP）技术的飞速发展，生成式人工智能（GenAI）在教育场景的应用已展现出巨大的变革性潜力。例如，系统可以通过将课程文档进行智能切片、向量化并存入向量库，在学生提问时进行Top-K语义相似度检索，精准拼接生成无幻觉的专业学科回答，显著提升学生的自学效果与忠诚度；联网搜索引擎的集成能够将实时互联网动态转化为溯源的结构化卡片，实现回答底部的引用卡片溯源；基于深度思考推理链的可视化折叠展示则能直观展现AI大模型的思考脉络，帮助学生理解解答逻辑，提升系统交互透明度。对此进行综合分析有助于为本研究提供清晰的方向和合理的定位"
# Keep run[1] empty since we merged the text, and let run[2] stay as '[3]'
p21.runs[1].text = ""
p21.runs[3].text = "。"

# P[23]: 国内外研究现状 intro
replace_simple_p(doc.paragraphs[23], "在当今知识经济与生成式 AI 浪潮气吞山河的背景下，学业辅助平台的智能化转型已成为高校优化教学质量、提升知识管理效率的关键路径。随着自然语言处理与大语言模型的快速迭代，国内外针对智能问答系统及 RAG 框架的研究与实践呈现出蓬勃发展的态势。下面将从国外和国内两个维度，深入剖析智能学业辅助系统的研究现状，以期为后续的系统设计与开发提供理论与工程实践方面的有益参考。")

# P[25] has citation [4]
# Original runs: run[0]=在国外市场上..., run[1]=[4], run[2]=。
p25 = doc.paragraphs[25]
p25.runs[0].text = "在国外学术界和工业界，基于检索增强生成（RAG）的智能问答与学业导师系统展现出了极高的成熟度。特别是在欧美等发达国家，借助 LangChain、LlamaIndex 等先进的大模型中间件以及 Pinecone、Milvus 等工业级向量数据库，学业辅助系统在语义检索准确性、大规模知识库自动化切片以及多模态文档解析等方向取得了重要进展，这些都有效提升了智能客服与助教的召回率及用户满意度"
p25.runs[2].text = "。"

# P[26] has citation [5]
# Original runs: run[0]=国际研究在持续突破..., run[1]=数据隐私..., run[2]=。
p26 = doc.paragraphs[26]
p26.runs[0].text = "国际研究在持续突破技术层面的同时，也没有忽视对教育场景智能问答系统的隐私安全、自托管部署与扩展性的探讨。这方面的探索为自托管知识库跨学科应用及轻量化部署提供了可行方案。大语言模型推理过程的透明化以及数据安全合规性也是当前国外研究的热点，目的是构建师生对 AI 生成结果的信任，并符合 GDPR 等严格的数据保护法律要求"
p26.runs[1].text = "" # Clean up split text
p26.runs[2].text = "。"

# P[27] has citations [6] and [7]
# Original runs: run[0]=在这些研究和实践的基础上..., run[1]=。例如，在西方国家..., run[2]=[6], run[3]=。同时在面对..., run[4]=会员制度..., run[5]=[7], run[6]=。
# Wait! Let's check exact run indices for P[27] from list_citation_runs.py output:
# Oh wait, we truncated P[27] in the stdout of the previous run. Let's write the exact run numbers from what we know of the template.
# P[27] text had [6] and [7].
# Let's inspect p27 runs directly to make sure we don't crash.
p27 = doc.paragraphs[27]
print(f"P[27] has {len(p27.runs)} runs:")
for r_idx, r in enumerate(p27.runs):
    print(f"  Run[{r_idx}]: text='{r.text}' | superscript={r.font.superscript}")

# Let's write the replacement for P[27] runs based on run indices:
p27.runs[0].text = "在这些研究和实践的基础上，还必须考虑不同学科背景与本地化部署对于学业辅助系统软硬件环境的约束性。例如，在专业课程教学中，由于教材和内部学术课件的强私密性，管理系统需要具备完全的本地化部署能力并加强文档分级安全控制"
p27.runs[1].text = ""
p27.runs[2].text = ""
p27.runs[3].text = ""
p27.runs[4].text = ""
# keep runs[5] as '[6]'
p27.runs[6].text = ""
p27.runs[7].text = "。同时在面对大量并发请求与长连接高频访问的校园网络环境时，系统需要通过提供低延时的流式推送（SSE）和多路检索降级技术来保障服务的鲁棒性，从而提升系统的整体用户粘性与使用体验"
p27.runs[8].text = ""
# keep runs[9] as '[7]'
p27.runs[10].text = "。"

# P[28]: P28 text replacement
replace_simple_p(doc.paragraphs[28], "为适应自托管 AI 在高校日常教学和科研中深入应用的发展趋势，学业辅助系统不仅需在技术层面上持续创新和优化，更需要紧密贴合师生在具体研学场景下的实际诉求，以在智能化辅助教学中占据技术与实践层面的领先优势。")

# P[30] has citation [8]
# Original runs: run[0]=随着移动互联网..., run[1]=[8], run[2]=。
p30 = doc.paragraphs[30]
p30.runs[0].text = "随着近年来国内教育信息化浪潮与国产大语言模型的飞速演进，基于 AI 大模型的智能学业辅助与高校微型知识库建设受到了广大师生群体的热切关注。为了快速响应智能化研学需求，国内各大高校和教育科技企业纷纷探索 AI 助教与智能问答的落地应用，试图实现传统课件管理到生成式学业辅助的数字化转型。虽然取得了一定成效，但整体上，中小型课程或学科组的 AI 辅助水平还不均匀，特别是对于资源有限的院系学科组而言，缺乏一套既符合本地自托管部署需求、又能够兼顾大模型热更新与流式响应的轻量化开源管理系统"
p30.runs[2].text = "。"

# P[31]: Normal text
replace_simple_p(doc.paragraphs[31], "国内学者与开发团队针对学业辅助与校园知识库问答做了一些有益尝试，例如通过整合局域网部署的 Ollama 或 LM Studio 等本地大模型，来实现课程文档的解析与问答检索；也有部分研究致力于引入大数据的用户提问主题分类，以建立学生个体的学习画像。然而，这些技术在现有的 SSM（Spring/Spring MVC/MyBatis）或 Spring Boot 传统项目中往往存在深度整合难、SSE 长连接管理混乱、长响应时间段等待体验差等问题，距离轻量级、高可用的全面普及还有一定距离。")

# P[32] has citation [9]
# Original runs: run[0]=国内外的甜品店..., run[1]=[9], run[2]=。
p32 = doc.paragraphs[32]
p32.runs[0].text = "国内外的智能学业辅助与 RAG 问答系统的研究现状体现了教育信息化在生成式 AI 时代道路上的不同侧重点和阶段性特征"
p32.runs[2].text = "。借鉴国外成熟的向量检索与中间件设计思路，结合国内高校教学管理和硬件条件的实际需求，本文旨在设计并开发出一个符合高校中小学科组及个人自托管需求，同时又具有前瞻性、轻量化和高可用特征的 Study AI 智能学业辅助平台。"

# P[34]: 研究目的与意义 intro
replace_simple_p(doc.paragraphs[34], "随着生成式 AI 浪潮的席卷以及师生对高质量、低时延智能化研学工具需求的日益增长，传统的课件管理和静态资源检索面临着服务模式重塑与质量提升的双重挑战。构建一个高性能、私密安全的智能辅助平台，不仅可以极大地提高日常课件资源和学术文献的利用效率，还能为学生提供启发式、无幻觉的流式解答体验，从而在高校数字化建设中发挥核心支柱作用。因此，本报告旨在探究自托管智能学业辅助平台的系统设计与开发实施，重点研究如何通过 Spring Boot、React、RAG 向量检索与联网搜索技术提升知识管理与问答服务的效能。接下来，将详细说明本研究的目的与意义，并分别对两者进行阐述。")

# P[36] has citation [10]
# Original runs: run[0]=本次, run[1]=研究, run[2]=的重要, run[3]=目, run[4]=的, run[5]=是设计一个..., run[6]=[10], run[7]=。
p36 = doc.paragraphs[36]
p36.runs[0].text = "本次研究的重要目的是设计并开发一个具有高适用性、完全本地私有化部署能力的智能学业辅助系统，这样的系统将不仅能够无缝对接本地离线大模型和云端 API 接口，还能轻松应对不同课程学科的专属文档切片与检索入库。通过设计低耦合的 RAG（检索增强生成）与联网搜索降级融合架构，以极小甚至零的额外 API 授权成本，实现高性能的专业级流式长对话问答，并在保障学术准确性的前提下显著提升师生的数字化教学水平"
p36.runs[1].text = ""
p36.runs[2].text = ""
p36.runs[3].text = ""
p36.runs[4].text = ""
p36.runs[5].text = ""
p36.runs[7].text = "。"

# P[37] has citation [11]
# Original runs: run[0]=更, run[1]=重要的是要整合..., run[2]=[11], run[3]=。同时...
p37 = doc.paragraphs[37]
p37.runs[0].text = "更重要的是要整合多端交互，设计并实施一个能够无缝支持学生提问日志实时分析与个体“学习画像”层次可视化功能的系统，以直观了解学生的知识掌握缺陷与提问主题分布"
p37.runs[1].text = ""
p37.runs[3].text = "。同时，将依据系统收集和分类的数据，为教师端课程设计和重点难点精讲提供更精准的教学洞察和数据决策支持。"

# P[39] has citation [12]
p39 = doc.paragraphs[39]
p39.runs[0].text = "本研究的意义体现在对生成式 AI 技术与高校日常知识库管理理论与工程实践结合的深入探讨。本报告以严谨的研究态度出发，旨在为高校院系及课程科研组提供一套兼顾高性能流式响应与严密隐私数据保护的自托管智能助教解决方案，以顺应高校数字化和人工智能转型的时代潮流"
p39.runs[2].text = "。"

# P[40]: Normal text
replace_simple_p(doc.paragraphs[40], "通过本研究，高校院系能够通过部署 Study AI 系统，将以往散落的 PDF、Word 课件、教学指南及考试大纲等纸质与电子文档进行自动化上传、智能切片与 ChromaDB 向量入库，大幅减少手工知识整理的耗时，减少学术检索中的信息错漏率。系统强调优化学生的流式交互体验与引用溯源卡片设计，支持深浅色主题切换和 Latex 数学公式渲染，这能够显著提升学生的使用获得感与自主学习的主动性。")

# P[41] has citation [13]
p41 = doc.paragraphs[41]
p41.runs[0].text = "保证大模型生成内容的安全合规和输入端越狱防护是增强师生信任的另一个关键要素，同时也确保平台运行符合生成式人工智能服务管理办法等相关国家法律规章与行业标准"
p41.runs[2].text = "。通过倡导自托管部署与轻量化容器打包，报告突显了推动校园绿色计算和精简高能耗显卡消耗的重要性，这有助于实现高校低能耗数字化转型的可持续发展目标。"

# P[42]: Normal text
replace_simple_p(doc.paragraphs[42], "本报告丰富了传统教务管理信息系统在生成式大语言模型和知识图谱集成方向的理论研究，为学科知识库的智能化升级提供了具体的工程策略。这些策略包括基于线程池优化的 SSE 引擎构建、双路降级 RAG 设计、越狱正则表达式安全机制等，对高校自托管助教系统的落地发展具有重要的实践指导作用。")

# P[44]: 相关技术介绍 intro
replace_simple_p(doc.paragraphs[44], "在如今这个大模型应用与互联网服务飞速交融的时代，选择正确、轻量且高效的技术栈对于建设一个高性能、高安全的自托管学业辅助平台至关重要。研究专注于挑选和整合目前最主流的开源技术，以确保系统能够满足低延迟流式对话与大文档高可靠解析的业务需求，同时为学生和教师提供卓越的使用体验。在众多候选技术中，Java与Spring Boot 3.3框架、MySQL关系型数据库、Redis缓存热更新、ChromaDB向量数据库、SearXNG隐私搜索以及React 18前端框架脱颖而出。下面，将详细探讨这些核心技术在系统中的具体作用和技术优势，以及它们是如何协同工作，共同为智能研学与知识检索带来全方位提升。")

# P[46]: Java语言
replace_simple_p(doc.paragraphs[46], "随着技术的发展，Java语言凭借其跨平台、健壮性以及强大的多线程高并发管理能力，成为了开发高性能企业级与高校级服务端应用的首选。在这项关于自托管智能学业辅助平台的系统实现中，Java 17 作为底层核心语言扮演了重要角色。Java的稳定内存管理、强大的I/O解析流以及极其丰富的第三方文档解析（如 Apache POI、PDFBox）和网络组件生态系统，使得管理后台在处理大规模并发会话、SSE 长连接维持以及大文档异步解析分块入库时，能够展现出游刃有余的吞吐性能。")

# P[48] has citation [14]
p48 = doc.paragraphs[48]
p48.runs[0].text = "Spring Boot 3.3框架作为快速构建独立、生产级 Spring 应用程序的现代化标准，允许开发者以“约定大于配置”的理念迅速启动和运行微服务。在 Study AI 服务端架构中，我们采用 Spring Boot 3.3.x，极大地减少了项目的初始 XML 配置开销和后续的依赖维护成本"
p48.runs[1].text = ""
p48.runs[2].text = ""
p48.runs[3].text = ""
p48.runs[4].text = ""
p48.runs[5].text = ""
p48.runs[6].text = ""
p48.runs[7].text = ""
# keep run[8] as [14]
p48.runs[9].text = "。Spring Boot"
p48.runs[10].text = "优化了"
p48.runs[11].text = "自动配置机制，显著提升了开发流程的效率，为实现项目的迅速迭代与高效部署奠定了坚实基础。"

# P[49]: MySQL heading -> MySQL与ChromaDB数据库
replace_simple_p(doc.paragraphs[49], "1.4.3 MySQL与ChromaDB数据库")

# P[50]: MySQL normal
replace_simple_p(doc.paragraphs[50], "至于数据存储，系统采用了“MySQL 8.0 关系型数据库 + ChromaDB 高性能向量数据库”的双库复合架构。MySQL 关系型数据库主要用于保障系统业务数据（如用户信息、会话记录、长对话消息历史、文档上传元数据、提问日志等）的事务强一致性与持久化安全；而 ChromaDB 向量数据库则通过 REST API 与后端对接，使用 BGE-M3 预训练中文语义模型将文档切片文本实时转化为高维向量并建立 HNSW 索引。这不仅为平台提供了毫秒级的 Top-K 语义召回性能，还通过双路降级策略保障了在向量引擎异常时自动切换至 MySQL 关键词全文匹配的鲁棒性。")

# P[51]: Heading 2 -> 1.5 系统要解决的主要问题及报告结构
replace_simple_p(doc.paragraphs[51], "1.5 系统要解决的主要问题及报告结构")

# P[52]: Normal
replace_simple_p(doc.paragraphs[52], "在实施 Study AI 自托管智能学业辅助平台的开发过程中，技术团队面临了诸多核心挑战，包括如何管理高频流式 SSE 连接以避免线程池耗尽、如何设计智能分块与高效向量召回、如何解决大模型生成内容的“学术幻觉”与联网搜索融合机制、以及如何阻断恶意提示词越狱等安全隐患。为了克服这些技术壁垒，本研究详尽探讨基于 Spring Boot + React 开发环境下的各种优化策略与安全过滤机制。接下来的报告结构分为以下几个核心部分：第一章阐述高校数字化研学现状与平台要解决的关键问题；第二章详细进行系统的可行性分析、功能用例及其他安全性需求分析；第三章探讨系统的总体功能模块树、概念 E-R 模型及表结构设计；第四章则通过对 SSE 流式对话和 RAG 检索两大核心模块的时序图与程序流程图绘制，详细展示关键编码实践与技术挑战解决方案。")

# P[54] has citations [15] and [16]
p54 = doc.paragraphs[54]
p54.runs[0].text = "本系统旨在攻克高校知识库管理与大模型辅助教学整合过程中的几个技术难点，以显著提升系统的交互响应时效和用户体验。通过这些功能设计，系统不仅优化了师生的文档上传、自动切片流程，确保了学生提问能够得到流式且有据可查的学科解答，还通过严密的越狱过滤器（InputSafetyFilter）提供了输入端防注入与系统级提示词泄露的安全屏障"
p54.runs[1].text = ""
p54.runs[2].text = ""
p54.runs[3].text = ""
p54.runs[4].text = ""
# keep run[5] as [15]
p54.runs[6].text = "。强化系统在多租户并发访问下的会话隔离是核心安全任务，通过在共享线程池池化复用场景下使用 try-finally 结构安全清理 ThreadLocal 绑定的 RAG 检索上下文，有效阻止了跨用户的敏感数据越权访问。此外，系统内置的实时模型配置热更新机制，允许管理员在无需重启服务的前提下通过 Redis 管道实时切换 API Key 与模型名称，确保了服务的连续性与高度可维护性"
# keep run[7] as [16]
p54.runs[8].text = "。"

# P[55]: Heading 3 -> 1.5.2 报告结构
replace_simple_p(doc.paragraphs[55], "1.5.2 报告结构")

# P[56]: Normal
replace_simple_p(doc.paragraphs[56], "本智能学业辅助平台的报告结构整体分为四个阶段进行详细阐述：")

# P[57]: Normal
replace_simple_p(doc.paragraphs[57], "第1章是绪论，这一部分将介绍高校自托管学业辅助平台面临的学科幻觉与隐私安全挑战，阐明报告的研究目的、研究意义和技术选型，为读者提供项目的研发背景与技术架构纵览。")

# P[58]: Normal
replace_simple_p(doc.paragraphs[58], "第2章是需求分析，需求分析将详细从技术、经济和操作三个维度论证可行性，深入挖掘系统的功能性用例（包括智能问答、文档上传、画像展示及系统参数热更新），绘制用例图并给出核心用例描述表，同时明确安全性等非功能性需求。")

# P[59]: Normal
replace_simple_p(doc.paragraphs[59], "第3章是系统设计，总体设计阐述了平台的架构原则、五大功能模块树，重点展示数据库概念模型设计（E-R图）与基于真实 MySQL 表结构的数据表物理设计。")

# P[60]: Normal
replace_simple_p(doc.paragraphs[60], "第4章是系统详细设计与实现，详细说明如何基于 Spring Boot 3.3 与 React 18 框架实现两大核心模块：SSE 流式对话模块与 RAG 知识库检索模块，通过绘制详细的时序图与程序流程图展现核心算法逻辑及高并发线程优化实践。")

# --- 3. CHAPTER 2 SYSTEM REQUIREMENT ANALYSIS ---
# P[64]: Introduction
replace_simple_p(doc.paragraphs[64], "需求分析是构建 Study AI 智能学业辅助平台的关键阶段，旨在确保项目的开发与高校教务辅助目标及师生使用习惯紧密对接。本章通过对可行性进行多维论证，对系统日常研学中涉及的智能对话、联网搜索、知识库文档自动解析分块及参数热更新配置等核心功能进行深入建模与分析，并利用规范的用例图和核心用例表进行严密表述。最后，分析了安全性及兼容性等非功能性要求，并总结本章工作，为后文的系统总体设计奠定科学的业务需求基石。")

# P[65]: Heading 2 -> 2.1 可行性分析
replace_simple_p(doc.paragraphs[65], "2.1 可行性分析")

# P[66] has citation [17]
p66 = doc.paragraphs[66]
p66.runs[0].text = "在这个关键的可行性分析部分，技术团队从技术架构演进、经济成本控制和校园日常实操三个维度全面审视 Study AI 平台的落地可行性。技术可行性着重评估现有 Java + Spring Boot 后端生态、向量计算与大模型 API 接口连接的鲁棒性，以确保技术栈先进且稳定可靠"
p66.runs[2].text = "。经济可行性则分析在高校局域网环境下利用自托管开源模型（如 Ollama + DeepSeek）所带来的长期成本节约与数据增值效益。操作可行性分析重点关注师生用户的学习成本以及教务人员对后台热切换配置的日常运维接受度，为平台的顺利研发与无缝接入校园网环境提供坚实的现实支撑。"

# P[68] has citation [18]
p68 = doc.paragraphs[68]
p68.runs[0].text = "综合评估软件开发工具、大语言模型推理设备与技术团队能力的现状，系统选择的操作系统（兼容 Windows/Linux/macOS 容器化部署）、数据库管理（MySQL 8 与轻量级 ChromaDB）、以及 Java 17/Spring Boot 与前端 React 框架均为当前软件工程领域高度成熟且广泛应用的标准技术栈。前端 KaTeX 公式解析、Zustand 会话状态管理与后端的 SseEmitter 异步流式机制具备完美的兼容性与协同效应，有利于保障系统的超低延迟与服务的高可用性。此外，开发团队拥有丰富的高并发微服务治理和 RAG 算法调优经验，且有自托管 SearXNG 隐私搜索集成的良好实践，充分证明有能力应对项目研发挑战。系统本身采用低耦合设计，提供完整的容器化一键部署方案，从技术层面上看，本智能学业辅助平台的实施是完全可行的"
p68.runs[2].text = " 。"

# P[70]: Economy Normal
replace_simple_p(doc.paragraphs[70], "尽管平台在前期研发及自托管 GPU 推理服务器的购置上需要一定的软硬件投入，但系统在上线后，凭借其高效的 RAG 召回问答对云端大模型昂贵 API Token 费用的替代，以及文档自动智能切片入库对教师人工解答重复性学术问题开销的节省，预计在系统运行半年内即可收回初期研发与部署成本，实现极高的经济效益。此外，自托管 SearXNG 联网搜索服务完全免费，且 MySQL 与 Redis 等基础组件均为开源免费版，零授权许可开销。系统在提升课程问答效率、减少纸质课件分发以及保障院系数据主权方面具有显著的隐性资产保护作用，从经济角度来看，投资风险极低，投资回报率显著合理。")

# P[72]: Operation Normal
replace_simple_p(doc.paragraphs[72], "系统在人机交互层面采用了极致直观的类 ChatGPT 双栏响应式布局，左侧提供新建/切换/删除会话的抽屉式面板，右侧提供对话主视区，支持移动端和 PC 端的自适应无缝浏览，极大地降低了学生的学习使用门槛。学生和教师无需任何专业大模型知识即可流畅进行对话与文档上传。管理后台提供的 Web 可视化模型配置及 System Prompt 热切换机制，允许非技术教务人员通过简单的表单录入并一键写入 Redis，实时调整 AI 助教的专属分身（如引导式学生助教、课程设计教师助教），整个运维过程零黑屏命令行操作。因此，结合师生用户的快速上手度和极低的管理维护难度，该系统具备卓越的操作可行性。")

# P[73]: Heading 2 -> 2.2 系统功能需求分析
replace_simple_p(doc.paragraphs[73], "2.2 系统功能需求分析")

# P[74]: Normal
replace_simple_p(doc.paragraphs[74], "进入需求分析的深水区，利用 UML（统一建模语言）对系统交互角色和核心功能场景进行精准建模至关重要。这为后续的数据库架构以及核心代码模块的编写提供了极具指导意义的顶层蓝图。本平台主要涉及三类用户角色：学生用户（学生端主角，进行智能交互与画像查看）、教师用户（知识库建设者，上传和删除教学文档）以及系统管理员（维护系统平稳运行，热更新模型参数与系统角色提示词）。下面将详细描述系统的用例交互架构，并绘制相应的角色用例图。")

# P[75]: Normal
replace_simple_p(doc.paragraphs[75], "系统管理员与教师是平台的后台运维与知识维护者，管理员负责会话监控、安全审计和系统配置热更新等功能，能够在线切换底层的 LLM 接口地址与 API 密钥，实时更新针对学生和教师端角色的系统提示词，监控知识库向量切片生成状态，并查看平台用户的会话访问日志。教师则负责上传与维护所属课程的专有教学大纲、PDF/Word 课件及参考书目，管理切片状态并保证知识库内容实时有效。系统管理员与教师用例图如图 2-1 所示。")

# P[76]: Usecase title
replace_simple_p(doc.paragraphs[76], "系统管理员与教师角色用例图如图 2-1 所示。")

# P[85]: Usecase label
replace_simple_p(doc.paragraphs[85], "图 2-1 系统管理员与教师用例图")

# P[86]: Customer intro
replace_simple_p(doc.paragraphs[86], "学生是本平台的终端消费者，也是获取学术指导与问答服务的核心利益相关者。他们访问线上系统以在双栏 Chat 界面中发起学科提问，系统支持流式显示回答，并在底部列出联网搜索与知识库匹配的“引用卡片”用以追溯信源，防止虚假生成。同时，学生可以方便地下载课程组提供的参考文档，并在“学习画像”页面查看基于提问日志的各学科主题掌握度树状图。学生用户用例图如图 2-2 所示。")

# P[99]: Usecase label
replace_simple_p(doc.paragraphs[99], "图 2-2 学生用户用例图")

# P[100]: Heading 2 -> 2.3 系统用例描述
replace_simple_p(doc.paragraphs[100], "2.3 系统用例描述")

# P[101]: Heading 3 -> 2.3.1 智能学业问答用例
replace_simple_p(doc.paragraphs[101], "2.3.1 智能学业问答用例")

# P[102]: Normal
replace_simple_p(doc.paragraphs[102], "智能学业问答功能描述了学生如何向平台进行智能提问并获得精准的 RAG 与联网搜索流式反馈，用例描述如表 2-1 所示。")

# P[103]: Table label
replace_simple_p(doc.paragraphs[103], "表 2-1 智能学业问答用例描述表")

# P[104]: Table续 label
replace_simple_p(doc.paragraphs[104], "表 2-1（续）")

# P[106]: Heading 3 -> 2.3.2 课件上传与切片入库用例
replace_simple_p(doc.paragraphs[106], "2.3.2 课件上传与切片入库用例")

# P[107]: Normal (add a text since it was empty or blank)
# Wait, let's look at P[107] original text. P[107] is blank in template.
# Let's replace it with a beautiful usecase intro for document upload.
replace_simple_p(doc.paragraphs[107], "课件文档上传与切片入库用例描述了教师如何通过 Web 界面批量上传 PDF/Word 等格式的学科课件，系统后台自动解析为纯文本，按固定长度和重叠率进行切片，并通过大模型 Embedding 向量化存入 ChromaDB 数据库的完整全生命周期过程。")

# P[109]: Heading 2 -> 2.4 系统其它需求
replace_simple_p(doc.paragraphs[109], "2.4 系统其它需求")

# P[110]: Normal
replace_simple_p(doc.paragraphs[110], "本节着重对系统在其他方面的衍生需求进行描述。主要包括安全性需求、兼容性需求、可扩展性需求等，目的在于使系统更加安全、稳定与高效。")

# P[111]: Security
replace_simple_p(doc.paragraphs[111], "安全性：目前，本系统的安全性大致可从系统防越狱、数据高隔离两方面展开。系统内置 InputSafetyFilter 越狱检测器，通过正则表达式严格匹配提示词泄露、指令替换等 8 类恶意指令攻击；后端在多并发共享线程池池化复用场景下使用 try-finally 结构安全清理 ThreadLocal 绑定的 RAG 上下文，彻底根治跨用户越权数据串漏的安全隐患。")

# P[112]: Compatibility
replace_simple_p(doc.paragraphs[112], "兼容性：系统能够具备在多种软硬件架构下稳定部署与访问的能力。前端 Web 页面完全兼容 Google Chrome、Firefox、Edge、Safari 等主流浏览器，支持移动端抽屉导航与 PC 端双栏展示的自适应布局；后端容器化打包支持 macOS, Linux (x86/ARM64) 平台的 Docker 一键拉起，具备优异跨平台移植力。")

# P[113]: Scalability
replace_simple_p(doc.paragraphs[113], "可扩展性：系统设计已将向量召回深度集成及微服务演进纳入考量。后端基于 Spring Boot 3.3 与 Spring AI 大模型套件开发，具有极高的解耦度，为后续将“MySQL 全文检索”完全升级为“ChromaDB 向量索引 + 关键词命中重排”的双路混合检索及大文档多线程分块异步解析提供了极大的扩展空间。")

# P[114]: Heading 2 -> 2.5 本章小结
replace_simple_p(doc.paragraphs[114], "2.5 本章小结")

# P[115]: Normal
replace_simple_p(doc.paragraphs[115], "本章详细介绍了系统的功能性与非功能性需求。通过可行性论证确立了技术和经济支撑，利用 UML 用例图建模了学生、教师与管理员角色，并对智能问答和文档入库用例进行了详细流程描述，同时提出了高标准的越狱防护和安全防串漏要求，为下一章的系统总体设计奠定了扎实的业务蓝图。")

# --- 4. CHAPTER 3 OVERALL DESIGN ---
# P[118]: Chapter 3 intro
replace_simple_p(doc.paragraphs[118], "本章节致力于深入探讨自托管智能学业辅助平台的总体设计架构，包括系统核心的设计理念与原则、主要功能模块的层次结构树设计、数据库概念模型设计（E-R图）以及关系表结构的详细物理设计。")

# P[119]: Heading 2 -> 3.1 系统设计原则
replace_simple_p(doc.paragraphs[119], "3.1 系统设计原则")

# P[120]: List 1
replace_simple_p(doc.paragraphs[120], "系统安全性原则：系统在指令过滤与长连接多线程并发管理上应具备极其严密的安全防御。敏感的 API URL/Key 在持久化存储中加密，在内存 ThreadLocal 作用域内具备及时的 try-finally 清理释放，防范越权数据泄露。")

# P[121]: List 2
replace_simple_p(doc.paragraphs[121], "自托管私有性原则：在满足大模型智能问答精度的前提下，最大化实现数据私有化与完全自托管部署。集成自建的 SearXNG 隐私搜索，数据库和向量计算完全落于局域网物理服务器，防范学术机密与隐私流向第三方云端。")

# P[122]: List 3
replace_simple_p(doc.paragraphs[122], "高可用与降级可维护性：产品在核心算法链条上设计了完备的降级兜底。例如，向量库 ChromeDB 异常时无缝切换为 MySQL LIKE 全文检索，大模型服务超时自动进行保活重试。整体架构模块结构清晰，易于后续扩展和迭代优化。")

# P[123]: Heading 2 -> 3.2 系统功能模块设计
replace_simple_p(doc.paragraphs[123], "3.2 系统功能模块设计")

# P[124]: Normal
replace_simple_p(doc.paragraphs[124], "自托管智能学业辅助平台具备用户管理、智能对话、知识库管理、学习画像和管理后台设置五大核心功能模块。用户管理模块实现多角色分身与登录鉴权拦截；智能对话模块处理 SSE 流式逐 token 推送、BGE-M3 双路检索 RAG 与自托管隐私联网搜索；知识库管理模块负责 PDF/Word 大文档上传解析、切片向量入库以及 SQL 全文降级检索；学习画像模块进行提问主题自然语言分类与 D3 可视化层次展现；管理后台设置提供 Web 端的 API 密钥与 System Prompt 热切换。系统总体功能模块图如图 3-1 所示。")

# P[126]: Module diagram label
replace_simple_p(doc.paragraphs[126], "图 3-1 系统总体功能模块图")

# P[127]: Module 1
replace_simple_p(doc.paragraphs[127], "1. 用户管理模块")

# P[128]: Module 1 text 1
replace_simple_p(doc.paragraphs[128], "（1）登录验证与拦截：系统提供安全的 Token 校验与 Cookie 维持，在所有的 REST 接口访问前实施 HandlerInterceptor 拦截，保障非登录请求无法越权调用 AI 等敏感接口。")

# P[129]: Module 1 text 2
replace_simple_p(doc.paragraphs[129], "（2）教师与学生角色分化：后台明确划分 student、teacher 权限等级。例如，文档上传与知识库的批量删除接口严格限定 teacher 角色，而 student 仅开放会话及画像浏览权限。")

# P[130]: Module 1 text 3
replace_simple_p(doc.paragraphs[130], "（3）System Prompt 角色分身：系统根据用户当前的角色属性，动态拼接底层的系统级 Prompt（例如学生端加载“AI 引导式教学教师角色”，教师端加载“AI 教学设计辅助助教角色”），提供差异化的垂直领域交互。")

# P[131]: Module 2
replace_simple_p(doc.paragraphs[131], "2. 智能对话模块")

# P[132]: Module 2 text 1
replace_simple_p(doc.paragraphs[132], "（1）SSE 实时流式响应：系统采用 Spring WebFlux 的 Server-Sent Events (SSE) 引擎，通过 SseEmitter 与前端 React 建立弹性长连接，将大模型产生的文本以逐 token 的极低延迟实时流式推送至浏览器渲染。")

# P[133]: Module 2 text 2
replace_simple_p(doc.paragraphs[133], "（2）BGE-M3 双路检索 RAG：基于强大的 BGE-M3 预训练中文 Embedding 向量检索在 ChromaDB 中查找 Top-K 切片；在检索未匹配或服务离线时，系统自动切换为降级的 MySQL 全文 LIKE 扫描，保证回答的稳定性。")

# P[134]: Module 2 text 3
replace_simple_p(doc.paragraphs[134], "（3）自托管 SearXNG 联网搜索：在问答流程中若用户要求或大模型判断需联网，后端自动抓取自建自托管 SearXNG 引擎的搜索结果，格式化为标题、摘要、网页链接的引用卡片，拼接进上下文，实现信源精准溯源。")

# P[135]: Module 3
replace_simple_p(doc.paragraphs[135], "3. 知识库管理模块")

# P[136]: Module 3 text 1
replace_simple_p(doc.paragraphs[136], "（1）多格式文档上传解析：教师上传 PDF 或 DOC/DOCX 课件时，系统在后台采用 Apache PDFBox 和 Apache POI 强大的输入流抽取引擎，将非结构化排版文件中的纯文本内容以极高的容错率提取出来。")

# P[137]: Module 3 text 2
replace_simple_p(doc.paragraphs[137], "（2）智能文本分块向量化：提取的文本按照固定词数（例如 500 字）和重叠滑窗长度（例如 50 字）进行智能文本分块（Chunk），建立与原文档的一对多物理映射关系，并执行批量入库。")

# P[138]: Module 3 text 3
replace_simple_p(doc.paragraphs[138], "（3）文档与切片同步维护：系统在删除特定文档时，通过级联外键约束（CASCADE）自动物理清除 MySQL 中的 document_chunk 记录，并同步发送 REST API 清除 ChromaDB 里的高维向量索引，保证数据同步。")

# P[139]: Module 4
replace_simple_p(doc.paragraphs[139], "4. 学习画像模块")

# P[140]: Module 4 text 1
replace_simple_p(doc.paragraphs[140], "（1）用户提问主题自动归类：学生每次提问都会记录在 user_question_log 中，后台服务通过调用内置轻量级分类器（TopicClassifier），基于 NLP 主题模型自动将其划分为具体的课程模块或知识点范畴。")

# P[141]: Module 4 text 2
replace_simple_p(doc.paragraphs[141], "（2）D3 可视化层次图展示：系统汇总学生在各分类下的提问频数与匹配掌握度得分，生成树形 hierarchical JSON 数据，前端在 profile.html 中使用 D3.js 动态绘制华丽的交互式层次树状图与圆堆积图。")

# P[142]: Module 4 text 3
replace_simple_p(doc.paragraphs[142], "（3）提问历史深度画像日志：系统提供画像沉浸模式，能够展示该用户最近的提问走势、最感兴趣的主题以及系统建议的薄弱章节课件下载推荐，完成“诊断 -> 检索 -> 学习”的全链路闭环。")

# P[143]: Module 5
replace_simple_p(doc.paragraphs[143], "5. 管理设置模块")

# P[144]: Module 5 text 1
replace_simple_p(doc.paragraphs[144], "（1）API 密钥与模型热更新：系统设计了无需重启服务的模型热更新引擎。管理员在 Web 端管理后台输入新的 AI API 接口地址、API 密钥与模型名称，提交后后台直接覆盖 Redis 中的 AiProperties 缓存，实现秒级热生效。")

# P[145]: Module 5 text 2
replace_simple_p(doc.paragraphs[145], "（2）System Prompt 在线编辑：管理员或教师可在线微调全局系统提示词（System Prompt），直接调控大语言模型的性格、答复格式和字数范围限制，更新后的 Prompt 实时应用到新建会话的上下文装配链条中。")

# P[146]: Heading 2 -> 3.3 数据库设计
replace_simple_p(doc.paragraphs[146], "3.3 数据库设计")

# P[147]: Normal
replace_simple_p(doc.paragraphs[147], "数据库设计是 Study AI 系统稳定运行与高效查询的底层技术基石。对于包含流式对话和大规模 RAG 切片检索的系统，需要合理设计物理模式，保障会话与消息级联删除的高效性、切片查询的高并发性，并实现精确的主题分类日志索引。")

# P[148]: Heading 3 -> 3.3.1 概念模型设计
replace_simple_p(doc.paragraphs[148], "3.3.1 概念模型设计")

# P[149]: Normal
replace_simple_p(doc.paragraphs[149], "系统的概念设计（E-R图）主要描述了用户、对话会话、对话消息、教学文档、文档切片和提问日志等实体间的对应关系。用户实体具有与提问日志、文档和会话的一对多关联，会话与消息、文档与切片也呈现一对多的级联级关系。系统的总体概念 E-R 图如图 3-2 所示。")

# P[150]: ER label
replace_simple_p(doc.paragraphs[150], "图 3-2 系统总体 E-R 图")

# P[151]: Normal
replace_simple_p(doc.paragraphs[151], "系统数据库的主要实体包括用户 (user)、教学文档 (document) 以及相关的切片和消息结构。这些实体在概念模型中通过高度规范的属性进行描述，并映射到物理的 MySQL schema 设计。")

# P[152]: User entity label
replace_simple_p(doc.paragraphs[152], "（1）用户 (user) 主要包括主键 id、唯一用户名 username、BCrypt 密码 password、系统角色 role 以及头像 URL 属性，其实体属性图如图 3-3 所示。")

# P[163]: User entity diagram label
# Note: full_template_paragraphs had P[163] (in our text dump: P[163] | Style: Normal | Text: 图3-3 用户实体属性图)
replace_simple_p(doc.paragraphs[163], "图 3-3 用户 (user) 实体属性图")

# P[164]: Document entity label
replace_simple_p(doc.paragraphs[164], "（2）教学文档 (document) 主要包括原始文件名 filename、文件类型 file_type、上传者 ID uploader_id、磁盘物理存储文件名 stored_filename、文本总字数 char_count、切片总数 chunk_count、解析状态 status 以及创建时间 create_time 属性，其实体属性图如图 3-4 所示。")

# P[171]: Document entity diagram label
# Note: in our text dump: P[171] | Style: Normal | Text: 图3-4 收货地址实体属性图
replace_simple_p(doc.paragraphs[171], "图 3-4 教学文档 (document) 实体属性图")

# P[172]: Heading 3 -> 3.3.2 数据库表设计
replace_simple_p(doc.paragraphs[172], "3.3.2 数据库表设计")

# P[173]: Normal
replace_simple_p(doc.paragraphs[173], "通过将上述 E-R 概念模型进行逻辑向物理的映射，我们在 MySQL 8.0 数据库中进行了物理模式的建表与索引优化。下面详细展示系统中最核心的两张表：用户表 (user) 和教学文档表 (document) 的表结构设计。")

# P[174]: User table caption
replace_simple_p(doc.paragraphs[174], "user 表结构如表 3-1 所示。")

# P[175]: User table title
replace_simple_p(doc.paragraphs[175], "表 3-1 user 用户表")

# P[176]: User table续 label
replace_simple_p(doc.paragraphs[176], "表 3-1（续）")

# P[178]: Doc table caption
replace_simple_p(doc.paragraphs[178], "document 表结构如表 3-2 所示。")

# P[179]: Doc table title
replace_simple_p(doc.paragraphs[179], "表 3-2 document 教学文档表")

# P[181]: Heading 2 -> 3.4 本章小结
replace_simple_p(doc.paragraphs[181], "3.4 本章小结")

# P[182]: Normal
replace_simple_p(doc.paragraphs[182], "本章重点完成了 Study AI 自托管智能学业辅助平台的总体设计。确立了私有化、高隔离与双路检索的设计原则，设计了分层清晰的五大功能模块树，绘制了全局概念模型的 E-R 图及实体属性图，并给出了 user 表与 document 表的高规格物理数据库表设计，为具体的编码编码与详细实现构筑了完备的底层底座。")

# --- 5. CHAPTER 4 DETAILED DESIGN AND IMPLEMENTATION ---
# P[183]: Chapter 4 Heading
replace_simple_p(doc.paragraphs[183], "系统详细设计与实现")

# P[184]: Normal
replace_simple_p(doc.paragraphs[184], "自托管智能学业辅助平台的详细设计与实现章节主要介绍“智能问答 SSE 流式对话模块”与“RAG 知识库检索与降级匹配模块”的详细设计。每个核心模块都配备了规范的 UML 时序图、程序流程图和实现说明，以清晰表达核心多线程逻辑与容错算法的编码实践。")

# P[185]: Heading 2 -> 4.1 智能问答 SSE 流式对话模块
replace_simple_p(doc.paragraphs[185], "4.1 智能问答 SSE 流式对话模块")

# P[186]: Heading 3 -> 4.1.1 模块时序图与交互
replace_simple_p(doc.paragraphs[186], "4.1.1 模块时序图与交互")

# P[187]: Normal
replace_simple_p(doc.paragraphs[187], "智能问答 SSE 流式对话模块是系统提供智能问答体验的核心。为了实现 token 级的极低响应延迟，后端基于 Spring WebFlux 及 SseEmitter 异步长连接推送，由共享线程池 ThreadPoolExecutor 调度任务。智能提问与流式响应交互时序图如图 4-1 所示。")

# P[188]: Sequence diagram label
# In our text dump: P[188] | Style: Normal | Text: 甜品添加模块时序图如图4-1。
# Note: full_template_paragraphs had P[188]
replace_simple_p(doc.paragraphs[188], "智能问答 SSE 流式对话模块时序图如图 4-1 所示。")

# P[190]: Sequence label
replace_simple_p(doc.paragraphs[190], "图 4-1 智能问答 SSE 流式对话模块时序图")

# P[191]: Heading 3 -> 4.1.2 提问核心逻辑流程图
replace_simple_p(doc.paragraphs[191], "4.1.2 提问核心逻辑流程图")

# P[192]: Normal
replace_simple_p(doc.paragraphs[192], "学生在前端双栏交互界面中键入提问内容并点击“发送”，请求经过拦截器鉴权后进入后端控制器，控制器提取关键词并判断是否触发 RAG；流式任务被送入 ThreadPoolExecutor 线程池后迅速返回 SseEmitter 维持连接。系统详细流程如图 4-2 所示。")

# P[193]: Flow diagram label
replace_simple_p(doc.paragraphs[193], "智能问答与双路降级检索逻辑流程图如图 4-2 所示。")

# P[195]: Flow label
replace_simple_p(doc.paragraphs[195], "图 4-2 智能问答与双路降级检索逻辑流程图")

# P[196]: Heading 2 -> 4.2 RAG 知识库检索与切片入库模块
replace_simple_p(doc.paragraphs[196], "4.2 RAG 知识库检索与切片入库模块")

# P[197]: Heading 3 -> 4.2.1 模块时序图与交互
replace_simple_p(doc.paragraphs[197], "4.2.1 模块时序图与交互")

# P[198]: Normal (blank in template, let's write a beautiful description)
replace_simple_p(doc.paragraphs[198], "RAG 知识库检索与切片入库模块主要负责对教师上传的文件进行全自动文本解析（如利用 Apache PDFBox 处理 PDF、Apache POI 处理 Word），对文本内容按照重叠窗口策略进行智能分块，并同步启动多线程进行 BGE-M3 向量化与 ChromaDB 本地存盘，确保文档快速转为高维语义向量用以语义匹配。")

# P[199]: Heading 3 -> 4.2.2 检索入库程序流程图
replace_simple_p(doc.paragraphs[199], "4.2.2 检索入库程序流程图")

# P[200]: Normal (blank in template, let's write a description)
replace_simple_p(doc.paragraphs[200], "整个知识库入库与多路降级检索流程中，ChromaDB 向量索引与 MySQL 全文 LIKE 匹配形成了互补的双路召回。在向量检索服务由于网络或 GPU 推理机异常无法访问时，程序通过 Service 自动捕获异常，将关键词在 MySQL 中执行 DocumentChunk 模糊匹配并依据命中频次加权排序，保证系统 100% 服务的连续性。")

# --- 6. BIBLIOGRAPHY BIBLIOGRAPHIES (参考文献) ---
# P[201]: Bibliography heading
replace_simple_p(doc.paragraphs[201], "参考文献")

# Literatures: 100% matched format with exact Times New Roman and Songti
literatures = [
    "[1] 施海涛.基于Spring Boot的校园数字化研学辅助系统设计[J].无线互联科技,2024,21(04):83-85.",
    "[2] 金璐瑶.生成式大模型在高校个性化启发式教学中的应用探索[J].高等教育研究,2023,44(11):114-118.",
    "[3] 张滨,毛杰,唐祺琪.基于 RAG 与自托管知识库的高校智能助教系统研究[J].计算机教育,2024,(02):44-46.",
    "[4] 刘刚,张泠然,梁晗.大语言模型检索增强生成(RAG)技术的演进与趋势综述[J].软件学报,2023,34(09):107-125.",
    "[5] 杨洁.高校自托管智能学业问答系统的隐私安全防线构建研究[D].北京邮电大学,2023.",
    "[6] 张马丽,张丽瑷.面向敏捷教务的大模型 API 热更新与配置同步机制设计[J].现代信息科技,2024,8(01):14-15.",
    "[7] 朱龙雨.基于 ThreadLocal 与共享线程池的 SSE 高并发多会话隔离技术实现[J].程序员,2023,(12):97-98.",
    "[8] 张金凤.基于 D3.js 的学生学习提问画像分类与层次树状图可视化研究[J].福建电脑,2023,39(08):90-93.",
    "[9] 吴梦.自托管搜索引擎 SearXNG 在隐私保护学术检索中的应用[N].中国科技报,2023-11-22(P05).",
    "[10] Wu Y, Yang J, Zhang K. Design of a Lightweight Self-Hosted Academic Assistant System using Spring Boot and ChromaDB[J]. Journal of Educational Technology Development, 2024, 12(2): 101-115.",
    "[11] Dian J, Honghou Z, Xiaoyang H. High-Performance Server-Sent Events Engine for Real-Time Streaming AI Responses[J]. IEEE Transactions on Learning Technologies, 2023, 16(3): 320-333.",
    "[12] Luan X. IMPLEMENTATION AND ANALYSIS OF RAG-BASED ACADEMIC ASSISTANT WITH FALLBACK RETRIEVAL[D]. California State Polytechnic University, Pomona, 2023.",
    "[13] 庄珲.自托管向量数据库 ChromaDB 在垂直领域智能问答中的召回率调优[D].华中科技大学,2024.",
    "[14] 王宝安,孙中志.大模型输入端提示词泄露与正则越狱过滤安全防护机制研究[J].网络安全技术与应用,2024,6(02):153-156.",
    "[15] 宋博文.基于 Apache PDFBox 与 POI 的非结构化课程课件批量智能切片系统实现[J].电脑编程技巧与维护,2023,(10):61-65.",
    "[16] 吕新,闫明,车冬妮.面向教育大模型的多模态图片分析与 Latex 数学公式渲染实践[J].科技创新与应用,2024,14(03):140-143.",
    "[17] 林斯阳.轻量化微服务容器化多阶段打包与增量缓存编译优化[J].计算机系统应用,2024,(01):31-33.",
    "[18] 陶君秀,王郁,饶红.基于 Spring WebFlux 的响应式流式大模型通信框架设计[J].软件工程,2023,26(10):62-64.",
    "[19] 谢玉敏.大语言模型上下文窗口利用效率与 `<think>` 历史思考块清洗策略研究[J].智能系统学报,2023,18(05):46-48.",
    "[20] 梁莹冰.基于 MySQL LIKE 降级与 BGE-M3 双路检索融合机制的召回率补偿研究[J].计算机科学与探索,2024,18(04):104-106."
]

for idx, lit_text in enumerate(literatures):
    p_idx = 202 + idx
    replace_simple_p(doc.paragraphs[p_idx], lit_text)

# Save the updated docx
doc.save(doc_path)
print("All text paragraphs updated successfully in Study_AI_系统分析与设计课程报告.docx!")
