import docx

doc_path = "/private/var/folders/m1/y8wt5fvd3hl_jb3c97rwk3w80000gn/T/MicrosoftEdgeDownloads/638411b0-51e9-4d19-91b8-d7cdff285eae/23级-系统分析与设计-课程报告模板V1.3.docx"
doc = docx.Document(doc_path)

print(f"Total Paragraphs: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") or p.text.strip().startswith("第") or len(p.text.strip()) < 50 and idx < 100:
        text = p.text.strip()
        if text:
            print(f"P[{idx}] | Style: {p.style.name} | Text: {text}")

# Check references or footnotes in paragraphs
print("\n--- Searching for square brackets reference style (e.g. [1]) ---")
ref_count = 0
for idx, p in enumerate(doc.paragraphs):
    if "[" in p.text and "]" in p.text:
        ref_count += 1
        if ref_count < 20:
            print(f"P[{idx}] | {p.text.strip()[:100]}")

print(f"Total paragraphs with potential references: {ref_count}")

# Check the footnotes in the document part
try:
    footnotes_part = doc.part.package.get_part('/word/footnotes.xml')
    print("Found footnotes.xml!")
    import xml.etree.ElementTree as ET
    root = ET.fromstring(footnotes_part.blob)
    # Define namespaces
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    footnotes = root.findall('.//w:footnote', namespaces)
    print(f"Number of XML footnotes: {len(footnotes)}")
    for i, fn in enumerate(footnotes[:10]):
        fn_id = fn.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
        texts = [t.text for t in fn.findall('.//w:t', namespaces) if t.text]
        print(f"  Footnote id={fn_id}: {''.join(texts)}")
except Exception as e:
    print(f"Error accessing footnotes.xml: {e}")
