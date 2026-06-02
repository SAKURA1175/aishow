import sys
import docx

doc_path = "/private/var/folders/m1/y8wt5fvd3hl_jb3c97rwk3w80000gn/T/MicrosoftEdgeDownloads/638411b0-51e9-4d19-91b8-d7cdff285eae/23级-系统分析与设计-课程报告模板V1.3.docx"

try:
    doc = docx.Document(doc_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    # Check footnotes if possible (python-docx doesn't support footnotes out of the box easily, but we can look at document.part.related_parts or custom XML parsing)
    print("\n--- Listing first 50 paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:50]):
        text = p.text.strip()
        if text:
            print(f"[{i}] Style: {p.style.name} | Text: {text[:100]}")
            
    # Check tables
    print("\n--- Listing tables ---")
    for i, t in enumerate(doc.tables):
        print(f"Table {i}: rows={len(t.rows)}, cols={len(t.columns)}")
        # print first row cells
        row_texts = [cell.text.strip() for cell in t.rows[0].cells]
        print(f"  Row 0: {row_texts}")

except Exception as e:
    print(f"Error: {e}")
